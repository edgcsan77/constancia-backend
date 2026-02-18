# rfc.py
# -*- coding: utf-8 -*-
import os
import sys
import re
import ssl
import tempfile
import json
import jwt
import hashlib
import random
import base64
import time
import traceback
import csv
import html
import math
import hmac
import unicodedata

from zoneinfo import ZoneInfo
from io import BytesIO
from barcode import Code128
from barcode.writer import ImageWriter
from zipfile import ZipFile, ZIP_DEFLATED

import qrcode
import requests
from bs4 import BeautifulSoup
from docx import Document
from flask import Flask, request, send_file, jsonify, Response, abort
from flask_cors import CORS
from requests.adapters import HTTPAdapter
from urllib3.poolmanager import PoolManager

import secrets
from werkzeug.security import generate_password_hash, check_password_hash
from docx_to_pdf_aspose import docx_to_pdf_aspose, docx_to_pdf_aspose_web

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)
    
from stats_store import get_and_update, get_state, _now_iso
from datetime import datetime, timedelta, date

from PIL import Image
import numpy as np
import cv2
import pytesseract
import urllib.parse

from collections import deque, defaultdict
import threading

try:
    from pyzbar.pyzbar import decode as zbar_decode
    PYZBAR_OK = True
except Exception as e:
    print("pyzbar disabled:", e)
    PYZBAR_OK = False
    zbar_decode = None

from cache_store import cache_get, cache_set, cache_del
from rfc_cli_pf_solo_completo_pro import rfc_pf_13

from core_sat import consultar_curp_bot

# ===== SATPI =====
SATPI_API_KEY = (os.getenv("SATPI_API_KEY") or "").strip()
SATPI_BASE = "https://satpi.mx/api/search" 

def _rfc_only_fallback_satpi(rfc: str) -> dict:
    datos = satpi_lookup_rfc(rfc)

    # SATPI puede no traer RFC porque ya consultaste POR RFC
    if not isinstance(datos, dict) or not datos:
        raise RuntimeError("SATPI_NO_DATA")

    # 🔒 Garantiza RFC para capas superiores
    datos = dict(datos)  # copia defensiva
    datos.setdefault("RFC", rfc)
    datos.setdefault("rfc", rfc)

    return datos

def _curp_to_fecha_nac(curp: str) -> str:
    c = (curp or "").strip().upper()
    if len(c) < 10:
        return ""
    yy, mm, dd = c[4:6], c[6:8], c[8:10]
    if not (yy.isdigit() and mm.isdigit() and dd.isdigit()):
        return ""

    y2 = int(yy)
    current_y2 = datetime.now().year % 100  # 26 en 2026

    # si YY <= año actual -> 2000s, si no -> 1900s
    century = 2000 if y2 <= current_y2 else 1900
    yyyy = century + y2

    try:
        datetime(yyyy, int(mm), int(dd))
    except Exception:
        return ""

    return f"{dd}-{mm}-{yyyy:04d}"

def _fecha_ddmmaaaa_to_iso(fecha_ddmmaaaa: str) -> str:
    """
    "DD-MM-AAAA" -> "AAAA-MM-DD" (para tu rfc_pf_13)
    """
    s = (fecha_ddmmaaaa or "").strip()
    m = re.match(r"^(\d{2})-(\d{2})-(\d{4})$", s)
    if not m:
        return ""
    return f"{m.group(3)}-{m.group(2)}-{m.group(1)}"

def _split_nombre_completo(nombre: str) -> dict:
    s = (nombre or "").strip().upper()
    s = s.replace("’", "'")
    s = s.replace(".", "")            # "J." -> "J"
    s = s.replace("-", " ")
    s = re.sub(r"\s+", " ", s).strip()

    parts = [p for p in s.split(" ") if p]
    if not parts:
        return {"NOMBRE": "", "PRIMER_APELLIDO": "", "SEGUNDO_APELLIDO": ""}

    if len(parts) == 1:
        return {"NOMBRE": parts[0], "PRIMER_APELLIDO": "", "SEGUNDO_APELLIDO": ""}

    if len(parts) == 2:
        return {"NOMBRE": parts[0], "PRIMER_APELLIDO": parts[1], "SEGUNDO_APELLIDO": ""}

    # partículas que se pueden pegar al apellido (especialmente antes del núcleo)
    P = {"DE", "DEL", "LA", "LAS", "LOS", "DA", "DAS", "DO", "DOS", "DI", "DU", "VON", "VAN"}
    # conectores que a veces van dentro del apellido compuesto
    C = {"Y"}

    def _merge_lastname_tokens(tokens):
        out = []
        i = 0
        while i < len(tokens):
            t = tokens[i]

            # une secuencias de partículas: DE LA / DE LOS / DEL ...
            if t in P:
                j = i
                buf = []
                while j < len(tokens) and tokens[j] in P:
                    buf.append(tokens[j])
                    j += 1
                # si hay núcleo después, lo pegamos también
                if j < len(tokens):
                    buf.append(tokens[j])
                    j += 1
                    out.append(" ".join(buf))
                    i = j
                    continue
                else:
                    # solo partículas (raro), se apilan
                    out.append(" ".join(buf))
                    i = j
                    continue

            # une "X Y Z" como apellido compuesto
            if t not in C and (i + 2) < len(tokens) and tokens[i+1] in C and tokens[i+2] not in C:
                out.append(f"{tokens[i]} {tokens[i+1]} {tokens[i+2]}")
                i += 3
                continue

            out.append(t)
            i += 1

        # si quedó algo tipo ["DE LA CRUZ","PEREZ"] se respeta
        return out

    # 1) Por defecto: los últimos 2 "bloques" son apellidos,
    #    pero antes de eso, fusionamos tokens de apellidos con partículas.
    #    Para eso trabajamos desde el final.
    tail = parts[:]  # copia

    # Vamos a formar AP2 desde el final, respetando partículas previas.
    # Tomamos 1 token mínimo
    ap2_tokens = [tail.pop()] if tail else []
    # Si antes hay partículas (DE/LA/DEL...) las anexamos hacia atrás
    while tail and tail[-1] in P:
        ap2_tokens.insert(0, tail.pop())
    # Si antes hay "Y" como conector interno, lo jalamos con el token anterior
    if tail and tail[-1] in C and len(tail) >= 2:
        ap2_tokens.insert(0, tail.pop())      # "Y"
        ap2_tokens.insert(0, tail.pop())      # token anterior

    # AP1 igual
    ap1_tokens = [tail.pop()] if tail else []
    while tail and tail[-1] in P:
        ap1_tokens.insert(0, tail.pop())
    if tail and tail[-1] in C and len(tail) >= 2:
        ap1_tokens.insert(0, tail.pop())
        ap1_tokens.insert(0, tail.pop())

    # El resto es NOMBRES
    nombres_tokens = tail

    # Fusión extra interna (por si quedaron combos raros)
    ap1 = " ".join(_merge_lastname_tokens(ap1_tokens)).strip()
    ap2 = " ".join(_merge_lastname_tokens(ap2_tokens)).strip()
    nombres = " ".join(nombres_tokens).strip()

    def _is_only_particles(txt: str) -> bool:
        toks = [t for t in (txt or "").split() if t]
        return bool(toks) and all(t in P for t in toks)

    if _is_only_particles(ap1):
        # pásalo a nombres
        nombres = (nombres + " " + ap1).strip()
        ap1 = ""

    if _is_only_particles(ap2):
        nombres = (nombres + " " + ap2).strip()
        ap2 = ""

    return {
        "NOMBRE": nombres,
        "PRIMER_APELLIDO": ap1,
        "SEGUNDO_APELLIDO": ap2,
    }

def normalize_satpi_rfc_only(sat: dict, rfc_query: str = "") -> dict:
    sat = sat or {}
    rfc = (sat.get("rfc") or sat.get("RFC") or rfc_query or "").strip().upper()
    curp = (sat.get("curp") or sat.get("CURP") or "").strip().upper()
    cp = re.sub(r"\D+", "", (sat.get("cp") or sat.get("CP") or sat.get("codigo_postal") or "")).strip()

    # nombre
    nombre_full = (sat.get("nombre") or sat.get("NOMBRE") or "").strip()
    name_parts = _split_nombre_completo(nombre_full) if nombre_full else {"NOMBRE": "", "PRIMER_APELLIDO": "", "SEGUNDO_APELLIDO": ""}

    # régimen
    reg_desc = ""
    reg_clave = ""
    reg = sat.get("regimen")
    if isinstance(reg, list) and reg:
        reg0 = reg[0] or {}
        reg_desc = (reg0.get("descripcion") or "").strip()
        reg_clave = (reg0.get("clave") or "").strip()
    else:
        reg_desc = (sat.get("regimen_desc") or sat.get("REGIMEN") or sat.get("regimen") or sat.get("regimenFiscal") or "").strip()

    reg_desc = limpiar_regimen((reg_desc or "").strip())

    fn = _curp_to_fecha_nac(curp) if curp else ""

    datos = {
        "RFC": rfc,
        "RFC_ETIQUETA": rfc,
        "CURP": curp,
        "CP": cp,
        "REGIMEN": reg_desc,
        "regimen": reg_desc,
        "REGIMEN_CLAVE": reg_clave,
        "NOMBRE": name_parts["NOMBRE"],
        "PRIMER_APELLIDO": name_parts["PRIMER_APELLIDO"],
        "SEGUNDO_APELLIDO": name_parts["SEGUNDO_APELLIDO"],
        "FECHA_NACIMIENTO": fn,
        "_ORIGEN": "SATPI_RFC_ONLY",
    }

    if cp:
        datos["_CP_SOURCE"] = "SATPI"
    if reg_desc:
        datos["_REG_SOURCE"] = "SATPI"
    if curp:
        datos["_CURP_SOURCE"] = "SATPI"
    if nombre_full:
        datos["_NAME_SOURCE"] = "SATPI"
    if fn:
        datos["_FN_SOURCE"] = "CURP_DERIVED"

    return datos

def satpi_lookup_rfc(rfc: str) -> dict:
    rfc = (rfc or "").strip().upper()
    if len(rfc) not in (12, 13):
        raise RuntimeError("SATPI_RFC_LEN")
    if not SATPI_API_KEY:
        raise RuntimeError("SATPI_NO_APIKEY")

    url = f"{SATPI_BASE}/{rfc}"
    headers = {"x-api-key": SATPI_API_KEY}

    try:
        r = requests.get(url, headers=headers, timeout=25)
    except requests.RequestException as e:
        raise RuntimeError(f"SATPI_NET:{type(e).__name__}") from e

    try:
        js = r.json()
    except Exception:
        js = {}

    st = js.get("status") or r.status_code

    # ✅ encontrado
    if st == 200:
        nombre = str(js.get("nombre") or "").strip().upper()
        cp = str(js.get("cp") or "").strip()
        curp = str(js.get("curp") or "").strip().upper()
        reg0 = js.get("regimen") or []

        # 🔥 IMPORTANTÍSIMO:
        # si viene 200 pero sin datos, trátalo como NO ENCONTRADO / NO INSCRITO
        # (ajusta la condición si SATPI siempre trae al menos "nombre" cuando existe)
        if (not nombre) and (not curp) and (not cp) and (not reg0):
            raise RuntimeError("SATPI_NOT_FOUND")

        reg_clave = ""
        reg_desc = ""
        if isinstance(reg0, list) and reg0:
            reg_clave = str(reg0[0].get("clave") or "").strip()
            reg_desc = str(reg0[0].get("descripcion") or "").strip()

        return {
            "cp": cp,
            "regimen_clave": reg_clave,
            "regimen_desc": reg_desc,
            "curp": curp,
            "nombre": nombre,
            # ✅ opcional: regresa RFC también (útil para tus capas superiores)
            "rfc": str(js.get("rfc") or rfc).strip().upper(),
            "RFC": str(js.get("rfc") or rfc).strip().upper(),
        }

    # sin consultas
    if st == 412:
        raise RuntimeError("SATPI_412")

    # RFC inválido (formato/estructura), aunque tenga 12/13 chars
    if st == 428:
        raise RuntimeError("SATPI_428")

    raise RuntimeError(f"SATPI_BAD:{st}")

def fecha_nacimiento_from_curp(curp: str) -> str:
    """
    Extrae fecha de nacimiento desde CURP (posiciones 5-10).
    Regla de siglo:
      - 00–26  -> 2000–2026
      - 27–99  -> 1927–1999
    Retorna dd-mm-aaaa o "" si inválida.
    """
    c = (curp or "").strip().upper()

    # CURP válida básica
    if not re.fullmatch(r"[A-Z0-9]{18}", c):
        return ""

    yy = c[4:6]
    mm = c[6:8]
    dd = c[8:10]

    if not (yy.isdigit() and mm.isdigit() and dd.isdigit()):
        return ""

    yy = int(yy)
    mm = int(mm)
    dd = int(dd)

    if not (1 <= mm <= 12 and 1 <= dd <= 31):
        return ""

    # 🔐 regla de siglo (la que tú definiste)
    if 0 <= yy <= 26:
        yyyy = 2000 + yy
    else:
        yyyy = 1900 + yy

    # valida fecha real
    try:
        datetime(yyyy, mm, dd)
    except ValueError:
        return ""

    return f"{dd:02d}-{mm:02d}-{yyyy}"

# ===== GOB CURP SCRAPER (usa tu core_sat.py) =====
def gobmx_curp_scrape(term: str) -> dict:
    curp = (term or "").strip().upper()
    d = consultar_curp_bot(curp)

    print("[GOB KEYS]", sorted(list((d or {}).keys()))[:60])
    print(
        "[GOB MUN CANDIDATES]",
        d.get("MUNICIPIO_REGISTRO"),
        d.get("MUNICIPIO"),
        d.get("LOCALIDAD"),
        d.get("MUNICIPIO_NACIMIENTO"),
    )

    fn_raw = (d.get("FECHA_NACIMIENTO") or "").strip()
    fn = fn_raw.replace("/", "-").strip()

    # dd-mm-aaaa desde gob.mx
    m = re.fullmatch(r"(\d{1,2})-(\d{1,2})-(\d{4})", fn)
    if m:
        fn = f"{m.group(1).zfill(2)}-{m.group(2).zfill(2)}-{m.group(3)}"
    else:
        fn = fecha_nacimiento_from_curp(curp)

    if not fn:
        raise RuntimeError("FECHA_NACIMIENTO_INVALIDA")

    dd, mm, yyyy = fn.split("-")
    fecha_iso = f"{yyyy}-{mm}-{dd}"  # "1979-03-07"

    rfc = rfc_pf_13(
        d.get("NOMBRE", ""),
        d.get("PRIMER_APELLIDO", ""),
        d.get("SEGUNDO_APELLIDO", ""),
        fecha_iso
    )

    # ✅ FIX: prioridad correcta (REGISTRO primero)
    mun = (
        d.get("MUNICIPIO_REGISTRO") or
        d.get("MUNICIPIO") or
        d.get("LOCALIDAD") or
        d.get("MUNICIPIO_NACIMIENTO") or
        ""
    )
    mun = (mun or "").strip().upper()
    ent = (d.get("ENTIDAD_REGISTRO") or d.get("ENTIDAD") or "").strip().upper()

    mun_lock = False

    ci = {
        "RFC": rfc,
        "CURP": d.get("CURP", ""),
        "NOMBRE": d.get("NOMBRE", ""),
        "APELLIDO_PATERNO": d.get("PRIMER_APELLIDO", ""),
        "APELLIDO_MATERNO": d.get("SEGUNDO_APELLIDO", ""),
        "FECHA_NACIMIENTO": fn,
        "ENTIDAD": ent,

        "LOCALIDAD": mun,
        "MUNICIPIO": mun,

        "CP": "",
        "COLONIA": "",

        "_MUN_LOCK": mun_lock,
        "_MUN_SOURCE": "GOBMX" if mun else "",
    }

    seed_key = (ci["RFC"] or ci["CURP"] or curp).strip().upper()
    datos = build_datos_final_from_ci(ci, seed_key=seed_key)

    datos["_ORIGEN"] = "GOBMX"
    return datos

def enrich_curp_with_rfc_and_satpi(datos: dict) -> dict:
    datos = datos or {}

    rfc = (datos.get("RFC") or datos.get("rfc") or "").strip().upper()
    if not rfc:
        return datos

    try:
        sat = satpi_lookup_rfc(rfc) or {}
    except Exception as e:
        print("[SATPI_SOFT_FAIL]", type(e).__name__, str(e))
        return datos

    def put_if_str(k_dst, v):
        if isinstance(v, str):
            v = v.strip()
            if v:
                datos[k_dst] = v

    # ======================
    # CP: si NO existe o NO es válido
    # ======================
    cp_raw = (datos.get("CP") or datos.get("cp") or "").strip()
    cp_digits = re.sub(r"\D+", "", cp_raw)
    
    sat_cp_raw = str(sat.get("cp") or sat.get("CP") or "").strip()
    sat_cp_digits = re.sub(r"\D+", "", sat_cp_raw)
    
    cp_src = (datos.get("_CP_SOURCE") or "").strip().upper()

    if (len(sat_cp_digits) == 5) and (
        (len(cp_digits) != 5) or (cp_src not in ("CHECKID", "SATPI"))
    ):
        datos["CP"] = sat_cp_digits
        datos.pop("cp", None)
        datos["_CP_SOURCE"] = "SATPI"

    # ==========================
    # Régimen: solo si NO existe
    # ==========================
    reg_actual = (datos.get("REGIMEN") or datos.get("regimen") or "").strip()
    reg_desc = (sat.get("regimen_desc") or sat.get("regimen") or sat.get("REGIMEN") or "").strip()

    if (not reg_actual) and reg_desc:
        reg_clean = limpiar_regimen(reg_desc)
        if reg_clean:
            datos["REGIMEN"] = reg_clean
            datos["regimen"] = reg_clean
            datos["_REG_SOURCE"] = "SATPI"

    # (opcional) guarda clave (no afecta)
    put_if_str("REGIMEN_CLAVE", sat.get("regimen_clave"))

    # (opcional) si SATPI trae curp/nombre y faltan, rellena
    if not (datos.get("CURP") or datos.get("curp") or "").strip():
        put_if_str("CURP", sat.get("curp"))
    put_if_str("NOMBRE_SATPI", sat.get("nombre"))

    return datos

WA_PROCESSED_MSG_IDS = set()
WA_PROCESSED_QUEUE = deque(maxlen=2000)
WA_LOCK = threading.Lock()

TEST_NUMBERS = set(x.strip() for x in (os.getenv("TEST_NUMBERS", "") or "").split(",") if x.strip())
PRICE_PER_OK_MXN = int(os.getenv("PRICE_PER_OK_MXN", "0") or "0")

CURP_RE = re.compile(r"^[A-Z][AEIOUX][A-Z]{2}\d{2}(0\d|1[0-2])(0\d|[12]\d|3[01])[HM][A-Z]{5}[0-9A-Z]\d$", re.I)
RFC_RE  = re.compile(r"^([A-ZÑ&]{3,4})\d{6}([A-Z0-9]{3})$", re.I)

ADMIN_KEY = os.getenv("ADMIN_KEY", "")

GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
GITHUB_OWNER = "edgcsan77"
GITHUB_REPO = "validacion-sat"
GITHUB_BRANCH = "main"
PERSONAS_PATH = "public/data/personas.json"

PDF_CACHE_TTL_SEC = int(os.getenv("PDF_CACHE_TTL_SEC", str(7 * 24 * 3600)))   # 7 días
PDF_CACHE_MAX_BYTES = int(os.getenv("PDF_CACHE_MAX_BYTES", str(900_000)))    # ~0.9MB

PUBLIC_BASE_URL = (os.getenv("PUBLIC_BASE_URL", "") or "").strip().rstrip("/")
DL_SECRET = (os.getenv("DL_SECRET", "") or "").strip()
DL_DIR = (os.getenv("DL_DIR", "") or "/app/data/downloads").strip()
DL_TTL_SEC = int(os.getenv("DL_TTL_SEC", "86400"))
    
def _dl_ensure_dir():
    os.makedirs(DL_DIR, exist_ok=True)

def _dl_sign(token: str) -> str:
    if not DL_SECRET:
        raise RuntimeError("DL_SECRET_MISSING")
    mac = hmac.new(DL_SECRET.encode("utf-8"), token.encode("utf-8"), hashlib.sha256).hexdigest()
    return mac

def _dl_put_bytes(file_bytes: bytes, filename: str, ttl_sec: int = None) -> str:
    """
    Guarda bytes como archivo descargable y regresa URL pública.
    """
    _dl_ensure_dir()

    if not PUBLIC_BASE_URL:
        raise RuntimeError("PUBLIC_BASE_URL_MISSING")

    ttl = int(ttl_sec or DL_TTL_SEC)
    exp = int(time.time()) + max(60, ttl)

    token = secrets.token_urlsafe(16)
    sig = _dl_sign(token)

    safe_name = (filename or "archivo.bin").replace("/", "_").replace("\\", "_")
    out_path = os.path.join(DL_DIR, f"{token}__{sig}__{exp}__{safe_name}")

    # write bytes
    with open(out_path, "wb") as f:
        f.write(file_bytes)

    # URL (escapar filename para WA)
    url = f"{PUBLIC_BASE_URL}/dl/{token}/{urllib.parse.quote(safe_name)}?sig={sig}&exp={exp}"
    return url

def _dl_find_file(token: str, sig: str, exp: str, filename: str) -> str:
    """
    Busca archivo en DL_DIR por patrón token__sig__exp__filename
    """
    _dl_ensure_dir()
    safe_name = (filename or "").replace("/", "_").replace("\\", "_")
    path = os.path.join(DL_DIR, f"{token}__{sig}__{exp}__{safe_name}")
    return path

def _file_cache_key(ok_key: str, kind: str) -> str:
    return f"FILECACHE:{kind}:{(ok_key or '').strip()}"

def filecache_get_bytes(ok_key: str, kind: str) -> tuple[str | None, bytes | None, str | None]:
    """
    return (filename, bytes, mime) or (None, None, None)
    """
    rec = cache_get(_file_cache_key(ok_key, kind))
    if not isinstance(rec, dict):
        return None, None, None
    b64 = rec.get("b64")
    if not b64:
        return None, None, None
    try:
        raw = base64.b64decode(b64.encode("utf-8"), validate=True)
    except Exception:
        return None, None, None
    return rec.get("filename"), raw, rec.get("mime")

def filecache_set_bytes(ok_key: str, kind: str, filename: str, raw: bytes, mime: str):
    if not ok_key or not raw:
        return
    if len(raw) > PDF_CACHE_MAX_BYTES:
        # muy grande para guardarlo en JSON
        return
    rec = {
        "filename": filename,
        "mime": mime,
        "b64": base64.b64encode(raw).decode("utf-8"),
        "ts": _now_iso() if "_now_iso" in globals() else ""
    }
    cache_set(_file_cache_key(ok_key, kind), rec, ttl=PDF_CACHE_TTL_SEC)

def github_upsert_persona_file(d3_key: str, persona: dict, max_retries: int = 3):
    headers = {
        "Authorization": f"Bearer {GITHUB_TOKEN}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
        "User-Agent": "constancia-backend",
    }

    # carpeta por persona
    path = f"public/data/personas/{d3_key}.json"
    base_url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{path}"
    get_url = base_url + f"?ref={GITHUB_BRANCH}"
    put_url = base_url

    def _safe_json(resp):
        t = (resp.text or "")
        if not t.strip():
            raise RuntimeError(f"GH_EMPTY_BODY status={resp.status_code}")
        try:
            return resp.json()
        except Exception:
            raise RuntimeError(f"GH_NON_JSON status={resp.status_code} head={t[:220]}")

    content_b64 = base64.b64encode(
        json.dumps(persona, ensure_ascii=False, indent=2).encode("utf-8")
    ).decode("utf-8")

    last = None
    for attempt in range(1, max_retries + 1):
        try:
            r = requests.get(get_url, headers=headers, timeout=12)
            sha = None
            if r.status_code == 200:
                data = _safe_json(r)
                sha = data.get("sha")
            elif r.status_code != 404:
                raise RuntimeError(f"GH_GET_FAIL status={r.status_code} head={(r.text or '')[:220]}")

            payload = {
                "message": f"upsert persona {d3_key}",
                "content": content_b64,
                "branch": GITHUB_BRANCH
            }
            if sha:
                payload["sha"] = sha

            r2 = requests.put(put_url, headers=headers, json=payload, timeout=12)
            if r2.status_code in (200, 201):
                return True

            if r2.status_code in (409, 422) and attempt < max_retries:
                time.sleep(0.4 * attempt)
                continue

            raise RuntimeError(f"GH_PUT_FAIL status={r2.status_code} head={(r2.text or '')[:260]}")

        except Exception as e:
            last = e
            if attempt < max_retries:
                time.sleep(0.5 * attempt)
                continue
            break

    raise RuntimeError(f"GH_UPSERT_PERSONA_FAILED last={repr(last)}")

def github_update_personas(d3_key: str, persona: dict, max_retries: int = 4):
    if not (GITHUB_TOKEN and GITHUB_OWNER and GITHUB_REPO and PERSONAS_PATH):
        raise RuntimeError("GITHUB_CONFIG_MISSING")

    headers = {
        "Authorization": f"Bearer {GITHUB_TOKEN}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
        "User-Agent": "constancia-backend",
    }

    base_url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{PERSONAS_PATH}"
    get_url = base_url + f"?ref={GITHUB_BRANCH}"
    put_url = base_url  # ✅ SIN ?ref=

    def _safe_resp_json(resp: requests.Response) -> dict:
        txt = (resp.text or "")
        if not txt.strip():
            raise RuntimeError(f"GH_EMPTY_BODY status={resp.status_code}")
        try:
            return resp.json()
        except Exception:
            raise RuntimeError(f"GH_NON_JSON status={resp.status_code} head={txt[:220]}")

    last_err = None

    for attempt in range(1, max_retries + 1):
        try:
            # 1) GET actual
            r = requests.get(get_url, headers=headers, timeout=12)

            if r.status_code == 404:
                current = {}
                sha = None

            elif r.status_code == 200:
                data = _safe_resp_json(r)
                sha = data.get("sha")

                # ⛔️ GUARD RAIL: si no viene content, NO podemos asegurar merge => NO sobrescribir
                content_b64 = (data.get("content") or "").strip()
                if not content_b64:
                    raise RuntimeError("GH_EMPTY_CONTENT_REFUSING_TO_OVERWRITE")

                raw = base64.b64decode(content_b64).decode("utf-8", errors="strict").strip()
                if not raw:
                    raise RuntimeError("GH_DECODED_EMPTY_REFUSING_TO_OVERWRITE")

                current = json.loads(raw)
                if not isinstance(current, dict):
                    raise RuntimeError("PERSONAS_JSON_NOT_OBJECT_REFUSING_TO_OVERWRITE")

            else:
                raise RuntimeError(f"GH_GET_FAIL status={r.status_code} head={(r.text or '')[:260]}")

            # 2) upsert
            current[d3_key] = persona

            new_content = base64.b64encode(
                json.dumps(current, indent=2, ensure_ascii=False).encode("utf-8")
            ).decode("utf-8")

            payload = {
                "message": f"update personas.json: {d3_key}",
                "content": new_content,
                "branch": GITHUB_BRANCH
            }
            if sha:
                payload["sha"] = sha

            # 3) PUT (commit)
            r2 = requests.put(put_url, headers=headers, json=payload, timeout=12)

            if r2.status_code in (200, 201):
                return True

            # conflictos (sha viejo) -> retry
            if r2.status_code in (409, 422):
                raise RuntimeError(f"GH_PUT_CONFLICT status={r2.status_code} head={(r2.text or '')[:260]}")

            raise RuntimeError(f"GH_PUT_FAIL status={r2.status_code} head={(r2.text or '')[:260]}")

        except Exception as e:
            last_err = e
            # backoff corto
            if attempt < max_retries:
                time.sleep(0.5 * attempt)
                continue
            break

    raise RuntimeError(f"GH_UPDATE_FAILED_AFTER_RETRIES last={repr(last_err)}")

def require_admin():
    sent = request.headers.get("X-Admin-Key", "")
    if not ADMIN_KEY or sent != ADMIN_KEY:
        return False
    return True

def normalize_text(s: str) -> str:
    return (s or "").strip()

def looks_like_qr_payload(s: str) -> bool:
    s = (s or "").lower()
    # QR del validador / parámetros D1 D2 D3
    if "validadorqr" in s or "faces/pages/mobile/validadorqr.jsf" in s:
        return True
    if "d1=" in s and "d2=" in s and "d3=" in s:
        return True
    return False

def looks_like_idcif(s: str) -> bool:
    # ajusta si tu IDCIF tiene un formato específico; esto es un heurístico seguro
    s = (s or "").upper()
    return ("IDCIF" in s) or ("CIF" in s and len(s) >= 8)

def classify_input_for_personas(raw_text: str):
    """
    Regresa:
      - ("only_curp", curp) si es SOLO CURP
      - ("only_rfc", rfc)   si es SOLO RFC
      - (None, None)        si NO aplica (qr, idcif, o múltiples datos)
    """
    t = normalize_text(raw_text)

    if not t:
        return (None, None)

    # Si parece QR / link QR => NO
    if looks_like_qr_payload(t):
        return (None, None)

    # Si trae IDCIF explícito => NO
    if looks_like_idcif(t):
        return (None, None)

    # Si el texto tiene espacios o separadores típicos de combos => NO (evita CURP+RFC)
    # ejemplo: "RFC:XXXX CURP:YYYY" o "XXXX|YYYY"
    if any(sep in t for sep in [" ", "\n", "\t", "|", ",", ";"]):
        # PERO si aun así es exactamente un CURP o exactamente un RFC, sí dejamos pasar
        # (ej: el usuario pegó con espacios al inicio/fin). Para eso validamos "token único":
        tokens = [x for x in re.split(r"[\s\|\.,;]+", t) if x]
        if len(tokens) != 1:
            return (None, None)
        t = tokens[0]

    # SOLO CURP
    if CURP_RE.match(t):
        return ("only_curp", t.upper())

    # SOLO RFC
    if RFC_RE.match(t):
        return ("only_rfc", t.upper())

    return (None, None)

def is_test_request(user_key: str, text_body: str = "") -> bool:
    if user_key in TEST_NUMBERS:
        return True
    t = (text_body or "").upper()
    return ("PRUEBA" in t) or ("TEST" in t)

def wa_seen_msg(msg_id: str) -> bool:
    if not msg_id:
        return False
    with WA_LOCK:
        if msg_id in WA_PROCESSED_MSG_IDS:
            return True
        WA_PROCESSED_MSG_IDS.add(msg_id)
        WA_PROCESSED_QUEUE.append(msg_id)
        if len(WA_PROCESSED_MSG_IDS) > 2500:
            WA_PROCESSED_MSG_IDS.clear()
            WA_PROCESSED_MSG_IDS.update(list(WA_PROCESSED_QUEUE))
    return False

# ====== STATS PERSISTENTES ======
# Render Disk: monta /data (Persistent Disk) y guarda ahí.
STATS_PATH = os.getenv("STATS_PATH", "/data/stats.json")
ADMIN_STATS_TOKEN = os.getenv("ADMIN_STATS_TOKEN", "")

# Asegurar carpeta para evitar que escriba en otro lado o falle silencioso
_stats_dir = os.path.dirname(STATS_PATH)
if _stats_dir and not os.path.exists(_stats_dir):
    os.makedirs(_stats_dir, exist_ok=True)

print("STATS_PATH ->", STATS_PATH, "exists?", os.path.exists(STATS_PATH))

# ================== RATE LIMIT (por usuario) ==================
WA_USER_COOLDOWN_SEC = int(os.getenv("WA_USER_COOLDOWN_SEC", "8"))   # 8s entre mensajes
WA_USER_PER_MINUTE   = int(os.getenv("WA_USER_PER_MINUTE", "12"))    # 12 por minuto

def _rl_key(wa_id: str, suffix: str) -> str:
    return f"RL:{suffix}:{(wa_id or '').strip()}"

def wa_check_rate_limit(wa_id: str) -> tuple[bool, str]:
    """
    return (allowed, reason)
    reason: "OK", "COOLDOWN", "PER_MINUTE"
    """
    uid = (wa_id or "").strip()
    if not uid:
        return True, "OK"

    now = time.time()

    # 1) cooldown simple
    k_cd = _rl_key(uid, "CD")
    rec = cache_get(k_cd)
    if isinstance(rec, dict) and rec.get("until"):
        try:
            until = float(rec["until"])
            if now < until:
                return False, "COOLDOWN"
        except Exception:
            pass

    # set cooldown
    cache_set(k_cd, {"until": now + WA_USER_COOLDOWN_SEC}, ttl=max(WA_USER_COOLDOWN_SEC, 2))

    # 2) per-minute counter
    k_pm = _rl_key(uid, "PM")
    rec2 = cache_get(k_pm)
    if not isinstance(rec2, dict):
        rec2 = {"start": now, "count": 0}

    start = float(rec2.get("start") or now)
    count = int(rec2.get("count") or 0)

    # ventana 60s
    if now - start >= 60:
        start = now
        count = 0

    count += 1
    cache_set(k_pm, {"start": start, "count": count}, ttl=70)

    if count > WA_USER_PER_MINUTE:
        return False, "PER_MINUTE"

    return True, "OK"

# ================== ADAPTADOR TLS SAT ==================

class SATAdapter(HTTPAdapter):
    """
    Adaptador para forzar un contexto TLS que no use DH de clave pequeña.
    """
    def init_poolmanager(self, *args, **kwargs):
        ctx = ssl.create_default_context()
        ctx.set_ciphers('HIGH:!DH:!aNULL')
        kwargs['ssl_context'] = ctx
        return super().init_poolmanager(*args, **kwargs)

# ================== CONSTANTES ==================

MESES_ES = {
    1: "ENERO",
    2: "FEBRERO",
    3: "MARZO",
    4: "ABRIL",
    5: "MAYO",
    6: "JUNIO",
    7: "JULIO",
    8: "AGOSTO",
    9: "SEPTIEMBRE",
    10: "OCTUBRE",
    11: "NOVIEMBRE",
    12: "DICIEMBRE",
}

_MESES = {
    "ENERO": "01", "FEBRERO": "02", "MARZO": "03", "ABRIL": "04",
    "MAYO": "05", "JUNIO": "06", "JULIO": "07", "AGOSTO": "08",
    "SEPTIEMBRE": "09", "SETIEMBRE": "09", "OCTUBRE": "10",
    "NOVIEMBRE": "11", "DICIEMBRE": "12",
}

def _z2(n) -> str:
    try:
        return f"{int(n):02d}"
    except Exception:
        return ""

def _to_dd_mm_aaaa_dash(value: str) -> str:
    """
    Convierte a dd-mm-aaaa desde:
      - ISO: 2007-12-03 / 2007-12-03T00:00:00
      - dd/mm/aaaa
      - dd-mm-aaaa (lo deja)
      - '03 DE DICIEMBRE DE 2007'
      - '03-12-2007' (lo deja)
    Si no puede, regresa "".
    """
    if not value:
        return ""

    s = str(value).strip()
    if not s:
        return ""

    # ISO yyyy-mm-dd...
    m = re.match(r"^\s*(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        y, mm, d = m.group(1), m.group(2), m.group(3)
        return f"{d}-{mm}-{y}"

    # dd/mm/aaaa
    m = re.match(r"^\s*(\d{1,2})/(\d{1,2})/(\d{4})\s*$", s)
    if m:
        d, mm, y = _z2(m.group(1)), _z2(m.group(2)), m.group(3)
        return f"{d}-{mm}-{y}"

    # dd-mm-aaaa
    m = re.match(r"^\s*(\d{1,2})-(\d{1,2})-(\d{4})\s*$", s)
    if m:
        d, mm, y = _z2(m.group(1)), _z2(m.group(2)), m.group(3)
        return f"{d}-{mm}-{y}"

    # 'DD DE MES DE AAAA'
    up = s.upper()
    m = re.match(r"^\s*(\d{1,2})\s+DE\s+([A-ZÁÉÍÓÚÜÑ]+)\s+DE\s+(\d{4})\s*$", up)
    if m:
        d = _z2(m.group(1))
        mes_txt = m.group(2).replace("Á","A").replace("É","E").replace("Í","I").replace("Ó","O").replace("Ú","U").replace("Ü","U")
        mm = _MESES.get(mes_txt, "")
        y = m.group(3)
        if mm:
            return f"{d}-{mm}-{y}"

    return ""

def _al_from_entidad(entidad: str) -> str:
    e = (entidad or "").strip().upper()
    if not e:
        return "CIUDAD DE MÉXICO 1"  # fallback seguro
    return f"{e} 1"

def ensure_default_status_and_dates(datos: dict, seed_key: str, tz: str = "America/Mexico_City") -> dict:
    datos = datos or {}
    seed_key = (seed_key or "").strip().upper() or (datos.get("RFC") or datos.get("CURP") or "SEED").strip().upper()

    ahora = datetime.now(ZoneInfo(tz))

    def _u(x: str) -> str:
        return (x or "").strip().upper()

    # ======================
    # ESTATUS
    # ======================
    if not (datos.get("ESTATUS") or "").strip():
        datos["ESTATUS"] = "ACTIVO"

    # ======================
    # FECHA_CORTA (hoy con hora)
    # ======================
    if not (datos.get("FECHA_CORTA") or "").strip():
        datos["FECHA_CORTA"] = ahora.strftime("%Y/%m/%d %H:%M:%S")

    # ===========================
    # NUMERO EXTERIOR (RESPETA "SIN NUMERO")
    # ===========================
    no_ext_in = (datos.get("NO_EXTERIOR") or datos.get("no_exterior") or "").strip()
    num_ext_in = (datos.get("NUMERO_EXTERIOR") or datos.get("numero_exterior") or "").strip()

    raw_ext = (num_ext_in or no_ext_in).strip()
    raw_ext_u = _u(raw_ext)

    SIN_NUMERO_TOKENS = {
        "SIN NUMERO", "SIN NÚMERO", "S/N", "SN", "S N", "S. N.", "S / N",
        "NO TIENE", "NO APLICA", "N/A", "-"
    }

    if raw_ext_u in SIN_NUMERO_TOKENS:
        datos["NO_EXTERIOR"] = "SIN NUMERO"
        datos["NUMERO_EXTERIOR"] = "SIN NUMERO"
        datos["_NOEXT_INVENTED"] = False
        datos["_NOEXT_SOURCE"] = datos.get("_NOEXT_SOURCE") or (datos.get("_ORIGEN") or "SAT")
        datos["_NOEXT_LOCK"] = True
    else:
        existing_ext = re.sub(r"\D+", "", raw_ext).strip()
        if existing_ext:
            datos["NO_EXTERIOR"] = existing_ext
            datos["NUMERO_EXTERIOR"] = existing_ext
            datos["_NOEXT_INVENTED"] = False
            datos["_NOEXT_SOURCE"] = datos.get("_NOEXT_SOURCE") or "INPUT"
        else:
            ext_fake = str(_det_rand_int("NOEXT|" + seed_key, 1, 999)).strip()
            datos["NO_EXTERIOR"] = ext_fake
            datos["NUMERO_EXTERIOR"] = ext_fake
            datos["_NOEXT_INVENTED"] = True
            datos["_NOEXT_SOURCE"] = "DERIVED"

    # ======================
    # MUNICIPIO/LOCALIDAD reconcile por CP (si NO está lockeado)
    # ======================
    try:
        if not bool(datos.get("_MUN_LOCK")):
            cp = re.sub(r"\D+", "", (datos.get("CP") or "")).strip()
            if len(cp) == 5:
                meta = sepomex_by_cp(cp) or {}
                mun_cp = _u(meta.get("municipio") or "")
                ent_cp = _u(meta.get("estado") or "")

                # Solo corrige si CP te trae algo consistente
                if ent_cp and not _u(datos.get("ENTIDAD")):
                    datos["ENTIDAD"] = ent_cp

                # Si municipio del CP existe, úsalo para evitar mezcla "La Piedad" con CP Guadalajara
                if mun_cp:
                    mun_cur = _u(datos.get("MUNICIPIO") or "")
                    loc_cur = _u(datos.get("LOCALIDAD") or "")
                    # Si falta o si está distinto, preferimos el del CP (cuando no hay lock)
                    if (not mun_cur) or (mun_cur != mun_cp):
                        datos["MUNICIPIO"] = mun_cp
                        datos["_MUN_SOURCE"] = datos.get("_MUN_SOURCE") or "SEPOMEX_BY_CP"
                    if (not loc_cur) or (loc_cur != mun_cp):
                        datos["LOCALIDAD"] = mun_cp
    except Exception:
        pass

    # ======================
    # FECHA (lugar y fecha de emisión)
    # ======================
    if not (datos.get("FECHA") or "").strip():
        mun = _u(datos.get("LOCALIDAD") or datos.get("MUNICIPIO") or "")
        ent = _u(datos.get("ENTIDAD") or "")
        if mun and ent:
            try:
                datos["FECHA"] = _fecha_lugar_mun_ent(mun, ent)
            except Exception:
                pass

    # ======================
    # Fechas derivadas (si faltan)
    # ======================
    need_any = any(
        not (datos.get(k) or "").strip()
        for k in ("FECHA_INICIO", "FECHA_ULTIMO", "FECHA_ALTA", "FECHA_INICIO_DOC", "FECHA_ULTIMO_DOC", "FECHA_ALTA_DOC")
    )

    if need_any:
        fn = (datos.get("FECHA_NACIMIENTO") or "").strip()
        birth_year = _parse_birth_year(fn)
        y0 = (birth_year + 18) if birth_year else (ahora.year - 5)

        d, m, y = _fake_date_components(y0, seed_key)

        fecha_inicio_raw = _fmt_dd_de_mes_de_aaaa(d, m, y)
        fecha_ultimo_raw = _fmt_dd_de_mes_de_aaaa(d, m, y)
        fecha_alta_raw   = _fmt_dd_mm_aaaa(d, m, y)

        fi_dash = _to_dd_mm_aaaa_dash(fecha_inicio_raw)
        fu_dash = _to_dd_mm_aaaa_dash(fecha_ultimo_raw)
        fa_dash = _to_dd_mm_aaaa_dash(fecha_alta_raw)

        if not fi_dash:
            fi_dash = _to_dd_mm_aaaa_dash(fn) or ""
        if not fu_dash:
            fu_dash = fi_dash
        if not fa_dash:
            fa_dash = fi_dash

        if not (datos.get("FECHA_INICIO") or "").strip() and fi_dash:
            datos["FECHA_INICIO"] = fi_dash
        if not (datos.get("FECHA_ULTIMO") or "").strip() and fu_dash:
            datos["FECHA_ULTIMO"] = fu_dash
        if not (datos.get("FECHA_ALTA") or "").strip() and fa_dash:
            datos["FECHA_ALTA"] = fa_dash

        if not (datos.get("FECHA_INICIO_DOC") or "").strip():
            datos["FECHA_INICIO_DOC"] = fecha_inicio_raw
        if not (datos.get("FECHA_ULTIMO_DOC") or "").strip():
            datos["FECHA_ULTIMO_DOC"] = fecha_ultimo_raw
        if not (datos.get("FECHA_ALTA_DOC") or "").strip():
            datos["FECHA_ALTA_DOC"] = fecha_alta_raw

    # ======================
    # ✅ Address defaults + alias (NUMERO_EXTERIOR SIEMPRE)
    # ======================

    # Alias de NO_EXTERIOR/NO_INTERIOR a NUMERO_EXTERIOR/NUMERO_INTERIOR
    no_ext = re.sub(r"\D+", "", (datos.get("NO_EXTERIOR") or datos.get("no_exterior") or "")).strip()
    no_int = re.sub(r"\D+", "", (datos.get("NO_INTERIOR") or datos.get("no_interior") or "")).strip()

    if no_ext and not (datos.get("NUMERO_EXTERIOR") or "").strip():
        datos["NUMERO_EXTERIOR"] = no_ext
    if no_int and not (datos.get("NUMERO_INTERIOR") or "").strip():
        datos["NUMERO_INTERIOR"] = no_int

    # Si sigue faltando NUMERO_EXTERIOR, inventa determinístico
    if (not (datos.get("NUMERO_EXTERIOR") or "").strip()) and (not bool(datos.get("_NOEXT_LOCK"))):
        try:
            if "_det_rand_int" in globals() or "_det_rand_int" in locals():
                n = _det_rand_int(f"NOEXT|{seed_key}", 1, 999)
            else:
                n = (abs(hash(f"NOEXT|{seed_key}")) % 999) + 1
            datos["NUMERO_EXTERIOR"] = str(int(n))
            datos["_NOEXT_INVENTED"] = True
        except Exception:
            pass

    return datos

def _parse_birth_year(fecha_nac: str) -> int | None:
    """
    Acepta:
      - '1978-09-20T00:00:00'
      - '1978-09-20'
      - '20/09/1978'
      - '20-09-1978'   ✅ ESTE ERA EL BUG
    """
    if not fecha_nac:
        return None

    s = str(fecha_nac).strip()

    try:
        # ISO: YYYY-MM-DD o YYYY-MM-DDTHH:MM:SS
        m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", s)
        if m:
            return int(m.group(1))

        # DD/MM/YYYY
        m = re.match(r"^(\d{2})/(\d{2})/(\d{4})$", s)
        if m:
            return int(m.group(3))

        # DD-MM-YYYY  ←🔥 ESTE FALTABA
        m = re.match(r"^(\d{2})-(\d{2})-(\d{4})$", s)
        if m:
            return int(m.group(3))

    except Exception:
        return None

    return None

def _det_rng(seed_text: str) -> random.Random:
    """
    Random determinístico para que NO cambie entre reintentos.
    """
    h = hashlib.sha256((seed_text or "").encode("utf-8")).hexdigest()
    seed_int = int(h[:16], 16)  # 64 bits
    return random.Random(seed_int)

def _fake_date_dd_de_mmmm_de_aaaa(year: int, seed_key: str, salt: str) -> str:
    """
    Día 1-30, mes 1-12, año fijo.
    """
    rng = _det_rng(f"{seed_key}|{salt}")
    day = rng.randint(1, 30)
    month = rng.randint(1, 12)
    mes = MESES_ES.get(month, str(month))
    return f"{day:02d} DE {mes} DE {year}"

def _fake_date_dd_mm_yyyy(year: int, seed_key: str, salt: str) -> str:
    """
    Para FECHA_ALTA si la quieres en formato  dd/mm/yyyy.
    """
    rng = _det_rng(f"{seed_key}|{salt}|alta")
    day = rng.randint(1, 30)
    month = rng.randint(1, 12)
    return f"{day:02d}/{month:02d}/{year}"

# ================== USUARIOS / SESIONES / IP / LÍMITES ==================
# CAMBIA ESTO por tus usuarios reales
USERS = {
    "admin": generate_password_hash("Loc0722E02?"),
    "test": generate_password_hash("TestIDCIF26"),
    #"graciela.barajas": generate_password_hash("BarajasCIF26"),
    "gerardo.calzada": generate_password_hash("CalzadaIDCIF26"),
    "gerardo.calzada.oficina": generate_password_hash("CalzadaIDCIF26"),
    #"tramites.monterrey": generate_password_hash("MonterreyCIF26"),
    #"monterrey.oficina1": generate_password_hash("MonterreyCIF26"),
    #"monterrey.oficina2": generate_password_hash("MonterreyCIF26"),
    "juan.gutierrez":generate_password_hash("GutierrezIDCIF26"),
    "angel.chavez":generate_password_hash("ChavezIDCIF26"),
    "daniel.gonzalez":generate_password_hash("GonzalezCIF26"),
    #"eos":generate_password_hash("EOScif26"),
    "omar.perez":generate_password_hash("PerezCIF26"),
    "omar.perez2":generate_password_hash("PerezCIF26"),
    "brandon.user":generate_password_hash("BrandonCIF26"),
    "alejandro.user":generate_password_hash("AlejandroIDCIF26"),
    "alejandro.user2":generate_password_hash("AlejandroIDCIF26"),
    "mariano.gonzalez":generate_password_hash("MarianoIDCIF26"),
    "gabriel.tavarez":generate_password_hash("TavarezCIF26"),
    "david.romero":generate_password_hash("RomeroCIF26"),
}

# Historial de logins por usuario
# {"usuario": [{"ip": "...", "fecha": "...", "ua": "..."}]}
HISTORIAL_LOGIN = {}

# Info de IP por usuario, para poder fijar una IP si quieres
# {"usuario": {"ip": "...", "bloquear_otras": True/False}}
USERS_IP_INFO = {}

# Si True, la primera IP que use el usuario se fija y se bloquean otras
BLOQUEAR_IP_POR_DEFAULT = False  # déjalo False mientras solo observas

# Límite diario de constancias por usuario
USO_POR_USUARIO = {}  # {"usuario": {"hoy": "2025-12-31", "count": 3}}
LIMITE_DIARIO = 1000    # cambia este número según tu plan

# ================== FUNCIONES AUXILIARES ==================

def hoy_mexico():
    try:
        return datetime.now(ZoneInfo("America/Mexico_City")).date()
    except Exception:
        return datetime.utcnow().date()

def formatear_fecha_dd_de_mmmm_de_aaaa(d_str, sep="-"):
    if not d_str:
        return ""
    partes = d_str.strip().split(sep)
    if len(partes) != 3:
        return d_str
    dd, mm, yyyy = partes
    try:
        dia = int(dd)
        mes = int(mm)
        anio = int(yyyy)
    except ValueError:
        return d_str
    nombre_mes = MESES_ES.get(mes, mm)
    return f"{dia:02d} DE {nombre_mes} DE {anio}"

def _fecha_lugar_mun_ent(
    municipio: str,
    entidad: str,
    *,
    year: int | None = None,
    month: int | None = None,
    day: int | None = None,
    d: date | None = None
) -> str:
    """
    Devuelve: 'MUNICIPIO , ENTIDAD A 29 DE ENERO DE 2026'

    Prioridad:
    1) Si viene d= (date), usa esa fecha completa.
    2) Si no, parte de hoy_mexico() y permite override de year/month/day.
    """
    base = d or hoy_mexico()

    y = int(year) if year is not None else base.year
    m = int(month) if month is not None else base.month
    dd = int(day) if day is not None else base.day

    # valida fecha (por si pasan 31 en febrero, etc.)
    try:
        final = date(y, m, dd)
    except Exception:
        # fallback: usa la base si se pasó algo inválido
        final = base

    fecha = f"{final.day:02d} DE {MESES_ES[final.month]} DE {final.year}"

    mun = (municipio or "").strip().upper()
    ent = (entidad or "").strip().upper()

    if mun and ent:
        return f"{mun} , {ent} A {fecha}"
    if mun:
        return f"{mun} A {fecha}"
    if ent:
        return f"{ent} A {fecha}"
    return fecha
    
def fecha_actual_lugar(localidad, entidad):
    hoy = hoy_mexico()
    dia = str(hoy.day).zfill(2)
    mes = MESES_ES[hoy.month]
    anio = hoy.year

    loc = (localidad or "").upper()
    ent = (entidad or "").upper()

    if not loc and not ent:
        lugar = ""
    elif loc and ent:
        lugar = f"{loc} , {ent} A "
    elif loc:
        lugar = f"{loc} A "
    else:
        lugar = f"{ent} A "

    return f"{lugar}{dia} DE {mes} DE {anio}"

def get_client_ip():
    """
    Intenta obtener la IP real del cliente, tomando en cuenta proxies (Render).
    """
    if "X-Forwarded-For" in request.headers:
        return request.headers["X-Forwarded-For"].split(",")[0].strip()
    return request.remote_addr or "0.0.0.0"

def _barcode_local_code128(rfc: str) -> bytes:
    """Barcode Code128 LOCAL, con texto abajo (SAT-like)."""
    buf = BytesIO()
    bc = Code128(rfc, writer=ImageWriter())
    bc.write(
        buf,
        options={
            "module_width": 0.32,
            "module_height": 12.0,
            "quiet_zone": 1.8,
            "font_size": 11,
            "text_distance": 4.4,
            "write_text": True,   # 👈 texto abajo
            "dpi": 300,
        }
    )
    return buf.getvalue()

def generar_qr_y_barcode(url_qr, rfc):
    # ---------- QR (SIEMPRE) ----------
    qr = qrcode.QRCode(
        version=None,
        box_size=8,
        border=2,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
    )
    qr.add_data(url_qr)
    qr.make(fit=True)

    qr_img = qr.make_image(fill_color="black", back_color="white")
    buf_qr = BytesIO()
    qr_img.save(buf_qr, format="PNG")
    qr_bytes = buf_qr.getvalue()

    # ---------- BARCODE ----------
    rfc_clean = (rfc or "").strip().upper()
    if not rfc_clean:
        return qr_bytes, None

    # 1) Cache (si existe)
    cache_key = f"BARCODE_TECIT:{rfc_clean}"
    try:
        cached = cache_get(cache_key)
        if cached:
            return qr_bytes, cached
    except Exception:
        pass

    # 2) Intentar TEC-IT (soft)
    rfc_encoded = urllib.parse.quote_plus(rfc_clean)
    url_barcode = (
        "https://barcode.tec-it.com/barcode.ashx"
        f"?data={rfc_encoded}"
        "&code=Code128"
        "&translate-esc=on"
        "&dpi=300"
    )

    TIMEOUT_SEC = 8
    MAX_ATTEMPTS = 2

    for attempt in range(1, MAX_ATTEMPTS + 1):
        try:
            resp = requests.get(url_barcode, timeout=TIMEOUT_SEC)
            if resp.ok and resp.content:
                try:
                    cache_set(cache_key, resp.content, ttl=60 * 60 * 24 * 7)
                except Exception:
                    pass
                return qr_bytes, resp.content
        except Exception as e:
            print("BARCODE TEC-IT FAIL (soft):", repr(e), "attempt", attempt)
            if attempt < MAX_ATTEMPTS:
                time.sleep(0.8)

    # 3) 🔥 FALLBACK LOCAL (NUNCA FALLA)
    try:
        barcode_local = _barcode_local_code128(rfc_clean)
        return qr_bytes, barcode_local
    except Exception as e:
        print("BARCODE LOCAL FAIL (very rare):", repr(e))
        return qr_bytes, None

D26_FOLIO_MIN = 300_000_000
D26_FOLIO_MAX = 399_999_999

D26_DOCNAME_CONST = "CONSTANCIA DE SITUACIÓN FISCAL"
D26_ID_CONST = "200001088888800000041"
D26_TOKEN_CONST = "U2FsdGVkX1/u6lyj56lir/HqUsYBpXK66xpeFKg5Qymp/ecS4Xweh/Iv+uVKzCMN"

def _d26_folio_deterministico(seed_key: str) -> int:
    """
    Folio determinístico por RFC (estable entre reintentos).
    Usa tu _det_rng existente.
    """
    rng = _det_rng(f"D26|{(seed_key or '').strip().upper()}")
    return int(rng.randint(D26_FOLIO_MIN, D26_FOLIO_MAX))

def _build_cadena_original_d26(fecha_corta: str, rfc: str) -> str:
    # Formato exacto solicitado
    fc = (fecha_corta or "").strip()
    rf = (rfc or "").strip().upper()
    return f"||{fc}|{rf}|{D26_DOCNAME_CONST}|{D26_ID_CONST}|{D26_TOKEN_CONST}||"

def _persona_d26_min(datos: dict, d3_key: str, rfc: str) -> dict:
    datos = datos or {}
    fecha_corta = (datos.get("FECHA_CORTA") or "").strip()  # ya lo generas como YYYY/MM/DD HH:MM:SS
    cadena = _build_cadena_original_d26(fecha_corta, rfc)

    return {
        "D1": "26",
        "D2": "1",
        "D3": d3_key,
        "rfc": (rfc or "").strip().upper(),
        "cadena_original": cadena
    }

def generar_solo_qr_png(url_qr: str) -> bytes:
    qr = qrcode.QRCode(
        version=None,
        box_size=8,
        border=2,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
    )
    qr.add_data(url_qr)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()

def obtener_mapa_trs(soup):
    filas = {}
    for tr in soup.find_all("tr"):
        tds = tr.find_all("td")
        if len(tds) >= 2:
            etiqueta = tds[0].get_text(strip=True)
            valor = tds[1].get_text(strip=True)
            if etiqueta:
                filas[etiqueta] = valor
    return filas

def construir_datos_manual(payload: dict, *, input_type: str = "MANUAL") -> dict:
    """
    payload: dict con tus datos ya completos (lo que tú capturas).
    Regresa el dict 'datos' con las mismas llaves que usa reemplazar_en_documento().
    NO consulta APIs.
    """
    p = payload or {}

    def U(x): return (x or "").strip().upper()
    def S(x): return (x or "").strip()
    def D(x): return _to_dd_mm_aaaa_dash(S(x))  # ya existe en tu archivo
    def DIG(x): return re.sub(r"\D+", "", S(x))

    rfc   = U(p.get("RFC") or p.get("rfc"))
    curp  = U(p.get("CURP") or p.get("curp"))
    nombre = U(p.get("NOMBRE") or p.get("nombre"))
    ap1    = U(p.get("PRIMER_APELLIDO") or p.get("apellido_paterno"))
    ap2    = U(p.get("SEGUNDO_APELLIDO") or p.get("apellido_materno"))

    nombre_etiqueta = " ".join(x for x in [nombre, ap1, ap2] if x).strip()

    # Fechas (acepta ISO / dd-mm-aaaa / dd/mm/aaaa / "03 DE ENERO DE 2026")
    fn = D(p.get("FECHA_NACIMIENTO") or p.get("fecha_nacimiento"))
    fi = D(p.get("FECHA_INICIO") or p.get("fecha_inicio_operaciones") or p.get("fecha_inicio"))
    fu = D(p.get("FECHA_ULTIMO") or p.get("fecha_ultimo_cambio") or p.get("fecha_ultimo"))
    fa = D(p.get("FECHA_ALTA") or p.get("fecha_alta"))

    entidad   = U(p.get("ENTIDAD") or p.get("entidad"))
    municipio = U(p.get("LOCALIDAD") or p.get("municipio"))
    colonia   = U(p.get("COLONIA") or p.get("colonia"))

    cp = DIG(p.get("CP") or p.get("cp"))

    # Dirección (si quieres “modo CURP/RFC_ONLY” fijo, deja defaults)
    tipo_vialidad = U(p.get("TIPO_VIALIDAD") or p.get("tipo_vialidad") or "CALLE") or "CALLE"
    vialidad      = U(p.get("VIALIDAD") or p.get("nombre_vialidad") or "SIN NOMBRE") or "SIN NOMBRE"
    no_ext        = DIG(p.get("NO_EXTERIOR") or p.get("numero_exterior"))
    no_int        = DIG(p.get("NO_INTERIOR") or p.get("numero_interior"))

    # Régimen / estatus
    regimen = limpiar_regimen(p.get("REGIMEN") or p.get("regimen") or "")
    estatus = U(p.get("ESTATUS") or p.get("situacion_contribuyente") or "ACTIVO") or "ACTIVO"

    # IDCIF (si tú lo vas a dar, pásalo; si no, puedes generar uno “fakey” determinístico)
    idcif = S(p.get("IDCIF") or p.get("IDCIF_ETIQUETA") or p.get("idcif") or "")
    if not idcif and rfc:
        # 11 dígitos determinísticos (para que QR/etiquetas sean estables)
        idcif = str(_det_rand_int(f"IDCIF|{rfc}", 10**10, 10**11 - 1))

    # FECHA / FECHA_CORTA como lo hace SAT
    ahora = datetime.now(ZoneInfo("America/Mexico_City"))
    fecha_corta = ahora.strftime("%Y/%m/%d %H:%M:%S")
    fecha_emision = _fecha_lugar_mun_ent(municipio, entidad)  # ya existe en tu archivo

    datos = {
        "RFC_ETIQUETA": rfc,
        "NOMBRE_ETIQUETA": nombre_etiqueta,
        "IDCIF_ETIQUETA": idcif,

        "RFC": rfc,
        "CURP": curp,
        "NOMBRE": nombre,
        "PRIMER_APELLIDO": ap1,
        "SEGUNDO_APELLIDO": ap2,

        # dd-mm-aaaa (para consistencia con tu normalizador)
        "FECHA_NACIMIENTO": fn,
        "FECHA_INICIO": fi,
        "FECHA_ULTIMO": fu,
        "FECHA_ALTA": fa,

        # versiones "DOC" (si tu plantilla las usa en texto largo)
        "FECHA_INICIO_DOC": formatear_fecha_dd_de_mmmm_de_aaaa(fi, sep="-") if fi else "",
        "FECHA_ULTIMO_DOC": formatear_fecha_dd_de_mmmm_de_aaaa(fu, sep="-") if fu else "",
        "FECHA_ALTA_DOC": fa.replace("-", "/") if fa else "",

        "ESTATUS": estatus,
        "REGIMEN": regimen,

        "CP": cp,
        "TIPO_VIALIDAD": tipo_vialidad,
        "VIALIDAD": vialidad,
        "NO_EXTERIOR": no_ext,
        "NO_INTERIOR": no_int,
        "COLONIA": colonia,
        "LOCALIDAD": municipio,
        "ENTIDAD": entidad,

        "FECHA": fecha_emision,
        "FECHA_CORTA": fecha_corta,

        # requerido por ti en otros flujos
        "AL": _al_from_entidad(entidad),
        "CORREO": S(p.get("CORREO") or p.get("correo")),
    }

    return datos

def extraer_datos_desde_sat(rfc, idcif):
    d3 = f"{idcif}_{rfc}"

    url = "https://siat.sat.gob.mx/app/qr/faces/pages/mobile/validadorqr.jsf"
    params = {"D1": "10", "D2": "1", "D3": d3}
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}

    session = requests.Session()
    session.mount("https://siat.sat.gob.mx", SATAdapter())

    resp = session.get(url, params=params, headers=headers, timeout=20)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    mapa = obtener_mapa_trs(soup)

    def get_val(*keys_posibles):
        for k in keys_posibles:
            if k in mapa:
                return mapa[k]
        return ""

    # ===== Persona física =====
    nombre = get_val("Nombre:", "Nombre (s):")
    ape1 = get_val("Apellido Paterno:", "Primer Apellido:")
    ape2 = get_val("Apellido Materno:", "Segundo Apellido:")
    nombre_etiqueta_pf = " ".join(x for x in [nombre, ape1, ape2] if x).strip()

    # ===== Persona moral =====
    denominacion = get_val("Denominación o Razón Social:")
    capital = get_val("Régimen de capital:")

    # Fechas
    fecha_inicio_raw = get_val("Fecha de Inicio de operaciones:", "Fecha inicio de operaciones:")
    fecha_ultimo_raw = get_val("Fecha del último cambio de situación:", "Fecha de último cambio de estado:")
    fecha_inicio_texto = formatear_fecha_dd_de_mmmm_de_aaaa(fecha_inicio_raw, sep="-")
    fecha_ultimo_texto = formatear_fecha_dd_de_mmmm_de_aaaa(fecha_ultimo_raw, sep="-")

    estatus = get_val("Situación del contribuyente:", "Estatus en el padrón:")
    curp = get_val("CURP:")

    # Domicilio
    cp = get_val("CP:", "Código Postal:")
    tipo_vialidad = get_val("Tipo de vialidad:")
    vialidad = get_val("Nombre de la vialidad:")
    no_ext = get_val("Número exterior:")
    no_int = get_val("Número interior:")
    colonia = get_val("Colonia:", "Nombre de la Colonia:")
    localidad = get_val("Municipio o delegación:", "Nombre del Municipio o Demarcación Territorial:")
    entidad = get_val("Entidad Federativa:", "Nombre de la Entidad Federativa:")

    regimen = get_val("Régimen:")
    fecha_alta_raw = get_val("Fecha de alta:")
    fecha_alta = fecha_alta_raw.replace("-", "/") if fecha_alta_raw else ""

    if not any([denominacion, nombre, ape1, ape2, curp, cp, regimen]):
        raise ValueError("SIN_DATOS_SAT")

    fecha_actual = fecha_actual_lugar(localidad, entidad)

    ahora = datetime.now(ZoneInfo("America/Mexico_City"))
    fecha_corta = ahora.strftime("%Y/%m/%d %H:%M:%S")

    nombre_etiqueta = (denominacion or nombre_etiqueta_pf).strip()

    datos = {
        "RFC_ETIQUETA": rfc,
        "NOMBRE_ETIQUETA": nombre_etiqueta,
        "IDCIF_ETIQUETA": idcif,

        "RFC": rfc,
        "CURP": curp,

        "DENOMINACION": denominacion,
        "CAPITAL": capital,
        
        "NOMBRE": (denominacion or nombre),
        
        "PRIMER_APELLIDO": ape1,
        "SEGUNDO_APELLIDO": ape2,
        
        "FECHA_INICIO_DOC": fecha_inicio_texto,
        "FECHA_ULTIMO_DOC": fecha_ultimo_texto,

        "ESTATUS": estatus,
        "FECHA": fecha_actual,
        "FECHA_CORTA": fecha_corta,

        "CP": cp,
        "TIPO_VIALIDAD": tipo_vialidad,
        "VIALIDAD": vialidad,
        "NO_EXTERIOR": no_ext,
        "NO_INTERIOR": no_int,
        "COLONIA": colonia,
        "LOCALIDAD": localidad,
        "ENTIDAD": entidad,

        "REGIMEN": regimen,
        "FECHA_ALTA_DOC": fecha_alta,
    }

    return datos

# ================== NUEVO: PUBLICAR DATOS EN VALIDACION-SAT ==================

VALIDACION_SAT_BASE = (os.getenv("VALIDACION_SAT_BASE", "") or "").rstrip("/")
VALIDACION_SAT_APIKEY = (os.getenv("VALIDACION_SAT_APIKEY", "") or "").strip()
VALIDACION_SAT_TIMEOUT = int(os.getenv("VALIDACION_SAT_TIMEOUT", "8") or "8")

def validacion_sat_enabled() -> bool:
    return bool(VALIDACION_SAT_BASE and VALIDACION_SAT_APIKEY)

def datos_to_persona_sat(datos: dict, d3: str, idcif: str, rfc: str, curp: str) -> dict:
    def U(x): return (x or "").strip().upper()
    def S(x): return (x or "").strip()

    nombre = U(datos.get("NOMBRE") or datos.get("nombre"))
    ap1 = U(datos.get("PRIMER_APELLIDO") or datos.get("apellido_paterno"))
    ap2 = U(datos.get("SEGUNDO_APELLIDO") or datos.get("apellido_materno"))

    # Fechas: tu backend ya las trae dd-mm-aaaa normalmente
    fn = S(datos.get("FECHA_NACIMIENTO") or datos.get("fecha_nacimiento"))
    fi = S(datos.get("FECHA_INICIO") or datos.get("fecha_inicio_operaciones"))
    fu = S(datos.get("FECHA_ULTIMO") or datos.get("fecha_ultimo_cambio"))
    fa = S(datos.get("FECHA_ALTA") or datos.get("fecha_alta"))

    entidad = U(datos.get("ENTIDAD") or datos.get("entidad"))
    municipio = U(datos.get("LOCALIDAD") or datos.get("municipio"))
    colonia = U(datos.get("COLONIA") or datos.get("colonia"))

    tipo_v = U(datos.get("TIPO_VIALIDAD") or datos.get("tipo_vialidad"))
    vial = U(datos.get("VIALIDAD") or datos.get("nombre_vialidad"))

    no_ext = S(datos.get("NO_EXTERIOR") or datos.get("numero_exterior"))
    no_int = S(datos.get("NO_INTERIOR") or datos.get("numero_interior"))

    cp = S(datos.get("CP") or datos.get("cp"))
    correo = S(datos.get("CORREO") or datos.get("correo"))

    estatus = U(datos.get("ESTATUS") or datos.get("situacion_contribuyente"))
    regimen = S(datos.get("REGIMEN") or datos.get("regimen"))

    nombre_etiqueta = U(datos.get("NOMBRE_ETIQUETA") or f"{nombre} {ap1} {ap2}".strip())
    al = S(datos.get("AL") or datos.get("al") or "")

    return {
        "D1": "10",
        "D2": "1",
        "D3": d3,

        "rfc": U(rfc),
        "curp": U(curp),

        "nombre": nombre,
        "apellido_paterno": ap1,
        "apellido_materno": ap2,

        "fecha_nacimiento": fn,
        "fecha_inicio_operaciones": fi,
        "situacion_contribuyente": estatus,
        "fecha_ultimo_cambio": fu,
        "regimen": regimen,
        "fecha_alta": fa,

        "entidad": entidad,
        "municipio": municipio,
        "colonia": colonia,

        "tipo_vialidad": tipo_v,
        "nombre_vialidad": vial,
        "numero_exterior": no_ext,
        "numero_interior": no_int,

        "cp": cp,
        "correo": correo,
        "al": al,

        "RFC_ETIQUETA": U(rfc),
        "NOMBRE_ETIQUETA": nombre_etiqueta,
        "IDCIF_ETIQUETA": S(idcif),
    }

def validacion_sat_publish(datos: dict, input_type: str) -> str | None:
    if not validacion_sat_enabled():
        return None

    rfc = (datos.get("RFC") or datos.get("rfc") or "").strip().upper()
    curp = (datos.get("CURP") or datos.get("curp") or "").strip().upper()
    idcif = (datos.get("IDCIF_ETIQUETA") or datos.get("IDCIF") or "").strip()

    if not (rfc and idcif):
        return None

    d3 = f"{idcif}_{rfc}"

    # ✅ objeto en formato SAT EXACTO
    persona = datos_to_persona_sat(datos, d3=d3, idcif=idcif, rfc=rfc, curp=curp)

    # ✅ publicar de verdad (commit a personas.json)
    github_update_personas(d3, persona)

    # ✅ URL de tu visor (ajusta si tu visor usa otra ruta)
    if VALIDACION_SAT_BASE:
        return f"{VALIDACION_SAT_BASE}/v?D1=10&D2=1&D3={urllib.parse.quote(d3)}"

    return None

def elegir_url_qr(datos: dict, input_type: str, rfc_val: str, idcif_val: str) -> str:
    input_type = (input_type or "").upper().strip()
    rfc_val = (rfc_val or "").strip().upper()
    idcif_val = (idcif_val or "").strip()

    # ✅ 0) Para SOLO CURP / SOLO RFC: SIEMPRE usar validadorqr.jsf con D3 = IDCIF_RFC
    if input_type in ("CURP", "RFC_ONLY", "MANUAL") and VALIDACION_SAT_BASE and idcif_val and rfc_val:

        d3 = f"{idcif_val}_{rfc_val}"
        return (
            f"{VALIDACION_SAT_BASE}/app/qr/faces/pages/mobile/validadorqr.jsf"
            f"?D1=10&D2=1&D3={urllib.parse.quote(d3)}"
        )

    # 1) Si ya se publicó en validacion-sat (otros casos), úsalo
    qr_url_pub = (datos.get("QR_URL") or "").strip()
    if qr_url_pub:
        return qr_url_pub

    # 2) QR oficial SAT SOLO cuando sea RFC_IDCIF y haya idCIF real
    if input_type == "RFC_IDCIF" and idcif_val:
        d3 = f"{idcif_val}_{rfc_val}"
        return (
            "https://siat.sat.gob.mx/app/qr/faces/pages/mobile/validadorqr.jsf"
            f"?D1=10&D2=1&D3={d3}"
        )

    # 3) Fallback seguro
    if VALIDACION_SAT_BASE:
        return f"{VALIDACION_SAT_BASE}/v?rfc={urllib.parse.quote_plus(rfc_val)}"

    return "https://siat.sat.validacion-sat.org"

def extraer_lista_rfc_idcif(text_body: str) -> list[tuple[str, str]]:
    """
    Lee un texto con muchas líneas tipo:
      RFC IDCIF
      RFC IDCIF
    y regresa una lista de (rfc, idcif).

    Acepta:
    - separador por espacios o tab
    - comas
    - guiones
    """
    if not text_body:
        return []

    pares = []
    seen = set()

    for raw in text_body.splitlines():
        line = (raw or "").strip().upper()
        if not line:
            continue

        # normaliza separadores
        line = line.replace(",", " ").replace("|", " ").replace(";", " ")
        line = re.sub(r"\s+", " ", line).strip()

        # intenta sacar RFC e IDCIF
        parts = line.split(" ")
        if len(parts) < 2:
            continue

        rfc = parts[0].strip()
        idcif = parts[1].strip()

        # validar longitudes mínimas
        if len(rfc) not in (12, 13):
            continue
        if len(idcif) != 11:
            continue

        # validar formato (si ya tienes is_valid_rfc)
        if not is_valid_rfc(rfc):
            continue
        if not idcif.isdigit():
            continue

        key = f"{rfc}|{idcif}"
        if key in seen:
            continue
        seen.add(key)

        pares.append((rfc, idcif))

    return pares

def parece_lista_rfc_idcif(text_body: str) -> bool:
    if not text_body:
        return False
    lines = [x.strip() for x in text_body.splitlines() if x.strip()]
    if len(lines) < 3:
        return False

    # si varias líneas tienen 2 tokens, es lista
    good = 0
    for ln in lines[:10]:
        ln = re.sub(r"\s+", " ", ln.upper()).strip()
        parts = ln.split(" ")
        if len(parts) >= 2:
            good += 1
    return good >= 3

def reemplazar_en_documento(ruta_entrada, ruta_salida, datos, input_type, qr2_bytes=None):
    # --- Asegurar llaves “DOC” aunque vengan sin sufijo ---
    datos = datos or {}
    if not datos.get("FECHA_INICIO_DOC"):
        datos["FECHA_INICIO_DOC"] = datos.get("FECHA_INICIO", "") or ""
    if not datos.get("FECHA_ULTIMO_DOC"):
        datos["FECHA_ULTIMO_DOC"] = datos.get("FECHA_ULTIMO", "") or ""
    if not datos.get("FECHA_ALTA_DOC"):
        datos["FECHA_ALTA_DOC"] = datos.get("FECHA_ALTA", "") or ""

    rfc_val = (datos.get("RFC_ETIQUETA") or datos.get("RFC", "")).strip()
    idcif_val = (datos.get("IDCIF_ETIQUETA") or "").strip()

    # ✅ aquí se decide el QR (UNA sola línea)
    url_qr = elegir_url_qr(datos, input_type, rfc_val, idcif_val)

    # ✅ generar una sola vez
    qr_bytes, barcode_bytes = generar_qr_y_barcode(url_qr, rfc_val)

    # hard rules por si llegan diferentes:
    if datos.get("COLONIA"):
        datos["COLONIA"] = str(datos["COLONIA"]).upper()

    if input_type in ("CURP","RFC_ONLY"):
        datos["TIPO_VIALIDAD"] = "CALLE"
        datos["VIALIDAD"] = "SIN NOMBRE"
        datos["NO_INTERIOR"] = ""

    placeholders = {
        "{{ RFC ETIQUETA }}": rfc_val,
        "{{ NOMBRE ETIQUETA }}": datos.get("NOMBRE_ETIQUETA", ""),
        "{{ idCIF }}": datos.get("IDCIF_ETIQUETA", ""),
        "{{ FECHA }}": datos.get("FECHA", ""),
        "{{ CORTA }}": datos.get("FECHA_CORTA", ""),
        "{{ DENOMINACION }}": datos.get("DENOMINACION", ""),
        "{{ CAPITAL }}": datos.get("CAPITAL", ""),
        "{{ RFC }}": datos.get("RFC", ""),
        "{{ CURP }}": datos.get("CURP", ""),
        "{{ NOMBRE }}": datos.get("NOMBRE", ""),
        "{{ PRIMER APELLIDO }}": datos.get("PRIMER_APELLIDO", ""),
        "{{ SEGUNDO APELLIDO }}": datos.get("SEGUNDO_APELLIDO", ""),
        "{{ INICIO }}": datos.get("FECHA_INICIO_DOC", ""),
        "{{ ESTATUS }}": datos.get("ESTATUS", ""),
        "{{ ULTIMO }}": datos.get("FECHA_ULTIMO_DOC", ""),
        "{{ CP }}": datos.get("CP", ""),
        "{{ TIPO VIALIDAD }}": datos.get("TIPO_VIALIDAD", ""),
        "{{ VIALIDAD }}": datos.get("VIALIDAD", ""),
        "{{ NO EXTERIOR }}": datos.get("NO_EXTERIOR", ""),
        "{{ NO INTERIOR }}": datos.get("NO_INTERIOR", ""),
        "{{ COLONIA }}": datos.get("COLONIA", ""),
        "{{ LOCALIDAD }}": datos.get("LOCALIDAD", ""),
        "{{ ENTIDAD }}": datos.get("ENTIDAD", ""),
        "{{ REGIMEN }}": datos.get("REGIMEN", ""),
        "{{ ALTA }}": datos.get("FECHA_ALTA_DOC", ""),
        "{{ FECHA NACIMIENTO }}": datos.get("FECHA_NACIMIENTO", ""),
    }

    with ZipFile(ruta_entrada, "r") as zin, ZipFile(ruta_salida, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if (
                item.filename == "word/document.xml"
                or item.filename.startswith("word/header")
                or item.filename.startswith("word/footer")
            ):
                try:
                    xml_text = data.decode("utf-8")
                except UnicodeDecodeError:
                    pass
                else:
                    if idcif_val:
                        patron_idcif = r"<w:t>{{</w:t>.*?<w:t>idCIF</w:t>.*?<w:t>}}</w:t>"
                        idcif_safe = html.escape(str(idcif_val or ""), quote=False)
                        xml_text, _ = re.subn(
                            patron_idcif,
                            f"<w:t>{idcif_safe}</w:t>",
                            xml_text,
                            flags=re.DOTALL
                        )
                    for k, v in placeholders.items():
                        safe_v = html.escape(str(v or ""), quote=False)
                    
                        # k viene como "{{ RFC }}", sacamos la KEY "RFC"
                        key = k.strip().lstrip("{").rstrip("}").strip()
                    
                        # 👇 Aguanta que Word parta el placeholder:
                        # reemplaza cualquier {{ ...KEY... }} (aunque tenga espacios)
                        patron = r"\{\{[^}]*" + re.escape(key) + r"[^}]*\}\}"
                        xml_text = re.sub(patron, safe_v, xml_text, flags=re.IGNORECASE)

                    data = xml_text.encode("utf-8")

            if item.filename == "word/media/image2.png":
                data = qr_bytes
            elif item.filename in ("word/media/image9.png", "word/media/image11.png"):
                if qr2_bytes:
                    data = qr2_bytes
            elif item.filename == "word/media/image6.png":
                if barcode_bytes:
                    data = barcode_bytes

            zout.writestr(item, data)

    doc = Document(ruta_salida)

    # ✅ Segundo pase (python-docx): reemplazo robusto aunque Word parta {{ ... }} en runs
    par_placeholders = {
        # con espacios
        "{{ NOMBRE ETIQUETA }}": datos.get("NOMBRE_ETIQUETA", ""),
        "{{ RFC ETIQUETA }}": rfc_val,
        "{{ idCIF }}": datos.get("IDCIF_ETIQUETA", ""),
        "{{ FECHA }}": datos.get("FECHA", ""),
        "{{ CORTA }}": datos.get("FECHA_CORTA", ""),

        "{{ DENOMINACION }}": datos.get("DENOMINACION", ""),
        "{{ CAPITAL }}": datos.get("CAPITAL", ""),
        "{{ RFC }}": datos.get("RFC", ""),
        "{{ CURP }}": datos.get("CURP", ""),
        "{{ NOMBRE }}": datos.get("NOMBRE", ""),
        "{{ PRIMER APELLIDO }}": datos.get("PRIMER_APELLIDO", ""),
        "{{ SEGUNDO APELLIDO }}": datos.get("SEGUNDO_APELLIDO", ""),

        "{{ INICIO }}": datos.get("FECHA_INICIO_DOC", ""),
        "{{ ESTATUS }}": datos.get("ESTATUS", ""),
        "{{ ULTIMO }}": datos.get("FECHA_ULTIMO_DOC", ""),

        "{{ CP }}": datos.get("CP", ""),
        "{{ TIPO VIALIDAD }}": datos.get("TIPO_VIALIDAD", ""),
        "{{ VIALIDAD }}": datos.get("VIALIDAD", ""),
        "{{ NO EXTERIOR }}": datos.get("NO_EXTERIOR", ""),
        "{{ NO INTERIOR }}": datos.get("NO_INTERIOR", ""),
        "{{ COLONIA }}": datos.get("COLONIA", ""),
        "{{ LOCALIDAD }}": datos.get("LOCALIDAD", ""),
        "{{ ENTIDAD }}": datos.get("ENTIDAD", ""),
        "{{ REGIMEN }}": datos.get("REGIMEN", ""),
        "{{ ALTA }}": datos.get("FECHA_ALTA_DOC", ""),
        "{{ FECHA NACIMIENTO }}": datos.get("FECHA_NACIMIENTO", ""),
    }

    # variantes sin espacios (por si tu docx trae algunas así)
    par_placeholders.update({
        "{{NOMBRE ETIQUETA}}": datos.get("NOMBRE_ETIQUETA", ""),
        "{{idCIF}}": datos.get("IDCIF_ETIQUETA", ""),
        "{{FECHA}}": datos.get("FECHA", ""),
        "{{CORTA}}": datos.get("FECHA_CORTA", ""),
        "{{DENOMINACION}}": datos.get("DENOMINACION", ""),
        "{{CAPITAL}}": datos.get("CAPITAL", ""),
        "{{RFC}}": datos.get("RFC", ""),
        "{{CURP}}": datos.get("CURP", ""),
        "{{NOMBRE}}": datos.get("NOMBRE", ""),
        "{{PRIMER APELLIDO}}": datos.get("PRIMER_APELLIDO", ""),
        "{{SEGUNDO APELLIDO}}": datos.get("SEGUNDO_APELLIDO", ""),
        "{{INICIO}}": datos.get("FECHA_INICIO_DOC", ""),
        "{{ESTATUS}}": datos.get("ESTATUS", ""),
        "{{ULTIMO}}": datos.get("FECHA_ULTIMO_DOC", ""),
        "{{CP}}": datos.get("CP", ""),
        "{{TIPO VIALIDAD}}": datos.get("TIPO_VIALIDAD", ""),
        "{{VIALIDAD}}": datos.get("VIALIDAD", ""),
        "{{NO EXTERIOR}}": datos.get("NO_EXTERIOR", ""),
        "{{NO INTERIOR}}": datos.get("NO_INTERIOR", ""),
        "{{COLONIA}}": datos.get("COLONIA", ""),
        "{{LOCALIDAD}}": datos.get("LOCALIDAD", ""),
        "{{ENTIDAD}}": datos.get("ENTIDAD", ""),
        "{{REGIMEN}}": datos.get("REGIMEN", ""),
        "{{ALTA}}": datos.get("FECHA_ALTA_DOC", ""),
        "{{FECHA NACIMIENTO}}": datos.get("FECHA_NACIMIENTO", ""),
    })

    def reemplazar_en_parrafos(paragraphs):
        for p in paragraphs:
            # ✅ Solo si aún existe un placeholder completo
            if "{{" not in p.text or "}}" not in p.text:
                continue
    
            # 1) texto completo del párrafo
            full = "".join(r.text for r in p.runs)
            if "{{" not in full or "}}" not in full:
                continue
    
            new_full = full
            for k, v in par_placeholders.items():
                if k in new_full:
                    new_full = new_full.replace(k, v)
    
            # Si no cambió, no tocar
            if new_full == full:
                continue
    
            # 2) ⚠️ Esto rompe formato, así que lo hacemos SOLO si sigue habiendo placeholders
            # (o sea, para resolver un caso donde Word partió el placeholder)
            if "{{" not in new_full and "}}" not in new_full:
                # ya no quedan placeholders -> mejor NO tocar (mantiene formato original)
                continue
    
            if p.runs:
                p.runs[0].text = new_full
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(new_full)

    reemplazar_en_parrafos(doc.paragraphs)
    for section in doc.sections:
        reemplazar_en_parrafos(section.header.paragraphs)
        reemplazar_en_parrafos(section.footer.paragraphs)
        try:
            reemplazar_en_parrafos(section.first_page_header.paragraphs)
            reemplazar_en_parrafos(section.first_page_footer.paragraphs)
            reemplazar_en_parrafos(section.even_page_header.paragraphs)
            reemplazar_en_parrafos(section.even_page_footer.paragraphs)
        except Exception:
            pass
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_en_parrafos(cell.paragraphs)

    doc.save(ruta_salida)

# ================== JWT AUTH (PRO) ==================
JWT_SECRET = os.getenv("JWT_SECRET", "")
JWT_DAYS = int(os.getenv("JWT_DAYS", "30") or "30")
ALLOW_KICKOUT = os.getenv("ALLOW_KICKOUT", "0") in ("1", "true", "TRUE", "yes", "YES")

# Guardamos "sesión actual" (jti) por usuario en /data para que sobreviva reinicios
SESSIONS_PATH = os.getenv("SESSIONS_PATH", "/data/sessions.json")

# ====== IDLE TIMEOUT (cierre por inactividad) ======
SESSION_IDLE_SECONDS = int(os.getenv("SESSION_IDLE_SECONDS", "1200") or "1200")  # 20 min default

def _read_json_file(path: str) -> dict:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f) or {}
    except FileNotFoundError:
        return {}
    except Exception as e:
        print("WARN read sessions json:", e)
        return {}

def _atomic_write_json(path: str, data: dict):
    # write seguro (verás esto mucho en Render)
    d = os.path.dirname(path)
    if d and not os.path.exists(d):
        os.makedirs(d, exist_ok=True)

    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)

def get_sessions_state() -> dict:
    state = _read_json_file(SESSIONS_PATH)
    if "user_session" not in state:
        # { "username": {"jti": "...", "exp": 1234567890} }
        state["user_session"] = {}
    return state

def set_user_session(username: str, jti: str | None, exp_ts: int | None, last_ts: int | None = None, device_id: str | None = None):
    st = get_sessions_state()
    if jti and exp_ts:
        st["user_session"][username] = {
            "jti": jti,
            "exp": int(exp_ts),
            "last": int(last_ts or int(datetime.utcnow().timestamp())),
            "device_id": (device_id or "UNKNOWN"),
        }
    else:
        st["user_session"].pop(username, None)
    _atomic_write_json(SESSIONS_PATH, st)

def get_user_session(username: str) -> dict | None:
    st = get_sessions_state()
    return (st.get("user_session") or {}).get(username)

def _now_ts() -> int:
    return int(datetime.utcnow().timestamp())

def session_is_active(sess: dict | None) -> bool:
    if not sess:
        return False

    now_ts = _now_ts()
    exp_ts = int(sess.get("exp") or 0)
    last_ts = int(sess.get("last") or 0)

    # expiración absoluta del JWT
    if exp_ts and exp_ts <= now_ts:
        return False

    # idle timeout (si no hay last, considérese inactiva)
    if not last_ts:
        return False

    if (now_ts - last_ts) > SESSION_IDLE_SECONDS:
        return False

    return True

def touch_user_session(username: str):
    st = get_sessions_state()
    us = (st.get("user_session") or {})
    sess = us.get(username)
    if not sess:
        return
    sess["last"] = _now_ts()
    us[username] = sess
    st["user_session"] = us
    _atomic_write_json(SESSIONS_PATH, st)

def get_user_jti(username: str) -> str | None:
    sess = get_user_session(username) or {}
    return sess.get("jti")

def is_user_session_expired(username: str) -> bool:
    sess = get_user_session(username)
    if not sess:
        return True
    exp_ts = sess.get("exp")
    if not exp_ts:
        return True
    return int(exp_ts) <= int(datetime.utcnow().timestamp())

def crear_jwt(username: str, device_id: str | None = None) -> str:
    if not JWT_SECRET:
        raise RuntimeError("Falta JWT_SECRET en variables de entorno.")

    jti = secrets.token_urlsafe(16)

    now = datetime.utcnow()
    exp = now + timedelta(days=JWT_DAYS)

    payload = {
        "sub": username,
        "jti": jti,
        "iat": int(now.timestamp()),
        "exp": int(exp.timestamp()),
    }

    set_user_session(username, jti, payload["exp"], last_ts=payload["iat"], device_id=device_id)

    token = jwt.encode(payload, JWT_SECRET, algorithm="HS256")
    return token

def verificar_jwt(token: str):
    if not token or not JWT_SECRET:
        return None, "NO_TOKEN"

    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
    except jwt.ExpiredSignatureError:
        return None, "JWT_EXPIRED"
    except Exception:
        return None, "JWT_INVALID"

    username = payload.get("sub")
    jti = payload.get("jti")
    if not username or not jti:
        return None, "JWT_MALFORMED"

    sess = get_user_session(username)
    if not sess:
        return None, "NO_SERVER_SESSION"

    # si expiró por idle o exp absoluto
    if not session_is_active(sess):
        set_user_session(username, None, None)
        return None, "IDLE_OR_SERVER_EXPIRED"

    current = sess.get("jti")
    if not current:
        return None, "NO_JTI"
    if current != jti:
        return None, "KICKED_BY_NEW_LOGIN"

    touch_user_session(username)
    return username, "OK"

def usuario_actual_o_none():
    auth_header = request.headers.get("Authorization", "") or ""
    if not auth_header.startswith("Bearer "):
        return None, "NO_AUTH_HEADER"
    token = auth_header.split(" ", 1)[1].strip()
    return verificar_jwt(token)

# ================== APP FLASK ==================

app = Flask(__name__)
CORS(
    app,
    resources={r"/*": {"origins": ["https://constancia-7xk29.vercel.app"]}},
    methods=["GET", "POST", "OPTIONS"],
    allow_headers="*",
    expose_headers=["Content-Disposition", "X-Output-Format"],
    max_age=86400,
)

REQUEST_TOTAL = 0
REQUEST_POR_DIA = {}
SUCCESS_COUNT = 0
SUCCESS_RFCS = []
WA_VERIFY_TOKEN = os.getenv("WA_VERIFY_TOKEN", "mi_token_wa_2026")

@app.route("/dl/<token>/<filename>", methods=["GET"])
def dl_get(token: str, filename: str):
    sig = (request.args.get("sig") or "").strip()
    exp = (request.args.get("exp") or "").strip()

    if not token or not sig or not exp:
        abort(404)

    # exp válido
    try:
        exp_i = int(exp)
    except Exception:
        abort(404)

    if time.time() > exp_i:
        abort(410)  # expirado

    # firma válida
    try:
        expected = _dl_sign(token)
    except Exception:
        abort(500)

    if not hmac.compare_digest(sig, expected):
        abort(403)

    path = _dl_find_file(token, sig, exp, filename)
    if not os.path.exists(path):
        abort(404)

    # servir como descarga (ZIP u otros)
    return send_file(path, as_attachment=True, download_name=filename, mimetype="application/octet-stream")

@app.route("/wa/webhook", methods=["GET"])
def wa_webhook_verify():
    mode = request.args.get("hub.mode", "")
    token = request.args.get("hub.verify_token", "")
    challenge = request.args.get("hub.challenge", "")

    print("WA VERIFY GET -> mode:", mode, "token:", token, "challenge:", challenge)

    if mode == "subscribe" and token == WA_VERIFY_TOKEN:
        return challenge, 200

    return "Forbidden", 403

# --- WhatsApp helpers mínimos ---
WA_TOKEN = os.getenv("WA_TOKEN", "")
WA_PHONE_NUMBER_ID = os.getenv("WA_PHONE_NUMBER_ID", "")
WA_GRAPH_VERSION = os.getenv("WA_GRAPH_VERSION", "v20.0")

print("WA CONFIG -> PHONE_NUMBER_ID:", WA_PHONE_NUMBER_ID)
print("WA CONFIG -> GRAPH_VERSION:", WA_GRAPH_VERSION)
print("WA CONFIG -> TOKEN_LEN:", len(WA_TOKEN))
print("WA CONFIG -> TOKEN_LAST6:", WA_TOKEN[-6:] if WA_TOKEN else "EMPTY")

def normalizar_wa_to(wa_id: str) -> str:
    """
    Meta a veces manda México como 521XXXXXXXXXX (formato viejo).
    En la consola de prueba, los destinatarios autorizados suelen estar como 52XXXXXXXXXX (sin el 1).
    Convertimos 521 + 10 dígitos -> 52 + 10 dígitos.
    """
    wa_id = (wa_id or "").strip()
    if wa_id.startswith("521") and len(wa_id) == 13:
        return "52" + wa_id[3:]
    return wa_id

def wa_api_url(path: str) -> str:
    return f"https://graph.facebook.com/{WA_GRAPH_VERSION}/{path.lstrip('/')}"

def wa_get_media_url(media_id: str) -> str:
    """
    1) Pide a Meta la URL temporal del media (image/audio/doc).
    """
    url = wa_api_url(f"{media_id}")
    headers = {"Authorization": f"Bearer {WA_TOKEN}"}
    r = requests.get(url, headers=headers, timeout=30)
    if not r.ok:
        print("WA GET MEDIA URL ERROR:", r.status_code, r.text)
    r.raise_for_status()
    data = r.json()
    return data.get("url")  # URL temporal

def wa_download_media_bytes(media_url: str) -> bytes:
    """
    2) Descarga el archivo desde la URL temporal.
    """
    headers = {"Authorization": f"Bearer {WA_TOKEN}"}
    r = requests.get(media_url, headers=headers, timeout=60)
    if not r.ok:
        print("WA DOWNLOAD MEDIA ERROR:", r.status_code, r.text[:4000])
    r.raise_for_status()
    return r.content

def _img_bytes_to_cv2(img_bytes: bytes):
    """
    Convierte bytes de imagen a matriz OpenCV (BGR).
    """
    arr = np.frombuffer(img_bytes, dtype=np.uint8)
    img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    return img

def decode_qr_from_image_bytes(img_bytes: bytes) -> list[str]:
    """
    QR robusto SOLO con OpenCV (sin pyzbar/zbar):
    - detectAndDecodeMulti
    - reescalado
    - binarización adaptativa
    """
    img = _img_bytes_to_cv2(img_bytes)
    if img is None:
        return []

    def _try_opencv_qr(bgr):
        try:
            det = cv2.QRCodeDetector()
            ok, decoded, _, _ = det.detectAndDecodeMulti(bgr)
            if ok and decoded:
                return [d for d in decoded if d]
            d, _, _ = det.detectAndDecode(bgr)
            return [d] if d else []
        except Exception:
            return []

    # 1) directo
    outs = _try_opencv_qr(img)
    if outs:
        return list(dict.fromkeys(outs))

    # 2) reescalados
    h, w = img.shape[:2]
    for scale in (1.5, 2.0, 2.5, 3.0):
        rs = cv2.resize(img, (int(w*scale), int(h*scale)), interpolation=cv2.INTER_CUBIC)
        outs = _try_opencv_qr(rs)
        if outs:
            return list(dict.fromkeys(outs))

    # 3) binarización
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.GaussianBlur(gray, (3, 3), 0)
    thr = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                cv2.THRESH_BINARY, 31, 7)
    thr_bgr = cv2.cvtColor(thr, cv2.COLOR_GRAY2BGR)
    outs = _try_opencv_qr(thr_bgr)
    return list(dict.fromkeys([o for o in outs if o]))

def parse_sat_qr_text_to_rfc_idcif(qr_text: str):
    """
    Si el QR trae URL del SAT tipo:
    ...validadorqr.jsf?D1=10&D2=1&D3=IDCIF_RFC
    Extrae RFC + IDCIF.
    """
    if not qr_text:
        return (None, None)

    t = qr_text.strip()

    # Si viene URL
    if "D3=" in t:
        try:
            parsed = urllib.parse.urlparse(t)
            qs = urllib.parse.parse_qs(parsed.query)
            d3 = (qs.get("D3") or [None])[0]
            if d3 and "_" in d3:
                idcif, rfc = d3.split("_", 1)
                rfc = (rfc or "").strip().upper()
                idcif = (idcif or "").strip()
                if rfc and idcif:
                    return (rfc, idcif)
        except Exception:
            pass

    # A veces el QR no trae URL, trae texto con RFC/IDCIF
    rfc, idcif = extraer_rfc_idcif(t)
    return (rfc, idcif)

def ocr_text_from_image_bytes(img_bytes: bytes, timeout_sec: int = 2) -> str:
    """
    OCR con timeout corto para NO tumbar gunicorn.
    Si tarda más, lo cortamos y regresamos vacío.
    """
    img = _img_bytes_to_cv2(img_bytes)
    if img is None:
        return ""

    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.bilateralFilter(gray, 9, 75, 75)
    thr = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                cv2.THRESH_BINARY, 31, 8)

    config = "--oem 1 --psm 6"

    try:
        # pytesseract soporta timeout=
        text = pytesseract.image_to_string(thr, lang="spa+eng", config=config, timeout=timeout_sec)
    except Exception:
        try:
            text = pytesseract.image_to_string(thr, config=config, timeout=timeout_sec)
        except Exception:
            return ""

    return (text or "").strip()

def extract_rfc_idcif_from_image_bytes(img_bytes: bytes):
    """
    1) QR robusto
    2) OCR opcional con timeout corto (controlado por env OCR_ENABLED)
    """
    # 1) QR
    qr_list = decode_qr_from_image_bytes(img_bytes)
    for qr in qr_list:
        rfc, idcif = parse_sat_qr_text_to_rfc_idcif(qr)
        if rfc and idcif:
            return rfc, idcif, "QR"

    # 2) OCR solo si lo habilitas
    OCR_ENABLED = os.getenv("OCR_ENABLED", "0") in ("1", "true", "TRUE", "yes", "YES")
    if not OCR_ENABLED:
        return None, None, "NO_QR"

    text = ocr_text_from_image_bytes(img_bytes, timeout_sec=2)
    rfc, idcif = extraer_rfc_idcif(text)
    if rfc and idcif:
        return rfc, idcif, "OCR"

    return rfc, idcif, "NONE"

def wa_send_text(to_wa_id: str, text: str):
    if not (WA_TOKEN and WA_PHONE_NUMBER_ID):
        raise RuntimeError("Faltan WA_TOKEN o WA_PHONE_NUMBER_ID.")

    url = wa_api_url(f"{WA_PHONE_NUMBER_ID}/messages")
    headers = {
        "Authorization": f"Bearer {WA_TOKEN}",
        "Content-Type": "application/json",
    }
    payload = {
        "messaging_product": "whatsapp",
        "to": to_wa_id,
        "type": "text",
        "text": {"body": text},
    }

    r = requests.post(url, headers=headers, json=payload, timeout=30)

    # 👇 esto te imprime el error exacto (muy importante)
    if not r.ok:
        print("WA SEND ERROR status:", r.status_code)
        print("WA SEND ERROR body:", r.text)

    r.raise_for_status()
    return r.json()

def extraer_rfc_idcif(texto: str):
    """
    Acepta formatos:
    - RFC: TOHJ640426XXX IDCIF: 19010347XXX
    - TOHJ640426XXX 19010347XXX
    - rfc TOHJ640426XXX idcif 19010347XXX
    """
    if not texto:
        return None, None

    t = texto.strip().upper()

    # RFC: 12 o 13 chars (persona física 13, moral 12)
    rfc_regex = r"\b([A-ZÑ&]{3,4}\d{6}[A-Z0-9]{3})\b"
    # idCIF: normalmente numérico largo (pero lo dejamos flexible)
    idcif_regex = r"\b(\d{8,20})\b"

    # Caso con etiquetas
    rfc_m = re.search(r"(RFC\s*[:\-]?\s*)" + rfc_regex, t)
    id_m = re.search(r"(IDCIF\s*[:\-]?\s*)" + idcif_regex, t)

    rfc = rfc_m.group(2) if rfc_m else None
    idcif = id_m.group(2) if id_m else None

    # Caso sin etiquetas: buscar 1 RFC y 1 número
    if not rfc:
        m = re.search(rfc_regex, t)
        if m:
            rfc = m.group(1)

    if not idcif:
        nums = re.findall(idcif_regex, t)
        if nums:
            # toma el primero “largo”
            idcif = nums[0]

    return rfc, idcif

# ================== NUEVO: CHECKID + DIPOMEX (FLUJO POR API) ==================

RFC_REGEX = r"\b([A-ZÑ&]{3,4}\d{6}[A-Z0-9]{3})\b"

def extraer_rfc_solo(texto: str):
    if not texto:
        return None
    t = (texto or "").strip().upper()
    m = re.search(RFC_REGEX, t)
    return m.group(1) if m else None

def throttle_checkid(min_interval_sec=1.2):
    k = "THROTTLE:CHECKID"
    now = time.time()
    try:
        last = float(cache_get(k) or 0.0)
    except Exception:
        last = 0.0
    wait = (last + float(min_interval_sec)) - now
    if wait > 0:
        time.sleep(wait)
    try:
        cache_set(k, time.time(), ttl=int(max(10, min_interval_sec * 5)))
    except Exception:
        pass

def checkid_lookup(curp_or_rfc: str) -> dict:
    url = "https://www.checkid.mx/api/Busqueda"

    apikey = (os.getenv("CHECKID_APIKEY", "") or "").strip()

    # timeout robusto (si viene mal, cae a 20)
    try:
        timeout = int((os.getenv("CHECKID_TIMEOUT", "20") or "20").strip())
    except Exception:
        timeout = 20

    if not apikey:
        raise RuntimeError("CHECKID_NO_APIKEY")

    try:
        if not cache_get("LOGGED:CHECKID_KEY_FPR"):
            print("[CHECKID] APIKEY_FPR", (apikey[:4] + "..." + apikey[-4:]), flush=True)
            cache_set("LOGGED:CHECKID_KEY_FPR", True, ttl=3600)
    except Exception:
        pass

    term = (curp_or_rfc or "").strip().upper()
    if not term:
        raise ValueError("CHECKID_EMPTY_TERM")

    # ✅ circuit breaker: si CheckID está fallando, no insistir por 60s
    try:
        cb_sec = int((os.getenv("CHECKID_CB_SEC", "60") or "60").strip())
    except Exception:
        cb_sec = 60
    
    cb_key = f"CB:CHECKID"
    cb = cache_get(cb_key) or {}
    if isinstance(cb, dict) and cb.get("until"):
        try:
            until = float(cb["until"])
            if time.time() < until:
                print("[CHECKID] CIRCUIT_OPEN until=", until, "now=", time.time(), "term=", term, flush=True)
                raise RuntimeError("CHECKID_CIRCUIT_OPEN")
        except (ValueError, TypeError):
            pass
    
    payload = {
        "ApiKey": apikey,
        "TerminoBusqueda": term,
        "ObtenerRFC": True,
        "ObtenerCURP": True,
        "ObtenerCP": True,
        "ObtenerRegimenFiscal": True,
        "ObtenerNSS": True,
        "Obtener69o69B": False,
    }

    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "User-Agent": "CSFDocs/1.0",
    }

    last_exc = None

    # ✅ intentos configurables por ENV
    try:
        max_attempts = int((os.getenv("CHECKID_MAX_ATTEMPTS", "2") or "2").strip())
    except Exception:
        max_attempts = 2
    max_attempts = max(1, min(3, max_attempts))  # 1..3
    
    for attempt in range(max_attempts):
        try:
            throttle_checkid(1.2)
            r = requests.post(url, json=payload, headers=headers, timeout=timeout)

            print(
                "[CHECKID] HTTP_RESP",
                "status=", r.status_code,
                "ok=", r.ok,
                "term=", term,
                "attempt=", attempt + 1,
                "ctype=", (r.headers.get("Content-Type") or ""),
                flush=True
            )

            # Log mínimo útil (solo cuando falla o 1er intento si quieres)
            if not r.ok:
                print("CHECKID HTTP:", r.status_code, "term:", term, "attempt:", attempt + 1)
                print("CHECKID BODY SNIP:", (r.text or "")[:500])

            if r.status_code == 404:
                raise RuntimeError("CHECKID_NOT_FOUND")

            if r.status_code == 429:
                # rate limit: espera y reintenta
                time.sleep(1.5 * (attempt + 1))
                raise RuntimeError("CHECKID_RATE_LIMIT")

            if not r.ok:
                raise RuntimeError(f"CHECKID_HTTP_{r.status_code}")

            # ✅ Puede venir algo que NO es JSON aunque status=200
            ctype = (r.headers.get("Content-Type") or "").lower()
            if "json" not in ctype:
                print("CHECKID NON-JSON:", ctype, "term:", term)
                print("CHECKID BODY SNIP:", (r.text or "")[:500])
                raise RuntimeError("CHECKID_NON_JSON")

            try:
                data = r.json() or {}
            except Exception as e:
                print("CHECKID JSON PARSE FAIL:", repr(e), "term:", term)
                print("CHECKID BODY SNIP:", (r.text or "")[:500])
                raise RuntimeError("CHECKID_BAD_JSON")

            print(
                "[CHECKID] JSON_OK",
                "exitoso=", data.get("exitoso"),
                "error=", data.get("error"),
                "codigoError=", data.get("codigoError"),
                "term=", term,
                flush=True
            )
            
            if isinstance(data, dict) and (data.get("exitoso") is False or data.get("error")):
                code = (data.get("codigoError") or "UNKNOWN").strip()
                msg = str(data.get("error") or "")
            
                if code == "E900":
                    # intenta extraer "hasta: dd/mm/aaaa HH:MM:SS"
                    until_ts = None
                    m = re.search(r"hasta:\s*(\d{2})/(\d{2})/(\d{4})\s+(\d{2}):(\d{2}):(\d{2})", msg)
                    if m:
                        dd, mm, yyyy, HH, MM, SS = map(int, m.groups())
                        try:
                            # asume hora CDMX (ajústalo si CheckID usa otra)
                            dt = datetime(yyyy, mm, dd, HH, MM, SS, tzinfo=ZoneInfo("America/Mexico_City"))
                            until_ts = dt.timestamp()
                        except Exception:
                            until_ts = None
            
                    # fallback: 10 min si no pudo parsear
                    if not until_ts:
                        until_ts = time.time() + 600
            
                    ttl = int(max(60, until_ts - time.time()))
                    try:
                        cache_set("CB:CHECKID", {"until": until_ts}, ttl=ttl + 10)
                    except Exception:
                        pass
            
                raise RuntimeError(f"CHECKID_{code}")

            return data

        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            last_exc = e
            print("CHECKID NET ERR:", type(e).__name__, repr(e), "term:", term, "attempt:", attempt + 1)
            
            if attempt < (max_attempts - 1):
                time.sleep(1.2)
                continue
        
            try:
                cache_set("CB:CHECKID", {"until": time.time() + cb_sec}, ttl=cb_sec + 10)
            except Exception:
                pass
        
            raise

        except RuntimeError as e:
            last_exc = e
            # reintenta solo en rate limit / non-json / bad-json (intermitentes)
            if attempt == 0 and str(e) in ("CHECKID_RATE_LIMIT", "CHECKID_NON_JSON", "CHECKID_BAD_JSON"):
                time.sleep(1.2)
                continue
            raise

    # por si acaso
    raise last_exc if last_exc else RuntimeError("CHECKID_UNKNOWN")

def _norm_regimenes(reg_obj) -> list[str]:
    """
    Devuelve lista de regímenes en formato string, soporta:
      - string
      - list[str]
      - list[dict] (clave/descripcion o similares)
      - dict (a veces viene dentro)
    """
    if not reg_obj:
        return []

    raw = reg_obj.get("regimenesFiscales") if isinstance(reg_obj, dict) else reg_obj

    def to_text(item) -> str:
        if item is None:
            return ""
        # dict -> intenta varias llaves comunes
        if isinstance(item, dict):
            # ejemplos posibles: {clave:"605", descripcion:"Sueldos..."} o {codigo:"605", nombre:"..."}
            clave = item.get("clave") or item.get("codigo") or item.get("id") or ""
            desc  = item.get("descripcion") or item.get("nombre") or item.get("regimen") or ""
            clave = str(clave).strip()
            desc  = str(desc).strip()
            if clave and desc:
                return f"{clave} - {desc}"
            return desc or clave

        # string o número
        return str(item).strip()

    items = []
    if isinstance(raw, list):
        items = [to_text(x) for x in raw]
    elif isinstance(raw, dict):
        items = [to_text(raw)]
    else:
        items = [to_text(raw)]

    # limpia vacíos y duplicados (manteniendo orden)
    out = []
    seen = set()
    for x in items:
        x = x.strip()
        if not x:
            continue
        key = x.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(x)
    return out

def satpi_regimen_to_str(ci: dict) -> str:
    """
    Convierte el régimen de SATPI (ci['regimen'] = lista de objetos) a string.
    Ej:
      [{"clave":"626","descripcion":"REGIMEN SIMPLIFICADO DE CONFIANZA"}]
      -> "Régimen Simplificado de Confianza"
    """
    reg = ci.get("REGIMEN") or ci.get("regimen") or ""

    # Caso 1: ya viene string
    if isinstance(reg, str):
        return limpiar_regimen(reg)

    # Caso 2: lista de objetos (SATPI)
    if isinstance(reg, list) and reg:
        first = reg[0] or {}
        if isinstance(first, dict):
            desc = (first.get("descripcion") or first.get("Descripcion") or "").strip()
            clave = (first.get("clave") or first.get("Clave") or "").strip()
            # si no hay desc pero hay clave, igual devolvemos algo útil
            if desc:
                return limpiar_regimen(desc)
            if clave:
                return f"Régimen {clave}"
        # lista pero no dicts -> intenta string directo
        return limpiar_regimen(str(reg[0]))

    # Caso 3: dict directo
    if isinstance(reg, dict):
        desc = (reg.get("descripcion") or "").strip()
        if desc:
            return limpiar_regimen(desc)

    return ""

REGIMENES_SAT_CANON = {
    "601": "Régimen General de Ley Personas Morales",
    "602": "Régimen Simplificado de Ley Personas Morales",
    "603": "Personas Morales con Fines No Lucrativos",
    "604": "Régimen de Pequeños Contribuyentes",
    "605": "Régimen de Sueldos y Salarios e Ingresos Asimilados a Salarios",
    "606": "Régimen de Arrendamiento",
    "607": "Régimen de Enajenación o Adquisición de Bienes",
    "608": "Régimen de los Demás Ingresos",
    "609": "Régimen de Consolidación",
    "610": "Régimen de Residentes en el Extranjero sin Establecimiento Permanente en México",
    "611": "Régimen de Ingresos por Dividendos (Socios y Accionistas)",
    "612": "Régimen de las Personas Físicas con Actividades Empresariales y Profesionales",
    "613": "Régimen Intermedio de las Personas Físicas con Actividades Empresariales",
    "614": "Régimen de los Ingresos por Intereses",
    "615": "Régimen de los Ingresos por Obtención de Premios",
    "616": "Sin obligaciones fiscales",
    "617": "PEMEX",
    "618": "Régimen Simplificado de Ley Personas Físicas",
    "619": "Ingresos por la Obtención de Préstamos",
    "620": "Sociedades Cooperativas de Producción que Optan por Diferir sus Ingresos",
    "621": "Régimen de Incorporación Fiscal",
    "622": "Régimen de Actividades Agrícolas, Ganaderas, Silvícolas y Pesqueras PM",
    "623": "Régimen Opcional para Grupos de Sociedades",
    "624": "Régimen de los Coordinados",
    "625": "Régimen de las Actividades Empresariales con Ingresos a través de Plataformas Tecnológicas",
    "626": "Régimen Simplificado de Confianza",
}

def _norm_txt(s: str) -> str:
    s = s.upper()
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    s = re.sub(r"\s+", " ", s).strip()
    return s

def limpiar_regimen(regimen: str, clave: str | None = None) -> str:
    if not regimen and not clave:
        return ""

    # ✅ Caso: "605 - Sueldos y Salarios..."
    if regimen:
        m = re.match(r"^\s*(\d{3})\s*-\s*(.+?)\s*$", str(regimen))
        if m:
            clave0 = m.group(1)
            desc0 = m.group(2)
            if clave0 in REGIMENES_SAT_CANON:
                return REGIMENES_SAT_CANON[clave0]
            regimen = desc0

    # 1) si viene clave válida, es rey
    if clave and clave in REGIMENES_SAT_CANON:
        return REGIMENES_SAT_CANON[clave]

    if not regimen:
        return ""

    reg_norm = _norm_txt(regimen)

    if ("NO TIENE" in reg_norm and "REGIMEN" in reg_norm and "VIGENTE" in reg_norm) or ("SIN REGIMEN" in reg_norm):
        return ""

    for k, v in REGIMENES_SAT_CANON.items():
        if _norm_txt(v) == reg_norm:
            return v

    reg_norm2 = reg_norm
    if reg_norm2.startswith("REGIMEN "):
        reg_norm2 = reg_norm2[8:]

    for k, v in REGIMENES_SAT_CANON.items():
        if _norm_txt(v).endswith(reg_norm2):
            return v

    t = reg_norm.lower().title()
    t = re.sub(r"\bDe\b", "de", t)
    t = re.sub(r"\bDel\b", "del", t)
    t = re.sub(r"\bY\b", "y", t)
    t = re.sub(r"\bE\b", "e", t)
    t = re.sub(r"\bA\b", "a", t)
    return f"Régimen de {t}"

def regimenes_to_list(reg_val) -> list[str]:
    """
    Acepta:
    - string (CheckID): "605 - Sueldos y Salarios..."
    - lista SATPI: [{"clave":"626","descripcion":"REGIMEN ..."}, ...]
    - lista strings: ["626 - ...", "601 - ..."]
    - dict: {"clave": "...", "descripcion": "..."}
    Devuelve lista de nombres canónicos "bonitos".
    """
    out: list[str] = []

    if not reg_val:
        return out

    # dict único
    if isinstance(reg_val, dict):
        clave = str(reg_val.get("clave") or reg_val.get("Clave") or "").strip()
        desc  = str(reg_val.get("descripcion") or reg_val.get("Descripcion") or "").strip()
        s = limpiar_regimen(desc or "", clave=clave or None)
        if s:
            out.append(s)
        return out

    # lista
    if isinstance(reg_val, list):
        for item in reg_val:
            out.extend(regimenes_to_list(item))
        # dedupe manteniendo orden
        seen = set()
        uniq = []
        for x in out:
            if x not in seen:
                seen.add(x)
                uniq.append(x)
        return uniq

    # string: puede venir con separadores
    sraw = str(reg_val).strip()
    # separadores comunes
    parts = re.split(r"\s*(?:\|\||\||;|,|\n)\s*", sraw)
    for p in parts:
        p = p.strip()
        if not p:
            continue
        out.append(limpiar_regimen(p))

    # dedupe
    seen = set()
    uniq = []
    for x in out:
        if x and x not in seen:
            seen.add(x)
            uniq.append(x)
    return uniq

def _norm_checkid_fields(ci_raw: dict) -> dict:
    ci_raw = ci_raw or {}
    res = ci_raw.get("resultado") or {}

    rfc_obj = res.get("rfc") or {}
    curp_obj = res.get("curp") or {}
    cp_obj = res.get("codigoPostal") or {}
    reg_obj = res.get("regimenFiscal") or {}
    nss_obj = res.get("nss") or {}
    e69_obj = res.get("estado69o69B") or {}

    def pick(*vals):
        for v in vals:
            if v not in (None, "", []):
                return str(v).strip()
        return ""

    razon = pick(rfc_obj.get("razonSocial"))
    nombres = pick(curp_obj.get("nombres"))
    ape1 = pick(curp_obj.get("primerApellido"))
    ape2 = pick(curp_obj.get("segundoApellido"))

    if not (nombres or ape1 or ape2) and razon:
        nombres = razon

        out_name_source = "RAZON_SOCIAL"
    else:
        out_name_source = ""

    fn_text = pick(curp_obj.get("fechaNacimientoText"))
    fn_iso = pick(curp_obj.get("fechaNacimiento"))
    fecha_nac = fn_text or fn_iso

    rfc = pick(rfc_obj.get("rfc"), rfc_obj.get("rfcRepresentante"))
    curp = pick(curp_obj.get("curp"), rfc_obj.get("curp"), rfc_obj.get("curpRepresentante"))

    cp = pick(cp_obj.get("codigoPostal"))

    # -------- DOMICILIO / UBICACIÓN (para NO caer en CDMX) --------
    # OJO: CheckID puede variar llaves; por eso probamos varias.
    entidad = pick(
        cp_obj.get("estado"),
        cp_obj.get("d_estado"),
        res.get("estado"),
        curp_obj.get("estado"),
        curp_obj.get("entidad"),
        curp_obj.get("entidadNacimientoText"),
    )

    municipio = pick(
        cp_obj.get("municipio"),
        cp_obj.get("D_mnpio"),
        cp_obj.get("d_mnpio"),
        cp_obj.get("alcaldia"),
        cp_obj.get("delegacion"),
        res.get("municipio"),
        curp_obj.get("municipioRegistro"),
        curp_obj.get("municipio"),
    )

    ciudad = pick(
        cp_obj.get("ciudad"),
        cp_obj.get("d_ciudad"),
        res.get("ciudad"),
        curp_obj.get("ciudad"),
    )

    colonia = pick(
        cp_obj.get("colonia"),
        cp_obj.get("asentamiento"),
        cp_obj.get("d_asenta"),
        res.get("colonia"),
    )

    # ✅ MULTI-RÉGIMEN (DOCX usa solo el primero)
    regimenes_list = _norm_regimenes(reg_obj)
    regimen_first_raw = regimenes_list[0] if regimenes_list else ""
    regimen_text = limpiar_regimen(regimen_first_raw)

    con_prob = bool(e69_obj.get("conProblema")) if isinstance(e69_obj, dict) else False
    estatus = "ACTIVO" if not con_prob else "CON_PROBLEMA_69B"

    out = {
        "RFC": rfc,
        "CURP": curp,
        "CP": cp,
        "NOMBRE": nombres,
        "APELLIDO_PATERNO": ape1,
        "APELLIDO_MATERNO": ape2,

        "REGIMEN": regimen_text,
        "REGIMENES": [limpiar_regimen(x) for x in regimenes_list],
        "FECHA_NACIMIENTO": fecha_nac,
        "ESTATUS": estatus,
        "NSS": pick(nss_obj.get("nss")),
        "RAZON_SOCIAL": razon,

        # ✅ CLAVE: esto evita que caigas en CDMX cuando CP viene vacío
        "ENTIDAD": entidad,
        "MUNICIPIO": municipio,
        "LOCALIDAD": (ciudad or municipio),
        "COLONIA": colonia,
    }

    if cp and re.sub(r"\D+", "", cp).strip().__len__() == 5:
        out["_CP_SOURCE"] = "CHECKID"
    if (regimen_text or "").strip():
        out["_REG_SOURCE"] = "CHECKID"

    if out_name_source:
        out["_NAME_SOURCE"] = out_name_source

    return out

def dipomex_by_cp(cp: str) -> dict:
    apikey = (os.getenv("DIPOMEX_APIKEY", "") or "").strip()
    timeout = int(os.getenv("DIPOMEX_TIMEOUT", "12") or "12")

    if not apikey:
        print("DIPOMEX WARN: falta DIPOMEX_APIKEY")
        return {}

    cp = re.sub(r"\D+", "", (cp or "")).strip()
    if len(cp) != 5:
        print("DIPOMEX WARN: CP inválido:", cp)
        return {}

    url = "https://api.tau.com.mx/dipomex/v1/codigo_postal"
    headers = {"APIKEY": apikey, "Accept": "application/json", "User-Agent": "CSFDocs/1.0"}

    for attempt in range(3):
        try:
            r = requests.get(url, headers=headers, params={"cp": cp}, timeout=timeout)
            print("DIPOMEX:", r.status_code, r.url)

            if r.status_code == 429:
                print("DIPOMEX WARN: rate limit 429", (r.text or "")[:200])
                time.sleep(1.2 * (attempt + 1))
                continue

            if r.status_code >= 500:
                print("DIPOMEX ERROR 5xx:", r.status_code, (r.text or "")[:200])
                time.sleep(0.6 * (attempt + 1))
                continue

            if not r.ok:
                print("DIPOMEX ERROR 4xx:", r.status_code, (r.text or "")[:400])
                return {}

            try:
                j = r.json() or {}
            except Exception:
                print("DIPOMEX ERROR: no JSON", (r.text or "")[:200])
                return {}

            codigo_postal = j.get("codigo_postal")
            return codigo_postal if isinstance(codigo_postal, dict) else {}

        except (requests.Timeout, requests.ConnectionError) as e:
            print("DIPOMEX WARN: timeout/conn", repr(e))
            time.sleep(0.6 * (attempt + 1))
        except Exception as e:
            print("DIPOMEX ERROR: exception", repr(e))
            return {}

    print("DIPOMEX WARN: servicio no disponible")
    return {}

# ===== SEPOMEX (CSV local) =====
SEPOMEX_CSV_PATH = (os.getenv("SEPOMEX_CSV_PATH", "") or "").strip()

_SEPOMEX_LOCK = threading.Lock()
_SEPOMEX_LOADED = False
_SEPOMEX_BY_CP = {}
_SEPOMEX_ERR = None

import re

def reconcile_location_by_cp(datos: dict, seed_key: str = "", force_mun: bool = False) -> dict:
    datos = datos or {}

    cp = re.sub(r"\D+", "", (datos.get("CP") or "")).strip()
    if len(cp) != 5:
        return datos

    meta = sepomex_by_cp(cp) or {}
    ent_meta = (meta.get("estado") or "").strip().upper()
    mun_meta = (meta.get("municipio") or "").strip().upper()

    colonias_meta = meta.get("colonias") or meta.get("asentamientos") or []
    colonias_u = set()
    for x in colonias_meta:
        if not x:
            continue
        if isinstance(x, dict):
            v = (x.get("colonia") or "").strip().upper()
        else:
            v = str(x).strip().upper()
        if v:
            colonias_u.add(v)

    # ENTIDAD: siempre al CP
    if ent_meta:
        datos["ENTIDAD"] = ent_meta
        datos["_ENT_SOURCE"] = "SEPOMEX"

    # MUNICIPIO/LOCALIDAD: solo si no está locked o si force_mun=True
    mun_locked = bool(datos.get("_MUN_LOCK"))
    if mun_meta and (force_mun or (not mun_locked)):
        datos["MUNICIPIO"] = mun_meta
        datos["LOCALIDAD"] = mun_meta
        datos["_MUN_SOURCE"] = "SEPOMEX"
        # ⚠️ NO tocar _MUN_LOCK aquí

    # COLONIA: si vacía o no pertenece al CP -> repick
    col_u = (datos.get("COLONIA") or "").strip().upper()
    if colonias_u and ((not col_u) or (col_u not in colonias_u)):
        col_pick = sepomex_pick_colonia_by_cp(cp, seed_key=seed_key)
        if col_pick:
            datos["COLONIA"] = (col_pick or "").strip().upper()
            datos["_COL_SOURCE"] = "SEPOMEX"

    return datos

def _sepomex_csv_default_path() -> str:
    base_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_dir, "sepomex.csv")

def _open_csv_robust(path: str):
    """
    Abre sepomex.csv SIN meter '�' (replacement char).
    Detecta BOM utf-8; si no hay, prefiere cp1252 (catálogos MX).
    """
    import codecs

    # 1) Detecta BOM UTF-8
    with open(path, "rb") as fb:
        head = fb.read(4)

    if head.startswith(codecs.BOM_UTF8):
        print("SEPOMEX: encoding=utf-8-sig (BOM detectado)")
        return open(path, "r", encoding="utf-8-sig", errors="strict", newline="")

    # 2) Sin BOM: normalmente es cp1252 / latin-1
    try:
        print("SEPOMEX: encoding=cp1252 (sin BOM)")
        return open(path, "r", encoding="cp1252", errors="strict", newline="")
    except UnicodeDecodeError:
        print("SEPOMEX: encoding=latin-1 (fallback)")
        return open(path, "r", encoding="latin-1", errors="strict", newline="")

def sepomex_load_once():
    """
    Carga sepomex.csv una sola vez a memoria.
    Robusto para CSVs donde el CP real puede venir en d_CP (col J) y d_codigo puede ser distinto.
    - Prefiere d_CP
    - Si d_CP no sirve, usa d_codigo
    - Si ambos existen y son distintos, indexa en AMBOS (alias)
    """
    global _SEPOMEX_LOADED, _SEPOMEX_BY_CP, _SEPOMEX_ERR

    with _SEPOMEX_LOCK:
        if _SEPOMEX_LOADED:
            return

        path = SEPOMEX_CSV_PATH or _sepomex_csv_default_path()

        try:
            if not os.path.exists(path):
                _SEPOMEX_ERR = f"SEPOMEX CSV no existe: {path}"
                print("SEPOMEX:", _SEPOMEX_ERR)
                _SEPOMEX_BY_CP = {}
                _SEPOMEX_LOADED = True
                return

            by_cp_cols = defaultdict(set)
            by_cp_meta = {}

            with _open_csv_robust(path) as f:
                # 1) Encuentra la fila header real
                reader = csv.reader(f)
                header = None

                for row in reader:
                    if not row:
                        continue
                    norm = [str(x or "").strip() for x in row]
                    if any((c or "").strip().lower() == "d_codigo" for c in norm):
                        header = norm
                        break

                if not header:
                    _SEPOMEX_ERR = "SEPOMEX: no encontré header con d_codigo"
                    print(_SEPOMEX_ERR)
                    _SEPOMEX_BY_CP = {}
                    _SEPOMEX_LOADED = True
                    return

                # 2) Lee como dict usando ese header
                dict_reader = csv.DictReader(f, fieldnames=header)

                # helper: obtener key case-insensitive
                def g(d, key):
                    if key in d:
                        return d.get(key)
                    lk = key.lower()
                    for k in d.keys():
                        if (k or "").strip().lower() == lk:
                            return d.get(k)
                    return None

                def _norm_cp(raw_val: str) -> str:
                    raw_val = (raw_val or "").strip()
                    raw_val = raw_val.replace(".0", "")  # catálogos con floats
                    return re.sub(r"\D+", "", raw_val)

                def _pick_cp_keys(d: dict) -> list:
                    """
                    Devuelve lista de CPs por los que se debe indexar esta fila.
                    - Prefiere d_CP como CP real
                    - Si d_CP y d_codigo son distintos, indexa ambos (alias)
                    """
                    cp_main = ""
                    cp_alt = ""

                    # CP real (en tu CSV viene aquí)
                    for key in ("d_CP", "D_CP", "d_cp", "dcp", "cp", "CP"):
                        cp_try = _norm_cp(g(d, key))
                        if len(cp_try) == 5:
                            cp_main = cp_try
                            break

                    # CP alterno
                    for key in ("d_codigo", "D_codigo", "dCodigo", "dcodigo"):
                        cp_try = _norm_cp(g(d, key))
                        if len(cp_try) == 5:
                            cp_alt = cp_try
                            break

                    cps = []
                    if len(cp_main) == 5:
                        cps.append(cp_main)

                    # si no hay main, usa alt
                    if not cps and len(cp_alt) == 5:
                        cps.append(cp_alt)

                    # si hay ambos y son diferentes, indexa ambos (alias)
                    if len(cp_main) == 5 and len(cp_alt) == 5 and cp_alt != cp_main:
                        cps.append(cp_alt)

                    # unique manteniendo orden
                    out = []
                    seen = set()
                    for x in cps:
                        if x not in seen:
                            seen.add(x)
                            out.append(x)
                    return out

                for d in dict_reader:
                    cps = _pick_cp_keys(d)
                    if not cps:
                        continue

                    col = (g(d, "d_asenta") or "").strip()
                    tipo = (g(d, "d_tipo_asenta") or "").strip()
                    mnpio = (g(d, "D_mnpio") or g(d, "d_mnpio") or "").strip()
                    estado = (g(d, "d_estado") or "").strip()
                    ciudad = (g(d, "d_ciudad") or "").strip()

                    for cp in cps:
                        if col:
                            by_cp_cols[cp].add(col)

                        # guarda meta si no existe, o rellena huecos
                        if cp not in by_cp_meta:
                            by_cp_meta[cp] = {
                                "estado": estado,
                                "municipio": mnpio,
                                "ciudad": ciudad,
                                "tipo_asenta": tipo,
                            }
                        else:
                            meta = by_cp_meta[cp]
                            if (not meta.get("estado")) and estado:
                                meta["estado"] = estado
                            if (not meta.get("municipio")) and mnpio:
                                meta["municipio"] = mnpio
                            if (not meta.get("ciudad")) and ciudad:
                                meta["ciudad"] = ciudad
                            if (not meta.get("tipo_asenta")) and tipo:
                                meta["tipo_asenta"] = tipo

            out = {}
            for cp, cols in by_cp_cols.items():
                meta = by_cp_meta.get(cp) or {}
                out[cp] = {
                    "codigo_postal": cp,
                    "estado": meta.get("estado", ""),
                    "municipio": meta.get("municipio", ""),
                    "ciudad": meta.get("ciudad", ""),
                    "tipo_asenta": meta.get("tipo_asenta", ""),
                    "colonias": [{"colonia": c} for c in sorted(cols)],
                }

            _SEPOMEX_BY_CP = out
            _SEPOMEX_ERR = None
            _SEPOMEX_LOADED = True
            print(f"SEPOMEX: loaded {len(_SEPOMEX_BY_CP)} CPs from {path}")

        except Exception as e:
            _SEPOMEX_ERR = f"SEPOMEX load error: {repr(e)}"
            print("SEPOMEX:", _SEPOMEX_ERR)
            _SEPOMEX_BY_CP = {}
            _SEPOMEX_LOADED = True

def sepomex_pick_cp_by_entidad_municipio(entidad: str, municipio: str, seed_key: str = "") -> str:
    """Compat: elige un CP usando ENTIDAD + MUNICIPIO.

    FIX:
    - La versión anterior dependía de _SEPOMEX_ROWS (no existe en tu carga real),
      por eso siempre regresaba '' y caía al fallback por ENTIDAD.
    - Esta versión usa el índice real cargado desde sepomex.csv vía sepomex_pick_cp_by_ent_mun().
    """
    entidad = (entidad or "").strip().upper()
    municipio = (municipio or "").strip().upper()
    if not entidad or not municipio:
        return ""
    return sepomex_pick_cp_by_ent_mun(entidad, municipio, seed_key=seed_key) or ""

def sepomex_fill_domicilio_desde_entidad(datos: dict, seed_key: str = "") -> dict:
    try:
        mun_locked = bool(datos.get("_MUN_LOCK"))
        no_cp_pick = bool(datos.get("_NO_SEPOMEX_CP_PICK"))

        cp_val = re.sub(r"\D+", "", (datos.get("CP") or "")).strip()
        
        if len(cp_val) == 5 and (datos.get("CP") or "") != cp_val:
            datos["CP"] = cp_val

        entidad_raw = (datos.get("ENTIDAD") or "")
        entidad = entidad_raw.strip().upper()

        mun = (datos.get("MUNICIPIO") or "").strip().upper()
        loc = (datos.get("LOCALIDAD") or "").strip().upper()
        col = (datos.get("COLONIA") or "").strip().upper()

        mun_pref = mun or loc

        cp_src = (datos.get("_CP_SOURCE") or "").strip().upper()
        # 1) Si CP no válido -> pick CP con más contexto posible
        if (not no_cp_pick) and len(cp_val) != 5 and entidad and cp_src not in ("SATPI", "CHECKID"):
            cp_pick = ""
            if mun_pref:
                cp_pick = sepomex_pick_cp_by_ent_mun(entidad, mun_pref, seed_key=seed_key)
            if not cp_pick:
                cp_pick = sepomex_pick_cp_by_entidad(entidad, seed_key=seed_key)
        
            if cp_pick:
                cp_val = re.sub(r"\D+", "", str(cp_pick)).strip()
                if len(cp_val) == 5:
                    datos["CP"] = cp_val
                    datos["_CP_SOURCE"] = "SEPOMEX_PICK"

        cp_val = re.sub(r"\D+", "", (datos.get("CP") or "")).strip()

        # 2) Con CP válido, rellena colonia y (solo si NO locked) municipio/localidad
        if len(cp_val) == 5:
            meta = sepomex_by_cp(cp_val) or {}

            mnpio = (meta.get("municipio") or "").strip().upper()
            ciudad = (meta.get("ciudad") or "").strip().upper()
            estado = (meta.get("estado") or "").strip().upper()

            # ✅ ENTIDAD: si está vacía, setéala.
            # ✅ Si NO está vacía pero NO coincide con la del CP, corrígela
            #    (PERO si viene de GOBMX, respétala).
            ent_src = (datos.get("_ENT_SOURCE") or "").strip().upper()
            if estado:
                if not entidad:
                    datos["ENTIDAD"] = estado
                else:
                    # mismatch típico: "MICHOACAN" vs "MICHOACÁN DE OCAMPO"
                    if ent_src != "GOBMX" and _norm_cmp(entidad) != _norm_cmp(estado):
                        datos["ENTIDAD"] = estado

            # ✅ NO tocar mun/loc si locked
            if not mun_locked:
                # si falta MUNICIPIO, usa el del CP
                if mnpio and not mun:
                    datos["MUNICIPIO"] = mnpio
                    mun = mnpio

                # si falta LOCALIDAD, usa ciudad o municipio
                if not loc:
                    datos["LOCALIDAD"] = ciudad or mnpio or mun

            # colonia si falta
            if not col:
                col_pick = sepomex_pick_colonia_by_cp(cp_val, seed_key=seed_key)
                if col_pick:
                    datos["COLONIA"] = str(col_pick).strip().upper()

        # ✅ Nunca dejar None en campos de texto
        for k in ("ENTIDAD", "MUNICIPIO", "LOCALIDAD", "COLONIA"):
            if datos.get(k) is None:
                datos[k] = ""

    except Exception as e:
        print("SEPOMEX FILL FAIL:", repr(e))

    return datos

def sepomex_by_cp(cp: str) -> dict:
    cp = re.sub(r"\D+", "", (cp or "")).strip()
    if len(cp) != 5:
        return {}
    sepomex_load_once()
    base = _SEPOMEX_BY_CP.get(cp) or {}
    return dict(base) if isinstance(base, dict) else {}

def _norm_cmp(s: str) -> str:
    """
    Normaliza para comparar: MAYUS, sin acentos, espacios limpios.
    """
    s = (s or "").strip().upper()
    # conserva Ñ como Ñ
    s = s.replace("Ñ", "__ENYE__")
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = s.replace("__ENYE__", "Ñ")
    s = re.sub(r"\s+", " ", s).strip()
    return s

def sepomex_pick_cp_by_ent_mun(entidad: str, municipio: str, seed_key: str = "") -> str:
    sepomex_load_once()
    items = list((_SEPOMEX_BY_CP or {}).items())

    ent = _norm_cmp(entidad)
    mun = _norm_cmp(municipio)

    # ------------------------------------------------------------
    # Heurística CDMX vs EDOMEX cuando la fuente manda "MEXICO"
    # Si el "municipio" parece ser ALCALDÍA, entonces es CDMX.
    # ------------------------------------------------------------
    CDMX_ALCALDIAS = {
        "ALVARO OBREGON", "AZCAPOTZALCO", "BENITO JUAREZ", "COYOACAN",
        "CUAJIMALPA DE MORELOS", "CUAUHTEMOC", "GUSTAVO A MADERO",
        "IZTACALCO", "IZTAPALAPA", "LA MAGDALENA CONTRERAS", "MIGUEL HIDALGO",
        "MILPA ALTA", "TLAHUAC", "TLALPAN", "VENUSTIANO CARRANZA", "XOCHIMILCO",
    }

    # ------------------------------------------------------------
    # ALIAS DE ENTIDADES (SEPOMEX vs fuentes externas)
    # OJO: "MEXICO" lo dejamos como "MÉXICO" (Edomex) por default.
    # ------------------------------------------------------------
    ENT_ALIASES = {
        # CDMX
        "CIUDAD DE MEXICO": "CIUDAD DE MÉXICO",
        "CDMX": "CIUDAD DE MÉXICO",
        "DISTRITO FEDERAL": "CIUDAD DE MÉXICO",

        # Edomex (mantener separado de CDMX)
        "ESTADO DE MEXICO": "MÉXICO",
        "EDOMEX": "MÉXICO",
        "MEXICO": "MÉXICO",

        # Estados con nombre largo
        "VERACRUZ": "VERACRUZ DE IGNACIO DE LA LLAVE",
        "MICHOACAN": "MICHOACÁN DE OCAMPO",
        "COAHUILA": "COAHUILA DE ZARAGOZA",
        "QUERETARO": "QUERÉTARO",
        "SAN LUIS POTOSI": "SAN LUIS POTOSÍ",
        "YUCATAN": "YUCATÁN",
        "NUEVO LEON": "NUEVO LEÓN",
    }

    ent = _norm_cmp(ENT_ALIASES.get(ent, ent))

    # Si quedó como MÉXICO (Edomex) pero el "municipio" es alcaldía -> realmente es CDMX
    if ent == "MÉXICO" and mun in CDMX_ALCALDIAS:
        ent = "CIUDAD DE MÉXICO"

    # ------------------------------------------------------------
    # MUNICIPIO BASURA (pero NO te vueles CDMX)
    # ------------------------------------------------------------
    if ent not in ("CIUDAD DE MÉXICO",):
        if mun in ("CIUDAD DE MEXICO", "CDMX", "DISTRITO FEDERAL"):
            mun = ""

    if not ent or not mun:
        return ""

    # 1er intento: match exacto ENT+MUN
    candidatos = []
    for cp, info in items:
        try:
            e = _norm_cmp(info.get("estado") or "")
            m = _norm_cmp(info.get("municipio") or "")
            if e == ent and m == mun:
                candidatos.append(cp)
        except Exception:
            pass

    # 2do intento: municipio con artículo al final (ej. "SALTO EL" -> "EL SALTO")
    if not candidatos:
        def _swap_articulo_final(m: str) -> str:
            toks = [t for t in (m or "").split() if t]
            if len(toks) < 2:
                return m
            last = toks[-1]
            if last in ("EL", "LA", "LOS", "LAS", "DEL", "DE", "AL"):
                return " ".join([last] + toks[:-1])
            return m

        mun2 = _swap_articulo_final(mun)
        if mun2 != mun:
            for cp, info in items:
                try:
                    e = _norm_cmp(info.get("estado") or "")
                    m = _norm_cmp(info.get("municipio") or "")
                    if e == ent and m == mun2:
                        candidatos.append(cp)
                except Exception:
                    pass

    if not candidatos:
        return ""

    idx = _det_rand_int(f"CP|{seed_key}|{ent}|{mun}", 0, len(candidatos) - 1)
    return candidatos[idx]

def sepomex_pick_cp_by_entidad(entidad: str, seed_key: str = "") -> str:
    sepomex_load_once()
    items = list((_SEPOMEX_BY_CP or {}).items())

    ent = _norm_cmp(entidad)
    if not ent:
        return ""

    candidatos = []
    for cp, info in items:
        try:
            e = _norm_cmp(info.get("estado") or "")
            if e == ent:
                candidatos.append(cp)
        except Exception:
            pass

    if not candidatos:
        return ""

    idx = _det_rand_int(f"CPENT|{seed_key}|{ent}", 0, len(candidatos) - 1)
    return candidatos[idx]

def sepomex_pick_colonia_by_cp(cp: str, seed_key: str = "") -> str:
    d = sepomex_by_cp(cp) or {}
    cols = d.get("colonias") or []
    if not cols:
        return ""
    idx = _det_rand_int(f"COL|{seed_key}|{cp}", 0, len(cols) - 1)
    try:
        return (cols[idx] or {}).get("colonia", "").strip().upper()
    except Exception:
        return ""

def _pick_first_colonia(dip: dict) -> str:
    cols = dip.get("colonias") or []
    if isinstance(cols, list) and cols:
        first = cols[0] or {}
        if isinstance(first, dict):
            return (first.get("colonia") or "").strip()
        if isinstance(first, str):
            return first.strip()
    return ""

def _det_rand_int(seed: str, lo: int, hi: int) -> int:
    """
    Aleatorio determinístico (no cambia en reintentos).
    """
    seed = (seed or "").encode("utf-8")
    h = hashlib.sha256(seed).hexdigest()
    n = int(h[:12], 16)  # suficiente
    return lo + (n % (hi - lo + 1))

def _fake_date_components(year: int, seed_key: str):
    rng = _det_rng(seed_key)
    day = rng.randint(1, 30)
    month = rng.randint(1, 12)
    return day, month, year

def _fmt_dd_de_mes_de_aaaa(day: int, month: int, year: int) -> str:
    return f"{day:02d} DE {MESES_ES[month]} DE {year}"

def _fmt_dd_mm_aaaa(day: int, month: int, year: int) -> str:
    return f"{day:02d}/{month:02d}/{year}"

# Conectores frecuentes en apellidos/nombres compuestos (MX)
_SURNAME_JOINERS = {
    "DE", "DEL", "LA", "LAS", "LOS", "Y", "MC", "MAC", "VON", "VAN",
    "SAN", "SANTA", "SANTO",
}

# Frases multi-token comunes
_MULTI_JOINERS = {
    ("DE", "LA"),
    ("DE", "LAS"),
    ("DE", "LOS"),
    ("DEL", "RIO"),
    ("DEL", "VALLE"),
    ("VON", "DER"),
    ("VAN", "DER"),
}

def _norm_spaces(s: str) -> str:
    s = (s or "").strip().upper()
    s = re.sub(r"\s+", " ", s)
    return s

def _tokenize_name(full_name: str) -> list[str]:
    full = _norm_spaces(full_name)
    if not full:
        return []
    # Quita signos raros pero conserva Ñ y letras
    full = re.sub(r"[^A-ZÑ\s]", " ", full)
    full = re.sub(r"\s+", " ", full).strip()
    return full.split(" ") if full else []

def _merge_multi_joiners(tokens: list[str]) -> list[str]:
    """
    Une frases multi-token conocidas en un solo token con guion bajo,
    para no romper apellidos tipo "DE LA".
    """
    out = []
    i = 0
    n = len(tokens)
    while i < n:
        if i + 1 < n and (tokens[i], tokens[i+1]) in _MULTI_JOINERS:
            out.append(tokens[i] + "_" + tokens[i+1])
            i += 2
            continue
        if i + 2 < n and (tokens[i], tokens[i+1], tokens[i+2]) == ("DE", "LA", "O"):
            # raro, pero por si acaso
            out.append("DE_LA_O")
            i += 3
            continue
        out.append(tokens[i])
        i += 1
    return out

def _unmerge(token: str) -> str:
    return token.replace("_", " ")

def _take_compound_from_end(tokens: list[str]) -> tuple[list[str], str]:
    """
    Toma un apellido desde el final considerando conectores.
    Ej:
      ["JUAN","DE_LA","CRUZ","HERNANDEZ"] -> toma "CRUZ HERNANDEZ" o "DE LA CRUZ" según patrón.
    Regla:
      - Siempre toma al menos 1 token (el final).
      - Si antes del último hay un conector (DE/DEL/DE_LA/etc), lo incluye.
      - Si hay cadena de conectores razonable, sigue incluyendo.
    """
    if not tokens:
        return [], ""

    take = [tokens[-1]]
    i = len(tokens) - 2

    # helper: considera también tokens unidos con "_" como conectores
    def is_joiner(tok: str) -> bool:
        if not tok:
            return False
        t = tok
        # token multi "DE_LA" cuenta como joiner
        if "_" in t:
            base = t.replace("_", " ")
            return base in {"DE LA", "DE LOS", "DE LAS", "VAN DER", "VON DER"} or t in {"DE_LA", "DE_LOS", "DE_LAS", "VAN_DER", "VON_DER"}
        return t in _SURNAME_JOINERS

    # incluye conectores previos (y/o parte compuesta)
    # Ej: "... DE_LA CRUZ" => joiner "DE_LA" + "CRUZ"
    # Ej: "... DEL RIO" => joiner "DEL" + "RIO"
    while i >= 0:
        prev = tokens[i]
        if is_joiner(prev):
            take.insert(0, prev)
            i -= 1
            continue

        # Caso: token normal antes del joiner ya tomado (ej "DE_LA CRUZ" => ya tomaste CRUZ,
        # si el anterior es DE_LA, lo tomas; pero si el anterior es "LA" (joiner) sin DE, también lo toma arriba)
        # Para apellidos dobles sin conectores (PEREZ LOPEZ) no seguimos tomando aquí.
        break

    remaining = tokens[: i + 1]
    surname = " ".join(_unmerge(t) for t in take).strip()
    return remaining, surname

def desglose_nombre_mex_pro(full_name: str) -> dict:
    """
    Heurística PRO MX:
    - Tokeniza y normaliza.
    - Extrae apellido materno desde el final (considerando conectores).
    - Extrae apellido paterno desde lo que queda al final:
        - Si quedan >= 2 tokens, intenta tomar 1 token + posibles conectores previos.
        - Si quedan 1 token, eso es el paterno y nombres vacío.
    - Lo restante son nombres.
    """
    toks = _tokenize_name(full_name)
    toks = _merge_multi_joiners(toks)

    if not toks:
        return {"NOMBRE": "", "APELLIDO_PATERNO": "", "APELLIDO_MATERNO": ""}

    # Si hay 1-2 tokens, no hay mucho que hacer
    if len(toks) == 1:
        return {"NOMBRE": _unmerge(toks[0]), "APELLIDO_PATERNO": "", "APELLIDO_MATERNO": ""}
    if len(toks) == 2:
        return {"NOMBRE": _unmerge(toks[0]), "APELLIDO_PATERNO": _unmerge(toks[1]), "APELLIDO_MATERNO": ""}

    # 1) Materno desde el final (con conectores)
    rem, am = _take_compound_from_end(toks)

    # Si al sacar materno nos quedamos sin nada, reacomoda
    if not rem:
        # todo era "apellido" raro; deja el original como nombre
        return {"NOMBRE": " ".join(_unmerge(t) for t in toks), "APELLIDO_PATERNO": "", "APELLIDO_MATERNO": ""}

    # 2) Paterno: por defecto, toma 1 token final (o compuesto si trae conectores)
    #    Para capturar "DE LA CRUZ" como paterno cuando el materno es "HERNANDEZ"
    #    necesitamos que el "DE_LA" se quede unido con "CRUZ".
    #    La función _take_compound_from_end ya hace eso.
    rem2, ap = _take_compound_from_end(rem)

    # 3) Nombres: lo que quede
    nombres = " ".join(_unmerge(t) for t in rem2).strip()

    # Edge: si nombres quedó vacío (ej "PEREZ LOPEZ GARCIA"), deja algo razonable
    return {
        "NOMBRE": nombres,
        "APELLIDO_PATERNO": ap.strip(),
        "APELLIDO_MATERNO": am.strip(),
    }

def build_datos_final_from_ci(ci: dict, seed_key: str = "") -> dict:
    ci = ci or {}

    # seed determinístico
    seed_key = (seed_key or (ci.get("RFC") or ci.get("CURP") or "")).strip().upper()

    # -----------------------
    # 1) Dirección base
    # -----------------------
    cp_final = re.sub(r"\D+", "", (ci.get("CP") or "")).strip()
    entidad = (ci.get("ENTIDAD") or "").strip().upper()
    municipio_in = (ci.get("MUNICIPIO") or ci.get("LOCALIDAD") or "").strip().upper()
    colonia = (ci.get("COLONIA") or "").strip().upper()

    mun_locked = bool(ci.get("_MUN_LOCK"))
    cp_picked = False

    # 1A) Si CP no es válido: pick por ENTIDAD+MUNICIPIO; si no, por ENTIDAD
    if len(cp_final) != 5 and entidad:
        cp_pick = ""
        if municipio_in:
            cp_pick = sepomex_pick_cp_by_ent_mun(entidad, municipio_in, seed_key=seed_key)
        if not cp_pick:
            cp_pick = sepomex_pick_cp_by_entidad(entidad, seed_key=seed_key)
        if cp_pick:
            cp_final = cp_pick
            cp_picked = True

    # 1B) Reconciliar por CP (si existe CP válido)
    municipio = municipio_in  # copia "original"
    if len(cp_final) == 5:
        cp_src = (ci.get("_CP_SOURCE") or "").strip().upper()

        # CP manda si:
        # - viene de SATPI/CHECKID, o
        # - lo pickeamos nosotros (cp_picked=True)
        force_mun = (cp_src in ("SATPI", "CHECKID", "SEPOMEX_PICK")) or cp_picked

        tmp = {
            "CP": cp_final,
            "ENTIDAD": entidad,
            "MUNICIPIO": municipio,
            "LOCALIDAD": municipio,
            "COLONIA": colonia,
            "_MUN_LOCK": mun_locked,
        }
        tmp = reconcile_location_by_cp(tmp, seed_key=seed_key, force_mun=force_mun)

        cp_final = re.sub(r"\D+", "", (tmp.get("CP") or cp_final)).strip()
        entidad = (tmp.get("ENTIDAD") or entidad).strip().upper()
        colonia = (tmp.get("COLONIA") or colonia).strip().upper()

        municipio = (tmp.get("MUNICIPIO") or tmp.get("LOCALIDAD") or municipio).strip().upper()

    # -----------------------
    # 2) Identidad / fakes
    # -----------------------
    rfc = (ci.get("RFC") or "").strip().upper()
    curp = (ci.get("CURP") or "").strip().upper()

    nombre = (ci.get("NOMBRE") or "").strip().upper()
    ap1 = (ci.get("APELLIDO_PATERNO") or ci.get("PRIMER_APELLIDO") or "").strip().upper()
    ap2 = (ci.get("APELLIDO_MATERNO") or ci.get("SEGUNDO_APELLIDO") or "").strip().upper()
    nombre_etiqueta = " ".join(x for x in [nombre, ap1, ap2] if x).strip()

    no_ext = str(_det_rand_int("NOEXT|" + seed_key, 1, 999)).strip()
    idcif_fake = str(_det_rand_int("IDCIF|" + seed_key, 10_000_000_000, 30_000_000_000)).strip()

    # -----------------------
    # 3) Fechas (MISMO formato que construir_datos_desde_apis)
    # -----------------------
    ahora = datetime.now(ZoneInfo("America/Mexico_City"))
    fecha_emision = _fecha_lugar_mun_ent(municipio, entidad)

    birth_year = _parse_birth_year(ci.get("FECHA_NACIMIENTO", ""))
    y0 = (birth_year + 18) if birth_year else (ahora.year - 5)

    d, m, y = _fake_date_components(y0, seed_key)

    fecha_inicio_raw = _fmt_dd_de_mes_de_aaaa(d, m, y)
    fecha_ultimo_raw = _fmt_dd_de_mes_de_aaaa(d, m, y)
    fecha_alta_raw   = _fmt_dd_mm_aaaa(d, m, y)

    fn_dash = _to_dd_mm_aaaa_dash(ci.get("FECHA_NACIMIENTO", ""))
    fi_dash = _to_dd_mm_aaaa_dash(fecha_inicio_raw)
    fu_dash = _to_dd_mm_aaaa_dash(fecha_ultimo_raw)
    fa_dash = _to_dd_mm_aaaa_dash(fecha_alta_raw)

    if not fn_dash:
        fn_dash = fi_dash
    if not fi_dash:
        fi_dash = fn_dash
    if not fu_dash:
        fu_dash = fi_dash
    if not fa_dash:
        fa_dash = fi_dash

    # -----------------------
    # 4) Régimen / AL
    # -----------------------
    reg_val = ci.get("REGIMEN", "")
    if isinstance(reg_val, list):
        reg_val = reg_val[0] if reg_val else ""
    reg_val = limpiar_regimen(reg_val)

    al_val = _al_from_entidad(entidad)

    # -----------------------
    # 5) Dict FINAL compatible con tu doc/WA
    # -----------------------
    datos = {
        "RFC_ETIQUETA": rfc,
        "NOMBRE_ETIQUETA": nombre_etiqueta,
        "IDCIF_ETIQUETA": idcif_fake,

        "RFC": rfc,
        "CURP": curp,
        "NOMBRE": nombre,
        "PRIMER_APELLIDO": ap1,
        "SEGUNDO_APELLIDO": ap2,

        "REGIMEN": reg_val,
        "ESTATUS": "ACTIVO",

        "FECHA_INICIO": fi_dash,
        "FECHA_ULTIMO": fu_dash,
        "FECHA_ALTA": fa_dash,

        "FECHA_INICIO_DOC": fecha_inicio_raw,
        "FECHA_ULTIMO_DOC": fecha_ultimo_raw,
        "FECHA_ALTA_DOC": fecha_alta_raw,

        "FECHA": fecha_emision,
        "FECHA_CORTA": ahora.strftime("%Y/%m/%d %H:%M:%S"),

        "CP": cp_final,
        "TIPO_VIALIDAD": "CALLE",
        "VIALIDAD": "SIN NOMBRE",

        # ✅ FIX: guarda ambas variantes (por compatibilidad con templates)
        "NO_EXTERIOR": no_ext,
        "NUMERO_EXTERIOR": no_ext,
        "NO_INTERIOR": "",
        "NUMERO_INTERIOR": "",

        "COLONIA": colonia,

        # ✅ FIX: expón ambos campos
        "MUNICIPIO": municipio,
        "LOCALIDAD": municipio,

        "ENTIDAD": entidad,

        "FECHA_NACIMIENTO": fn_dash,
        "AL": al_val,
    }

    # -----------------------
    # 6) Respeta locks / fuentes
    # -----------------------
    if ci.get("_MUN_LOCK"):
        datos["_MUN_LOCK"] = True
        if ci.get("_MUN_SOURCE"):
            datos["_MUN_SOURCE"] = ci.get("_MUN_SOURCE")

    if ci.get("_CP_SOURCE"):
        datos["_CP_SOURCE"] = ci.get("_CP_SOURCE")
    if ci.get("_REG_SOURCE"):
        datos["_REG_SOURCE"] = ci.get("_REG_SOURCE")

    return datos

def construir_datos_desde_apis(term: str) -> dict:
    term_norm = (term or "").strip().upper()
    if not term_norm:
        raise ValueError("TERM_EMPTY")

    key = f"CHECKID:{term_norm}"

    cached = cache_get(key)
    if cached:
        print("[CHECKID] CACHE_HIT", key, "term=", term_norm, flush=True)
        return cached

    print("[CHECKID] CACHE_MISS", key, "term=", term_norm, flush=True)

    # ---------- 1) CheckID ----------
    ci_raw = checkid_lookup(term_norm)
    print("[CHECKID] LOOKUP_RETURNED", "term=", term_norm, "keys=", list((ci_raw or {}).keys())[:6], flush=True)
    ci = _norm_checkid_fields(ci_raw)
    
    def _is_curp_pf(curp: str) -> bool:
        c = (curp or "").strip().upper()
        return len(c) == 18 and bool(re.match(r"^[A-Z]{4}\d{6}[A-Z]{6}\d{2}$", c))

    def _is_rfc_pf(rfc: str) -> bool:
        r = (rfc or "").strip().upper()
        return bool(re.match(r"^[A-Z&Ñ]{4}\d{6}[A-Z0-9]{3}$", r))
    
    if (ci.get("NOMBRE") and not ci.get("APELLIDO_PATERNO") and not ci.get("APELLIDO_MATERNO")):
        src = (ci.get("_NAME_SOURCE") or "").strip().upper()
        curp_pf = _is_curp_pf(ci.get("CURP", ""))
        rfc_pf = _is_rfc_pf(ci.get("RFC", ""))
    
        allow = (src != "RAZON_SOCIAL") or curp_pf or rfc_pf
    
        if allow:
            d = desglose_nombre_mex_pro(ci["NOMBRE"])
            ci["NOMBRE"] = d["NOMBRE"]
            ci["APELLIDO_PATERNO"] = d["APELLIDO_PATERNO"]
            ci["APELLIDO_MATERNO"] = d["APELLIDO_MATERNO"]

    if not (ci.get("RFC") or ci.get("CURP")):
        raise RuntimeError("CHECKID_SIN_DATOS")

    # ---------- 2) Dipomex (SOFT FAIL) + SEPOMEX fallback ----------
    dip = {}
    cp_val = re.sub(r"\D+", "", (ci.get("CP") or "")).strip()
    
    if len(cp_val) == 5:
        # 2.1) intenta Dipomex
        try:
            dip = dipomex_by_cp(cp_val) or {}
        except Exception as e:
            print("DIPOMEX FAILED (soft):", repr(e))
            dip = {}
    
        # 2.2) si Dipomex falló / vacío / sin colonias -> SEPOMEX local
        if (not dip) or (not (dip.get("colonias") or [])):
            dip = sepomex_by_cp(cp_val) or {}
            if dip:
                print("SEPOMEX fallback OK for CP:", cp_val)
            else:
                print("SEPOMEX: CP not found:", cp_val)

    # ---------- 3) Dirección (SIN duplicados / SIN pisarte) ----------
    # 3.0) de CheckID
    entidad_ci = (ci.get("ENTIDAD") or "").strip().upper()
    municipio_ci = (
        (ci.get("MUNICIPIO") or "") or
        (ci.get("LOCALIDAD") or "") or
        ""
    ).strip().upper()
    colonia_ci = (ci.get("COLONIA") or "").strip().upper()
    
    # Normaliza DF -> CDMX (si quieres)
    if entidad_ci in ("DISTRITO FEDERAL", "DF"):
        entidad_ci = "CIUDAD DE MÉXICO"
    
    # 3.1) de Dipomex/SEPOMEX (solo si hubo CP)
    entidad_dip   = (dip.get("estado") or "").strip().upper()
    municipio_dip = (dip.get("municipio") or "").strip().upper()
    colonia_dip   = (_pick_first_colonia(dip) or "").strip().upper()
    
    # 3.2) Decide entidad/municipio/colonia SIN defaults agresivos
    entidad = entidad_ci or entidad_dip
    municipio = municipio_ci or municipio_dip
    colonia = colonia_ci or colonia_dip
    
    # Si NO hay entidad, ya de plano default CDMX (último recurso)
    if not entidad:
        entidad = "CIUDAD DE MÉXICO"
    
    # Si entidad NO es CDMX, NO inventes Cuauhtémoc/Centro
    if entidad in ("CIUDAD DE MÉXICO", "CDMX"):
        if not municipio:
            municipio = "CUAUHTÉMOC"
        if not colonia:
            colonia = "CENTRO"
    else:
        # fuera de CDMX: deja vacío; después SEPOMEX (por CP) lo completa
        municipio = municipio or ""
        colonia = colonia or ""

    # Reglas fijas
    tipo_vialidad = "CALLE"
    vialidad = "SIN NOMBRE"

    # Semilla determinística
    seed_key = (ci.get("RFC") or ci.get("CURP") or term_norm).strip().upper()

    cp_final = re.sub(r"\D+", "", (ci.get("CP") or "")).strip()

    # ✅ Marca origen del CP
    if len(cp_final) == 5 and (ci.get("_CP_SOURCE") or "").strip().upper() == "CHECKID":
        ci["_CP_SOURCE"] = "CHECKID"
    
    # Si CP no vino, intenta escoger uno real dentro del estado
    if len(cp_final) != 5 and entidad:
        cp_pick = sepomex_pick_cp_by_entidad(entidad, seed_key=seed_key)
        if cp_pick:
            cp_final = cp_pick
            ci["_CP_SOURCE"] = "SEPOMEX_PICK"
    
    #  Si ya tenemos CP: forzamos que ENT/MUN/COL correspondan al CP
    if len(cp_final) == 5:
        tmp = {
            "CP": cp_final,
            "ENTIDAD": entidad,
            "MUNICIPIO": municipio,
            "LOCALIDAD": municipio,
            "COLONIA": colonia,
            "_MUN_LOCK": False,
        }
    
        cp_src = (ci.get("_CP_SOURCE") or "").strip().upper()
        force_mun = cp_src in ("SATPI", "CHECKID", "SEPOMEX_PICK")
        
        tmp = reconcile_location_by_cp(tmp, seed_key=seed_key, force_mun=force_mun)
    
        entidad = (tmp.get("ENTIDAD") or entidad).strip().upper()
        municipio = (tmp.get("MUNICIPIO") or tmp.get("LOCALIDAD") or municipio).strip().upper()
        colonia = (tmp.get("COLONIA") or colonia).strip().upper()

    no_ext = str(_det_rand_int("NOEXT|" + seed_key, 1, 999))
    idcif_fake = str(_det_rand_int("IDCIF|" + seed_key, 10_000_000_000, 30_000_000_000))

    nombre_etiqueta = " ".join(
        x for x in [ci.get("NOMBRE"), ci.get("APELLIDO_PATERNO"), ci.get("APELLIDO_MATERNO")] if x
    ).strip()

    ahora = datetime.now(ZoneInfo("America/Mexico_City"))
    fecha_emision = _fecha_lugar_mun_ent(municipio, entidad)

    # ---------- 4) Fechas (RAW primero) ----------
    birth_year = _parse_birth_year(ci.get("FECHA_NACIMIENTO", ""))
    if birth_year:
        y0 = birth_year + 18
    else:
        # fallback razonable: hoy - 5 años
        y0 = ahora.year - 5
    
    d, m, y = _fake_date_components(y0, seed_key)
    
    fecha_inicio_raw = _fmt_dd_de_mes_de_aaaa(d, m, y)
    fecha_ultimo_raw = _fmt_dd_de_mes_de_aaaa(d, m, y)
    fecha_alta_raw   = _fmt_dd_mm_aaaa(d, m, y)

    # ---------- 5) Normalización FINAL (lo importante) ----------
    al_val = _al_from_entidad(entidad)
    
    # Fechas para PÁGINA (dd-mm-aaaa)
    fn_dash = _to_dd_mm_aaaa_dash(ci.get("FECHA_NACIMIENTO", ""))
    
    # Convierte a dash desde los RAW ya calculados en el paso 4
    # OJO: fecha_inicio_raw y fecha_ultimo_raw están en "09 DE ENERO DE 2026"
    #      fecha_alta_raw está en "09/01/2026"
    fi_dash = _to_dd_mm_aaaa_dash(fecha_inicio_raw)
    fu_dash = _to_dd_mm_aaaa_dash(fecha_ultimo_raw)
    fa_dash = _to_dd_mm_aaaa_dash(fecha_alta_raw)
    
    # --- formatos PARA DOCUMENTO ---
    fi_doc = fecha_inicio_raw          # "09 DE ENERO DE 2026"
    fu_doc = fecha_ultimo_raw          # "09 DE ENERO DE 2026"
    fa_doc = fecha_alta_raw            # "09/01/2026"
    
    # fallbacks por si algo no parseó
    if not fn_dash:
        fn_dash = fi_dash

    if not fi_dash:
        fi_dash = fn_dash
    
    if not fu_dash:
        fu_dash = fi_dash
    
    if not fa_dash:
        fa_dash = fi_dash

    # Régimen: si ci trae lista o string, quédate con el primero
    # (si ci["REGIMEN"] ya viene limpio, esto solo lo asegura)
    reg_val = ci.get("REGIMEN", "")
    if isinstance(reg_val, list):
        reg_val = reg_val[0] if reg_val else ""
    reg_val = limpiar_regimen(reg_val)

    datos = {
        "RFC_ETIQUETA": (ci.get("RFC") or "").strip().upper(),
        "NOMBRE_ETIQUETA": nombre_etiqueta,
        "IDCIF_ETIQUETA": idcif_fake,

        "RFC": (ci.get("RFC") or "").strip().upper(),
        "CURP": (ci.get("CURP") or "").strip().upper(),
        "NOMBRE": (ci.get("NOMBRE") or "").strip().upper(),
        "PRIMER_APELLIDO": (ci.get("APELLIDO_PATERNO") or "").strip().upper(),
        "SEGUNDO_APELLIDO": (ci.get("APELLIDO_MATERNO") or "").strip().upper(),

        "REGIMEN": reg_val,
        "ESTATUS": "ACTIVO",

        # ✅ dd-mm-aaaa
        "FECHA_INICIO": fi_dash,        # DD-MM-AAAA
        "FECHA_ULTIMO": fu_dash,        # DD-MM-AAAA
        "FECHA_ALTA": fa_dash,          # DD-MM-AAAA

        "FECHA_INICIO_DOC": fi_doc,     # 09 DE ENERO DE 2026
        "FECHA_ULTIMO_DOC": fu_doc,     # 09 DE ENERO DE 2026
        "FECHA_ALTA_DOC": fa_doc,       # 09/01/2026

        "FECHA": fecha_emision,
        "FECHA_CORTA": ahora.strftime("%Y/%m/%d %H:%M:%S"),

        "CP": cp_final,

        "TIPO_VIALIDAD": tipo_vialidad,
        "VIALIDAD": vialidad,
        "NO_EXTERIOR": no_ext,
        "NUMERO_EXTERIOR": no_ext,
        "NO_INTERIOR": "",
        "NUMERO_INTERIOR": "",


        "COLONIA": colonia,
        "LOCALIDAD": municipio,
        "ENTIDAD": entidad,

        # ✅ dd-mm-aaaa
        "FECHA_NACIMIENTO": fn_dash,

        # ✅ requerido por ti
        "AL": al_val,
        "_CP_SOURCE": (ci.get("_CP_SOURCE") or "").strip().upper(),
    }

    # ---------- 6) Cache ----------
    cache_set(key, datos)

    if datos.get("RFC"):
        cache_set(f"CHECKID:{datos['RFC'].strip().upper()}", datos)
    if datos.get("CURP"):
        cache_set(f"CHECKID:{datos['CURP'].strip().upper()}", datos)

    return datos

# ================== DETECCIÓN INPUT TYPE + CURP ==================

CURP_REGEX = r"\b([A-Z][AEIOUX][A-Z]{2}\d{6}[HM][A-Z]{5}[A-Z0-9]\d)\b"

def extraer_curp(texto: str):
    if not texto:
        return None
    t = (texto or "").strip().upper()
    m = re.search(CURP_REGEX, t)
    return m.group(1) if m else None

def detect_input_type(text_body: str, had_image: bool, fuente_img: str = "") -> str:
    """
    Prioridad:
      - si viene de imagen y se detectó por QR/OCR => QR
      - si el texto contiene CURP => CURP
      - si contiene RFC+IDCIF => RFC_IDCIF
      - si contiene RFC (solo) => RFC_ONLY
      - default => RFC_IDCIF
    """
    if had_image and (fuente_img in ("QR", "OCR")):
        return "QR"

    curp = extraer_curp(text_body or "")
    if curp:
        return "CURP"

    rfc, idcif = extraer_rfc_idcif(text_body or "")
    if rfc and idcif:
        return "RFC_IDCIF"

    rfc_solo = extraer_rfc_solo(text_body or "")
    if rfc_solo:
        return "RFC_ONLY"

    return "RFC_IDCIF"

def make_ok_key(input_type: str, rfc: str | None, curp: str | None) -> str:
    input_type = (input_type or "").upper().strip()
    if input_type == "CURP":
        return f"CURP:{(curp or '').upper().strip()}"
    # RFC_IDCIF o QR => dedupe por RFC
    return f"RFC:{(rfc or '').upper().strip()}"

# ================== WA MSG_ID DEDUPE (PERSISTENTE) ==================

WA_MSG_TTL_DONE_SEC = int(os.getenv("WA_MSG_TTL_DONE_SEC", str(7 * 24 * 3600)))   # 7 días
WA_MSG_TTL_PROC_SEC = int(os.getenv("WA_MSG_TTL_PROC_SEC", str(10 * 60)))        # 10 minutos

def _wa_msg_key(msg_id: str) -> str:
    return f"WA_MSGID:{(msg_id or '').strip()}"

def wa_is_duplicate(msg_id: str) -> bool:
    """
    True si:
      - ya está DONE (procesado antes)
      - o está PROCESSING (se está procesando en paralelo / reintento inmediato)
    """
    mid = (msg_id or "").strip()
    if not mid:
        return False
    rec = cache_get(_wa_msg_key(mid))
    if not rec:
        return False
    st = (rec.get("status") if isinstance(rec, dict) else "") or ""
    st = str(st).upper().strip()
    return st in ("PROCESSING", "DONE")

def wa_mark_processing(msg_id: str, from_wa_id: str = ""):
    mid = (msg_id or "").strip()
    if not mid:
        return
    cache_set(
        _wa_msg_key(mid),
        {"status": "PROCESSING", "ts": _now_iso() if "_now_iso" in globals() else "", "from": (from_wa_id or "")},
        ttl=WA_MSG_TTL_PROC_SEC
    )

def wa_mark_done(msg_id: str, from_wa_id: str = ""):
    mid = (msg_id or "").strip()
    if not mid:
        return
    cache_set(
        _wa_msg_key(mid),
        {"status": "DONE", "ts": _now_iso() if "_now_iso" in globals() else "", "from": (from_wa_id or "")},
        ttl=WA_MSG_TTL_DONE_SEC
    )

def wa_unmark(msg_id: str):
    mid = (msg_id or "").strip()
    if not mid:
        return
    try:
        cache_del(_wa_msg_key(mid))
    except Exception:
        # si no existe, no pasa nada
        pass
    
# ========= CASOS CRÍTICOS: helpers =========

ERR_CURP_INVALID = "❌ La CURP ingresada no tiene un formato válido"
ERR_NO_RFC_FOR_CURP = "⚠️ No se encontró un RFC asociado a esta CURP en los sistemas de validación"
ERR_RFC_IDCIF_INVALID = "🚫 El RFC o el identificador (IDCIF) no tienen un formato válido"
ERR_SERVICE_DOWN = "🛠️ El servicio de validación no está disponible en este momento. Intenta más tarde"
ERR_SAT_NO_DATA = "⚠️ No fue posible obtener datos del SAT con ese RFC e identificador. Verifica que el identificador sea correcto"
MSG_IN_PROCESS = "⏳ Tu trámite está en proceso, espera unos momentos"

def is_valid_curp(curp: str) -> bool:
    curp = (curp or "").strip().upper()
    # usa tu CURP_RE (ya existe en tu archivo) si está arriba; si no, usa CURP_RE del inicio.
    try:
        return bool(CURP_RE.match(curp))
    except Exception:
        return False

def is_valid_rfc(rfc: str) -> bool:
    rfc = (rfc or "").strip().upper()
    try:
        return bool(RFC_RE.match(rfc))
    except Exception:
        return False

def is_valid_idcif(idcif: str) -> bool:
    s = (idcif or "").strip()
    # flexible pero evita basura: 8-20 dígitos (ajusta si tu idCIF tiene longitud fija)
    return bool(re.fullmatch(r"\d{8,20}", s))

def looks_like_user_typed_a_curp(text: str) -> bool:
    """
    Detecta intención de CURP para contestar 'formato inválido' y NO gastar API.
    """
    t = (text or "").strip().upper()
    if "CURP" in t:
        return True
    # si mandan un token de ~18 chars alfanum, suele ser intento de CURP
    tokens = [x for x in re.split(r"[\s\|\.,;:]+", t) if x]
    if len(tokens) == 1 and 16 <= len(tokens[0]) <= 19:
        return True
    return False

_RE_ALNUM = re.compile(r"[A-Z0-9]+", re.I)

def _extract_alnum_tokens(text: str):
    text = (text or "").upper()
    # captura tokens tipo CAVA5005... sin espacios
    toks = _RE_ALNUM.findall(text)
    # filtra tokens muy cortos para reducir ruido
    return [t for t in toks if len(t) >= 8]

def looks_like_user_typed_a_curp(text: str) -> bool:
    # “parece CURP” si hay un token alfanumérico largo cercano a 18 (ej 16-20)
    for t in _extract_alnum_tokens(text):
        if 14 <= len(t) <= 22:
            # heurística: CURP suele iniciar con 4 letras
            if t[:4].isalpha():
                return True
    return False

def looks_like_user_typed_an_rfc(text: str) -> bool:
    # “parece RFC” si hay token 10-15, empieza con letras y tiene números
    for t in _extract_alnum_tokens(text):
        if 10 <= len(t) <= 15 and t[:3].isalpha():
            has_digit = any(c.isdigit() for c in t)
            if has_digit:
                return True
    return False

def looks_like_user_typed_an_idcif(text: str) -> bool:
    # “parece IDCIF” si hay token numérico 9-13 (cerca de 11)
    for t in _extract_alnum_tokens(text):
        if t.isdigit() and 9 <= len(t) <= 13:
            return True
    return False

# ========= DEDUPE "EN PROCESO" (además de wa_seen_msg y ok_key) =========
_INFLIGHT_LOCK = threading.Lock()
_INFLIGHT = {}  # ok_key -> exp_epoch
_INFLIGHT_TTL_SEC = int(os.getenv("INFLIGHT_TTL_SEC", "120") or "120")

def inflight_start(ok_key: str) -> bool:
    """
    Regresa True si se pudo marcar como "en proceso".
    Regresa False si YA estaba en proceso (=> responder MSG_IN_PROCESS).
    """
    if not ok_key:
        return True
    now = int(time.time())
    with _INFLIGHT_LOCK:
        # limpia expirados
        dead = [k for k,v in _INFLIGHT.items() if int(v or 0) <= now]
        for k in dead:
            _INFLIGHT.pop(k, None)

        exp = _INFLIGHT.get(ok_key)
        if exp and int(exp) > now:
            return False

        _INFLIGHT[ok_key] = now + max(30, _INFLIGHT_TTL_SEC)
        return True

def inflight_end(ok_key: str):
    if not ok_key:
        return
    with _INFLIGHT_LOCK:
        _INFLIGHT.pop(ok_key, None)
    
def wa_upload_document(file_bytes: bytes, filename: str, mime: str):
    """
    Sube un archivo a WhatsApp y regresa media_id
    """
    if not (WA_TOKEN and WA_PHONE_NUMBER_ID):
        raise RuntimeError("Faltan WA_TOKEN o WA_PHONE_NUMBER_ID.")

    url = wa_api_url(f"{WA_PHONE_NUMBER_ID}/media")
    headers = {"Authorization": f"Bearer {WA_TOKEN}"}
    files = {
        "file": (filename, file_bytes, mime),
    }
    data = {"messaging_product": "whatsapp"}
    r = requests.post(url, headers=headers, files=files, data=data, timeout=60)

    if not r.ok:
        print("WA MEDIA UPLOAD ERROR status:", r.status_code)
        print("WA MEDIA UPLOAD ERROR body:", r.text)

    r.raise_for_status()
    return r.json().get("id")


def wa_send_document(to_wa_id: str, media_id: str, filename: str, caption: str = ""):
    """
    Envía un documento ya subido (media_id) por WhatsApp
    """
    if not (WA_TOKEN and WA_PHONE_NUMBER_ID):
        raise RuntimeError("Faltan WA_TOKEN o WA_PHONE_NUMBER_ID.")

    url = wa_api_url(f"{WA_PHONE_NUMBER_ID}/messages")
    headers = {
        "Authorization": f"Bearer {WA_TOKEN}",
        "Content-Type": "application/json",
    }
    payload = {
        "messaging_product": "whatsapp",
        "to": to_wa_id,
        "type": "document",
        "document": {
            "id": media_id,
            "filename": filename,
        },
    }
    if caption:
        payload["document"]["caption"] = caption

    r = requests.post(url, headers=headers, json=payload, timeout=30)

    if not r.ok:
        print("WA SEND DOC ERROR status:", r.status_code)
        print("WA SEND DOC ERROR body:", r.text)

    r.raise_for_status()
    return r.json()

# Si quieres limitar concurrencia:
from concurrent.futures import ThreadPoolExecutor
EXEC = ThreadPoolExecutor(max_workers=3)

# ================== BACKPRESSURE / LIMITADOR DE COLA ==================
WA_MAX_PENDING = int(os.getenv("WA_MAX_PENDING", "12"))  # máximo de jobs simultáneos/encolados
_WA_PENDING_SEM = threading.BoundedSemaphore(WA_MAX_PENDING)

def wa_try_acquire_slot() -> bool:
    try:
        return _WA_PENDING_SEM.acquire(blocking=False)
    except Exception:
        return False

def wa_release_slot():
    try:
        _WA_PENDING_SEM.release()
    except Exception:
        # si ya estaba liberado o error, no rompe flujo
        pass

def safe_submit(fn, *args, **kwargs):
    try:
        EXEC.submit(fn, *args, **kwargs)
    except Exception:
        # fallback: no tumbes el webhook
        traceback.print_exc()

@app.route("/wa/webhook", methods=["POST"])
def wa_webhook_receive():
    payload = request.get_json(silent=True) or {}
    print("WA WEBHOOK POST payload:", payload)

    try:
        entry = (payload.get("entry") or [])[0]
        changes = (entry.get("changes") or [])[0]
        value = changes.get("value") or {}

        messages = value.get("messages") or []
        if not messages:
            return "OK", 200

        msg = messages[0]
        msg_id = (msg.get("id") or "").strip()

        contacts = value.get("contacts") or []
        raw_wa_id = (contacts[0].get("wa_id") if contacts else None) or msg.get("from")
        from_wa_id = normalizar_wa_to(raw_wa_id)
        if not from_wa_id:
            return "OK", 200

        # ✅ DEDUPE PERSISTENTE (primero, baratísimo)
        if msg_id and wa_is_duplicate(msg_id):
            print("WA DUPLICATE(PERSISTENT) msg_id ignored:", msg_id)
            return "OK", 200

        # ✅ allow/block rápido
        try:
            st = get_state(STATS_PATH)
            from stats_store import is_allowed, is_blocked

            if not is_allowed(st, from_wa_id):
                print("WA NOT ALLOWED (ignored):", from_wa_id)
            
                # ✅ log en stats para auditoría
                def _ev(s):
                    from stats_store import log_event
                    log_event(s, "NOT_ALLOWED", from_wa_id, {"where": "wa_webhook"})
                try:
                    get_and_update(STATS_PATH, _ev)
                except Exception:
                    pass
            
                # ✅ respuesta opcional (actívala con env)
                if os.getenv("WA_REPLY_NOT_ALLOWED", "0") == "1":
                    try:
                        wa_send_text(from_wa_id, "⛔ Este número no está autorizado. Contacta al administrador.")
                    except Exception:
                        pass
            
                return "OK", 200
            
            if is_blocked(st, from_wa_id):
                # ✅ log en stats para auditoría
                def _ev2(s):
                    from stats_store import log_event
                    log_event(s, "BLOCKED", from_wa_id, {"where": "wa_webhook"})
                try:
                    get_and_update(STATS_PATH, _ev2)
                except Exception:
                    pass
            
                wa_send_text(from_wa_id, "⛔ Tu número está suspendido. Contacta al administrador.")
                return "OK", 200

        except Exception as e:
            print("Allow/block check error:", e)
            return "OK", 200

        # ✅ BACKPRESSURE: si no hay cupo, no encolar
        if not wa_try_acquire_slot():
            # Importante: NO marcamos PROCESSING, porque no vamos a procesar
            try:
                wa_send_text(from_wa_id, "⏳ Ahorita estoy saturado procesando solicitudes.\nIntenta de nuevo en 1 minuto.")
            except Exception:
                pass
            return "OK", 200

        # ✅ RATE LIMIT (anti-flood)
        try:
            ok_rl, why = wa_check_rate_limit(from_wa_id)
            if not ok_rl:
                if why == "COOLDOWN":
                    # no molestes con mensaje siempre (opcional)
                    if os.getenv("WA_REPLY_COOLDOWN", "0") == "1":
                        wa_send_text(from_wa_id, "⏳ Espera unos segundos y vuelve a intentar.")
                    return "OK", 200
        
                if why == "PER_MINUTE":
                    wa_send_text(from_wa_id, "⚠️ Estás enviando demasiados mensajes. Intenta de nuevo en 1 minuto.")
                    return "OK", 200
        except Exception as e:
            print("rate limit error:", e)
        
        # ✅ Marca PROCESSING antes de cualquier submit (evita dobles en ráfaga / reintentos)
        if msg_id:
            wa_mark_processing(msg_id, from_wa_id)

        # (Opcional) dedupe en memoria si lo quieres conservar (no estorba)
        if msg_id and wa_seen_msg(msg_id):
            print("WA DUPLICATE(in-memory) msg_id ignored:", msg_id)
            # ya está PROCESSING, lo dejamos y salimos
            return "OK", 200

        # ✅ marca visto el msg_id lo antes posible (si tu wa_seen_msg usa “set”/persistencia)
        try:
            if msg_id:
                wa_mark_seen(msg_id)  # si tienes helper; si no, omite
        except Exception:
            pass

        # ✅ dispara worker SIN bloquear
        job = {
            "from_wa_id": from_wa_id,
            "msg": msg,
            "value": value,
            "msg_id": msg_id,
            "received_at": time.time(),
            "bp_slot": True,
        }

        try:
            safe_submit(_process_wa_message, job)
        except Exception as e:
            print("safe_submit failed:", e)
            if msg_id:
                wa_unmark(msg_id)
            wa_release_slot()
            return "OK", 200

        return "OK", 200

    except Exception as e:
        print("Error WA webhook:", e)
        return "OK", 200

def wa_mark_seen(msg_id: str):
    if not (WA_TOKEN and WA_PHONE_NUMBER_ID and msg_id):
        return
    url = wa_api_url(f"{WA_PHONE_NUMBER_ID}/messages")
    headers = {"Authorization": f"Bearer {WA_TOKEN}", "Content-Type": "application/json"}
    payload = {"messaging_product": "whatsapp", "status": "read", "message_id": msg_id}
    r = requests.post(url, headers=headers, json=payload, timeout=20)
    if not r.ok:
        print("WA MARK SEEN ERROR:", r.status_code, r.text)

def ensure_idcif_fakey(datos: dict, seed_key: str = "") -> dict:
    """
    Garantiza que IDCIF e IDCIF_ETIQUETA existan y contengan
    SIEMPRE el valor numérico (nunca el string 'idCIF').
    """
    if datos is None:
        datos = {}

    # intenta recuperar un valor válido existente
    v = (
        datos.get("IDCIF")
        or datos.get("IDCIF_ETIQUETA")
        or datos.get("idCIF")
        or datos.get("idcif")
        or ""
    ).strip()

    # si es inválido (vacío o texto tipo 'idCIF'), genera uno nuevo
    if (not v) or (not v.isdigit()):
        if seed_key:
            v = str(_det_rand_int(
                "IDCIF|" + seed_key,
                10_000_000_000,
                30_000_000_000
            ))
        else:
            v = str(random.randint(10_000_000_000, 30_000_000_000))

    # 🔒 set consistente (mismo valor en todas)
    datos["IDCIF"] = v
    datos["IDCIF_ETIQUETA"] = v
    datos["idCIF"] = v
    datos["idcif"] = v

    print(f"[IDCIF_FAKEY_OK] IDCIF={v} ETIQUETA={v}")
    return datos

STRICT_NO_SEPOMEX_WA_IDS = {
    "528992146348",
    "527717584737",
}

def normalize_regimen_fields(datos: dict) -> dict:
    # 1) si ya existe REGIMEN/regimen, úsalo
    reg_up = (datos.get("REGIMEN") or "").strip()
    reg_lo = (datos.get("regimen") or "").strip()
    reg = reg_up or reg_lo

    # 2) si no existe, toma regimen_desc de SATPI
    if not reg:
        reg = (datos.get("regimen_desc") or "").strip()

    if reg:
        datos["REGIMEN"] = reg
        datos["regimen"] = reg

    return datos

def _strict_gate_or_abort(datos: dict, input_type: str) -> bool:
    input_type = (input_type or "").upper()

    rfc = (datos.get("RFC") or "").strip().upper()
    if not rfc:
        return False

    if bool(datos.get("_RFC_UNCONFIRMED")):
        return False

    cp_src = (datos.get("_CP_SOURCE") or "").strip().upper()
    reg_src = (datos.get("_REG_SOURCE") or "").strip().upper()

    cp_ok = cp_src in ("CHECKID", "SATPI")
    reg_ok = reg_src in ("CHECKID", "SATPI")

    if input_type == "RFC_ONLY":
        return cp_ok and reg_ok

    if input_type == "CURP":
        return True

    return cp_ok or reg_ok

def _strict_confirm_curp_with_satpi(datos: dict) -> dict:
    """
    En STRICT+CURP: si RFC está ausente o está marcado como DERIVED/UNCONFIRMED,
    intenta confirmarlo con SATPI usando el RFC actual.
    Si SATPI confirma, marca _RFC_UNCONFIRMED=False y sources SATPI.
    """
    rfc = (datos.get("RFC") or "").strip().upper()
    unconf = bool(datos.get("_RFC_UNCONFIRMED"))

    # nada qué confirmar
    if not rfc:
        return datos

    # si ya está confirmado, no hagas nada
    if not unconf:
        return datos

    satpi_d = _rfc_only_fallback_satpi(rfc) or {}

    rfc_sat = (satpi_d.get("rfc") or satpi_d.get("RFC") or "").strip().upper()
    cp_v = (satpi_d.get("cp") or satpi_d.get("CP") or "").strip()
    curp_v = (satpi_d.get("curp") or satpi_d.get("CURP") or "").strip().upper()
    reg_v = (satpi_d.get("regimen_desc") or satpi_d.get("REGIMEN") or satpi_d.get("regimen") or "").strip()

    # “confirmado” = trae RFC y al menos un dato útil
    if rfc_sat and (cp_v or curp_v or reg_v):
        datos.update(satpi_d)
        datos["RFC"] = rfc_sat
        datos["RFC_ETIQUETA"] = rfc_sat
        datos["_RFC_UNCONFIRMED"] = False
        datos["_RFC_SOURCE"] = "SATPI"

        if cp_v:
            datos["_CP_SOURCE"] = "SATPI"
        if reg_v:
            datos["_REG_SOURCE"] = "SATPI"

        datos = normalize_regimen_fields(datos)
        return datos

    # si no confirmó, deja como estaba (seguirá bloqueando en gate)
    return datos

def _norm_reg(s: str) -> str:
    s = (s or "").strip().upper()
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _regimen_no_vigente(datos: dict) -> bool:
    reg = _norm_reg(datos.get("REGIMEN") or datos.get("regimen") or "")
    if not reg:
        return True  # en strict, régimen vacío = mala señal

    # tokens ya sin acentos
    bad_tokens = (
        "NO TIENE REGIMEN VIGENTE",
        "SIN REGIMEN VIGENTE",
        "NO VIGENTE",
        "INEXISTENTE",
    )
    return any(t in reg for t in bad_tokens)

# Memoria simple por chat (no persistente): evita spamear al usuario con el mismo paso
_WA_UX_STATE = {}  # { from_wa_id: {"last_ts": float, "last_step": str, "rid": str} }

def _ux_rid(from_wa_id: str, msg_id: str = "") -> str:
    base = f"{from_wa_id}|{msg_id}|{int(time.time()//60)}"  # cambia cada minuto
    return hashlib.sha1(base.encode("utf-8")).hexdigest()[:8].upper()

def wa_step(from_wa_id: str, text: str, *, step: str, min_interval_sec: float = 3.5, force: bool = False):
    """
    Envía un mensaje de progreso al usuario, evitando repetir el mismo step y evitando spam.
    - step: identificador estable (ej. "PARSE", "VALIDATE", "CHECKID", "SATPI", "DOCS", "SEND", "DONE")
    """
    now = time.time()
    st = _WA_UX_STATE.get(from_wa_id) or {}
    last_ts = float(st.get("last_ts") or 0.0)
    last_step = str(st.get("last_step") or "")

    if (not force) and (step == last_step) and (now - last_ts) < min_interval_sec:
        return

    # si cambió step pero fue hace muy poquito, también frena un poco
    if (not force) and (step != last_step) and (now - last_ts) < 0.8:
        return

    _WA_UX_STATE[from_wa_id] = {"last_ts": now, "last_step": step, "rid": st.get("rid")}

    try:
        wa_send_text(from_wa_id, text)
    except Exception:
        # no revientes el flujo por UX
        pass

def _process_wa_message(job: dict):
    from_wa_id = job.get("from_wa_id")
    msg = job.get("msg") or {}
    msg_id = job.get("msg_id")

    err = None

    try:
        msg_type = msg.get("type")

        text_body = ""
        image_bytes = None
        fuente_img = ""

        # 1) Parse de contenido
        if msg_type == "text":
            text_body = ((msg.get("text") or {}).get("body") or "").strip()

        if msg_type == "text":
            t = (text_body or "").strip().lower()
            if t in ("hola", "ola", "buena", "buenas", "hi", "buen dia", "buen día", "buenos dias", "buenos días", "buenas tardes", "buenas noches", "hey"):
                wa_send_text(
                    from_wa_id,
                    "👋 Hola.\n\nEnvíame:\n"
                    "• RFC (13)\n"
                    "• CURP (18)\n"
                    "• RFC (12-13) + IDCIF (11)\n"
                    "• o una foto del QR\n\n"
                )
                return

        elif msg_type == "image":
            media_id = ((msg.get("image") or {}).get("id") or "").strip()
            if media_id:
                media_url = wa_get_media_url(media_id)
                image_bytes = wa_download_media_bytes(media_url)

        elif msg_type == "document":
            media_id = ((msg.get("document") or {}).get("id") or "").strip()
            mime = ((msg.get("document") or {}).get("mime_type") or "")
            if media_id and (mime.startswith("image/") or mime in ("application/octet-stream", "")):
                media_url = wa_get_media_url(media_id)
                image_bytes = wa_download_media_bytes(media_url)

        rid = _ux_rid(from_wa_id or "", msg_id or "")
        try:
            st = _WA_UX_STATE.get(from_wa_id) or {}
            st["rid"] = rid
            _WA_UX_STATE[from_wa_id] = st
        except Exception:
            pass

        # ====== DETECCIÓN AUTOMÁTICA DE JSON (MANUAL) ======
        payload = None
        text = (text_body or "").strip()

        if msg_type == "text" and text.startswith("{") and text.endswith("}"):
            try:
                obj = json.loads(text)
                # Para evitar falsas detecciones, exige RFC o CURP
                if isinstance(obj, dict) and (obj.get("RFC") or obj.get("rfc") or obj.get("CURP") or obj.get("curp")):
                    payload = obj
            except Exception:
                payload = None

        # 2) Si hay imagen, intenta extraer RFC/IDCIF
        if image_bytes:
            rfc_img, idcif_img, fuente_img = extract_rfc_idcif_from_image_bytes(image_bytes)
            if rfc_img and idcif_img:
                text_body = f"{rfc_img} {idcif_img}"
            else:
                wa_step(
                    from_wa_id,
                    "⚠️ Recibí tu imagen pero no pude leer el QR/texto.\n🧩 Te doy tips para corregirlo:",
                    step="DETECT_FAIL",
                    force=True
                )
                wa_send_text(
                    from_wa_id,
                    "• Manda la foto del QR lo más centrada posible\n"
                    "• Sin reflejos y con buena luz\n"
                    "• O escribe: RFC IDCIF\n\n"
                )
                return

        # 3) Si no hay nada, guía
        if not (text_body or "").strip():
            wa_send_text(
                from_wa_id,
                "📩 Envíame RFC e idCIF o una foto donde se vea el QR.\n\n"
                "Ejemplo texto:\nTOHJ640426XXX 19010347XXX"
            )
            return

        # 4) Detectar tipo de entrada (ROBUSTO)
        if payload is not None:
            input_type = "MANUAL"

        elif image_bytes and (fuente_img in ("QR", "OCR")):
            input_type = "QR"
            # QR siempre cae a RFC_IDCIF interno (tú ya lo manejas en extract)
        else:
            curp_tok = (extraer_curp(text_body) or "").strip().upper()
            rfc_tok, idcif_tok = extraer_rfc_idcif(text_body)
            rfc_only_tok = (extraer_rfc_solo(text_body) or "").strip().upper()

            if curp_tok and not idcif_tok:
                input_type = "CURP"
            elif rfc_tok and idcif_tok:
                input_type = "RFC_IDCIF"
            elif rfc_only_tok and not idcif_tok:
                input_type = "RFC_ONLY"
            else:
                # fallback, pero ya no es crítico
                kind, _ = classify_input_for_personas(text_body)
                if kind == "only_curp":
                    input_type = "CURP"
                elif kind == "only_rfc":
                    input_type = "RFC_ONLY"
                else:
                    input_type = "RFC_IDCIF"

        # UX: qué entendí del usuario
        tipo_humano = {
            "MANUAL": "JSON (manual)",
            "QR": "QR / imagen",
            "CURP": "CURP",
            "RFC_IDCIF": "RFC + idCIF",
            "RFC_ONLY": "Solo RFC",
        }.get(input_type, input_type)

        wa_step(
            from_wa_id,
            f"🔎 Detecté: *{tipo_humano}*\n⏳ Consultando información oficial...",
            step="DETECTED",
            force=True
        )
        
        # 5) Test mode (no cobro)
        test_mode = is_test_request(from_wa_id, "MANUAL" if payload is not None else text_body)

        # Helper: inc request (solo cuando sí vamos a consultar)
        def inc_req_if_needed():
            if test_mode:
                return
            def _inc_req(s):
                from stats_store import inc_request, inc_user_request
                inc_request(s)
                inc_user_request(s, from_wa_id)
            get_and_update(STATS_PATH, _inc_req)

        # ==========================
        # MODO LISTA RFC + IDCIF (varias líneas) + ZIP SI ES GRANDE (por link)
        # ==========================
        if payload is None and parece_lista_rfc_idcif(text_body):
            pares = extraer_lista_rfc_idcif(text_body)
        
            if not pares:
                wa_send_text(from_wa_id, "❌ No encontré pares RFC+IDCIF válidos.\nEnvíalos así (uno por línea):\nRFC IDCIF")
                return
        
            import math, csv
            from zipfile import ZipFile, ZIP_DEFLATED
        
            # ✅ límites
            MAX_BATCH = 300
        
            ZIP_THRESHOLD = 20        # para probar con 5
            CHUNK_SIZE = 6           # cuántos procesa antes de descansar
            PAUSE_BETWEEN_CHUNKS = 8 # descanso entre chunks
        
            # ✅ throttles (SAT suele bloquear si vas muy rápido)
            PER_REQUEST_SLEEP_OK = 1.2
            PER_REQUEST_SLEEP_FAIL = 2.5
        
            # ✅ retry controlado en SAT
            SAT_MAX_ATTEMPTS_PER_PAIR = 2  # 1..2 recomendado (no más)
            SAT_BACKOFF_BASE = 3.0         # segundos
        
            # ✅ corte por racha de fallos (si SAT te bloquea, ya no gastes)
            FAIL_STREAK_CUTOFF = 10
        
            if len(pares) > MAX_BATCH:
                wa_send_text(from_wa_id, f"⚠️ Me enviaste {len(pares)} pares. Máximo permitido: {MAX_BATCH}.")
                return
        
            batch_fingerprint = hashlib.sha1(
                ("\n".join([f"{r} {i}" for (r,i) in pares])).encode("utf-8")
            ).hexdigest()[:12]
            
            batch_key = make_ok_key("BATCH_RFC_IDCIF", rfc=batch_fingerprint, curp=None)
            if not inflight_start(batch_key):
                wa_send_text(from_wa_id, MSG_IN_PROCESS)
                return
        
            inc_req_if_needed()
        
            total = len(pares)
            use_zip = total >= ZIP_THRESHOLD
        
            ok = 0
            fail = 0
            fail_streak = 0
        
            # guardamos fallos para reporte
            failed_rows = []  # [(rfc, idcif, reason)]
        
            # wrapper: SAT con retry + backoff
            def _sat_fetch_with_retry(rfc: str, idcif: str):
                last_exc = None
                for attempt in range(1, SAT_MAX_ATTEMPTS_PER_PAIR + 1):
                    try:
                        return extraer_datos_desde_sat(rfc, idcif)
                    except ValueError as e:
                        # casos tipo "SIN_DATOS_SAT"
                        last_exc = e
                        # no reintentes si es "sin datos" (no va a cambiar)
                        if str(e) == "SIN_DATOS_SAT":
                            raise
                    except Exception as e:
                        last_exc = e
        
                    # backoff antes del reintento
                    if attempt < SAT_MAX_ATTEMPTS_PER_PAIR:
                        time.sleep(SAT_BACKOFF_BASE * attempt)
        
                # si agotó intentos, levanta el último error
                if last_exc:
                    raise last_exc
                raise RuntimeError("SAT_UNKNOWN")
        
            try:
                if use_zip:
                    wa_send_text(
                        from_wa_id,
                        f"✅ Recibí {total} pares RFC+IDCIF.\n"
                        f"📦 Te enviaré un enlace con un ZIP.\n"
                        f"⏳ Procesando..."
                    )
                else:
                    wa_step(
                        from_wa_id,
                        f"✅ Lote recibido.\n🧾 Folio: {rid}\n📦 Registros: {total}\n"
                        f"⏳ Procesaré por bloques para evitar bloqueo.\n"
                        f"Te iré avisando el avance.",
                        step="BATCH_START",
                        force=True
                    )
        
                with tempfile.TemporaryDirectory() as tmpdir:
                    zip_path = os.path.join(tmpdir, "constancias.zip")
        
                    zf = None
                    if use_zip:
                        zf = ZipFile(zip_path, "w", compression=ZIP_DEFLATED)
        
                    chunks = int(math.ceil(total / float(CHUNK_SIZE)))
        
                    for cidx in range(chunks):
                        start = cidx * CHUNK_SIZE
                        end = min(start + CHUNK_SIZE, total)
                        bloque = pares[start:end]
        
                        for (rfc, idcif) in bloque:
                            try:
                                # SAT
                                datos = _sat_fetch_with_retry(rfc, idcif)
                                datos = completar_campos_por_tipo(datos)
        
                                if use_zip:
                                    # genera PDF local + mete al ZIP (no manda WA individual)
                                    pdf_filename = generar_pdf_en_tmp(
                                        tmpdir=tmpdir,
                                        text_body=f"{rfc} {idcif}",
                                        datos=datos,
                                        input_type="RFC_IDCIF",
                                    )
                                    pdf_full = os.path.join(tmpdir, pdf_filename)
                                    zf.write(pdf_full, arcname=pdf_filename)
        
                                    _bill_and_log_ok(from_wa_id, "RFC_IDCIF", datos, test_mode)
                                    ok += 1
        
                                else:
                                    _generar_y_enviar_archivos(
                                        from_wa_id,
                                        f"{rfc} {idcif}",
                                        datos,
                                        "RFC_IDCIF",
                                        test_mode
                                    )
                                    ok += 1
        
                                fail_streak = 0
                                time.sleep(PER_REQUEST_SLEEP_OK + random.uniform(0.0, 0.6))
        
                            except ValueError as e:
                                fail += 1
                                fail_streak += 1
        
                                reason = str(e)
                                if reason == "SIN_DATOS_SAT":
                                    failed_rows.append((rfc, idcif, "SIN_DATOS_SAT"))
                                    if not use_zip:
                                        wa_send_text(from_wa_id, f"❌ {rfc} {idcif}: sin datos en SAT.")
                                else:
                                    failed_rows.append((rfc, idcif, f"VALUE_ERROR:{reason}"))
                                    if not use_zip:
                                        wa_send_text(from_wa_id, f"❌ {rfc} {idcif}: error ({repr(e)})")
        
                                time.sleep(PER_REQUEST_SLEEP_FAIL + random.uniform(0.0, 0.8))
        
                            except Exception as e:
                                fail += 1
                                fail_streak += 1
                                failed_rows.append((rfc, idcif, f"EXC:{type(e).__name__}"))
                                if not use_zip:
                                    wa_send_text(from_wa_id, f"❌ {rfc} {idcif}: error ({repr(e)})")
        
                                time.sleep(PER_REQUEST_SLEEP_FAIL + random.uniform(0.0, 0.8))
        
                            # ✅ si SAT te bloqueó y ya es racha larga, corta para no perder tiempo
                            if fail_streak >= FAIL_STREAK_CUTOFF:
                                if use_zip:
                                    wa_send_text(
                                        from_wa_id,
                                        f"⚠️ Detecté {fail_streak} fallos seguidos.\n"
                                        f"Probable bloqueo/limitación del SAT. Corto el lote para proteger el sistema."
                                    )
                                # fuerza salida de loops
                                break
        
                        if fail_streak >= FAIL_STREAK_CUTOFF:
                            break
        
                        # ✅ progreso (NO spam): cada 25 o al final
                        if end % 25 == 0 or end == total:
                            wa_step(
                                from_wa_id,
                                f"⏳ Avance: {end}/{total}\n✅ OK: {ok} | ❌ Fallas: {fail}",
                                step="BATCH_PROGRESS",
                                min_interval_sec=6.0
                            )

                        time.sleep(PAUSE_BETWEEN_CHUNKS)
                    
                    # ✅ si fue ZIP, meter reporte de fallos y publicar link
                    if use_zip:
                        # agrega CSV de fallos dentro del ZIP
                        csv_name = "fallidos.csv"
                        csv_path = os.path.join(tmpdir, csv_name)
                        with open(csv_path, "w", newline="", encoding="utf-8") as fcsv:
                            w = csv.writer(fcsv)
                            w.writerow(["RFC", "IDCIF", "MOTIVO"])
                            for row in failed_rows:
                                w.writerow(list(row))
                        zf.write(csv_path, arcname=csv_name)

                        try:
                            zf.close()
                        except Exception:
                            pass
        
                        zip_name = f"constancias_{total}.zip"
                        with open(zip_path, "rb") as f:
                            zip_bytes = f.read()
        
                        try:
                            url = _dl_put_bytes(zip_bytes, zip_name, ttl_sec=DL_TTL_SEC)
                        except Exception as e:
                            print("DL publish fail:", repr(e))
                            wa_send_text(from_wa_id, "⚠️ No pude generar el enlace de descarga. Intenta de nuevo.")
                            return
        
                        wa_send_text(
                            from_wa_id,
                            f"📦 Proceso terminado.\n\n"
                            f"✅ Constancias generadas: {ok}\n"
                            f"❌ Fallidas / no disponibles: {fail}\n\n"
                            f"🔗 ZIP (válido {int(DL_TTL_SEC/3600)}h):\n{url}\n\n"
                            f"📄 Incluí un archivo 'fallidos.csv' dentro del ZIP con el motivo."
                        )
                    else:
                        wa_send_text(from_wa_id, f"✅ Lote terminado.\nCorrectos: {ok}\nFallidos: {fail}")
        
                    return
        
            finally:
                inflight_end(batch_key)

        # 6) Ruteo por tipo
        if input_type == "MANUAL":
            rfc_m = (payload.get("RFC") or payload.get("rfc") or "").strip().upper()
            curp_m = (payload.get("CURP") or payload.get("curp") or "").strip().upper()
            if rfc_m and not is_valid_rfc(rfc_m):
                wa_send_text(from_wa_id, ERR_RFC_IDCIF_INVALID)
                return
            if curp_m and not is_valid_curp(curp_m):
                wa_send_text(from_wa_id, ERR_CURP_INVALID)
                return

            ok_key = make_ok_key("MANUAL", rfc=rfc_m or None, curp=curp_m or None)
            if not inflight_start(ok_key):
                wa_send_text(from_wa_id, MSG_IN_PROCESS)
                return

            inc_req_if_needed()

            try:
                wa_step(
                    from_wa_id,
                    f"🧩 Datos manuales recibidos.\nRFC: {rfc_m or '-'}\nCURP: {curp_m or '-'}\n⏳ Armando datos...",
                    step="MANUAL_BUILD",
                    force=True
                )

                datos = construir_datos_manual(payload, input_type="MANUAL")

                try:
                    pub_url = validacion_sat_publish(datos, "MANUAL")
                    if pub_url:
                        datos["QR_URL"] = pub_url
                except Exception as e:
                    print("validacion_sat_publish fail:", e)

                _generar_y_enviar_archivos(from_wa_id, text_body, datos, "MANUAL", test_mode)
                return

            finally:
                inflight_end(ok_key)

        if input_type in ("CURP", "RFC_ONLY"):

            if input_type == "CURP":
                query = (extraer_curp(text_body) or "").strip().upper()

                if not query and looks_like_user_typed_a_curp(text_body):
                    wa_send_text(from_wa_id, "La CURP ingresada no tiene un formato válido o está incompleta (debe tener 18 caracteres).")
                    return

                if not query:
                    wa_send_text(from_wa_id, "❌ No pude leer tu CURP. Envíala de nuevo (18 caracteres).")
                    return

                if len(query) != 18:
                    wa_send_text(from_wa_id, "La CURP debe tener 18 caracteres. Verifica y envíala de nuevo.")
                    return

                if not is_valid_curp(query):
                    wa_send_text(from_wa_id, ERR_CURP_INVALID)
                    return
                curp_original = query
                
            else:
                query = (extraer_rfc_solo(text_body) or "").strip().upper()

                if not query and looks_like_user_typed_an_rfc(text_body):
                    wa_send_text(from_wa_id, "El RFC ingresado parece incompleto o con formato incorrecto (debe tener 12 o 13 caracteres).")
                    return

                if not query:
                    wa_send_text(from_wa_id, "❌ No pude leer tu RFC. Envíalo de nuevo (12 o 13 caracteres).")
                    return

                if len(query) not in (12, 13):
                    wa_send_text(from_wa_id, "El RFC debe tener 12 (moral) o 13 (física) caracteres. Verifica y envíalo de nuevo.")
                    return

                if not is_valid_rfc(query):
                    wa_send_text(from_wa_id, ERR_RFC_IDCIF_INVALID)
                    return
                curp_original = ""

            ok_key = make_ok_key(
                input_type,
                rfc=query if input_type == "RFC_ONLY" else None,
                curp=query if input_type == "CURP" else None
            )
            if not inflight_start(ok_key):
                wa_send_text(from_wa_id, MSG_IN_PROCESS)
                return

            label = {
                "RFC_ONLY": "RFC",
                "RFC_IDCIF": "RFC",
                "CURP": "CURP",
                "MANUAL": "MANUAL",
            }.get(input_type, input_type)

            try:
                STRICT_NO_SEPOMEX_ESSENTIALS = (from_wa_id in STRICT_NO_SEPOMEX_WA_IDS)

                def _apply_strict(datos: dict) -> dict:
                    if not STRICT_NO_SEPOMEX_ESSENTIALS:
                        return datos
                
                    cp_src = (datos.get("_CP_SOURCE") or "").strip().upper()
                    reg_src = (datos.get("_REG_SOURCE") or "").strip().upper()
                
                    if cp_src not in ("CHECKID", "SATPI"):
                        datos["_NO_SEPOMEX_CP_PICK"] = True
                
                    if reg_src not in ("CHECKID", "SATPI"):
                        datos["_REG_UNTRUSTED"] = True
                
                    return datos

                def _curp_to_checkid_term(curp: str) -> tuple[dict, str]:
                    """
                    Intenta usar gobmx para derivar RFC (PF 13) y buscar en CheckID por RFC.
                    Si gobmx falla o viene incompleto, NO truena: regresa (gob, "").
                    """
                    try:
                        gob = gobmx_curp_scrape(curp) or {}
                    except Exception as e:
                        print("[GOBMX FAIL]", repr(e), flush=True)
                        return {}, ""
                
                    # 1) si gob ya trae RFC, úsalo
                    rfc_calc = (gob.get("RFC") or gob.get("rfc") or "").strip().upper()
                
                    # 2) si no, intenta derivarlo SOLO si hay datos mínimos
                    if not rfc_calc:
                        nombre = (gob.get("NOMBRE") or "").strip()
                        ap1 = (gob.get("PRIMER_APELLIDO") or "").strip()
                        ap2 = (gob.get("SEGUNDO_APELLIDO") or "").strip()
                        fn_raw = (gob.get("FECHA_NACIMIENTO") or "").strip()
                
                        # normaliza fecha a yyyy-mm-dd
                        fecha_iso = ""
                        m = re.match(r"^(\d{2})/(\d{2})/(\d{4})$", fn_raw)
                        if m:
                            fecha_iso = f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
                        else:
                            m = re.match(r"^(\d{2})-(\d{2})-(\d{4})$", fn_raw)
                            if m:
                                fecha_iso = f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
                            else:
                                m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", fn_raw)
                                if m:
                                    fecha_iso = f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
                
                        try:
                            # ✅ OJO: exige ap1 o ap2 para no reventar rfc_pf_13
                            if fecha_iso and nombre and (ap1 or ap2):
                                rfc_calc = rfc_pf_13(nombre, ap1, ap2, fecha_iso).strip().upper()
                        except Exception as e:
                            print("[CURP->RFC DERIVE SKIP]", repr(e), "curp=", curp, flush=True)
                            rfc_calc = ""
                
                    # valida final
                    if not rfc_calc or not is_valid_rfc(rfc_calc):
                        gob["_RFC_DERIVE_FAIL"] = True
                        return gob, ""
                
                    gob["RFC"] = rfc_calc
                    return gob, rfc_calc

                def _merge_gob_into_datos(datos: dict, gob: dict, curp: str) -> dict:
                    datos = datos or {}
                    datos["CURP"] = curp
                
                    # ✅ municipio/entidad "reales" de gob
                    ent_g = (gob.get("ENTIDAD") or "").strip().upper()
                    mun_g = (gob.get("LOCALIDAD") or gob.get("MUNICIPIO") or "").strip().upper()
                
                    if ent_g:
                        datos["ENTIDAD"] = ent_g
                        datos["_ENT_SOURCE"] = datos.get("_ENT_SOURCE") or "GOBMX"
                
                    if mun_g:
                        datos["MUNICIPIO"] = mun_g
                        datos["LOCALIDAD"] = mun_g
                        datos["_MUN_LOCK"] = True
                        datos["_MUN_SOURCE"] = "GOBMX"
                
                    # ✅ nombres/fecha SOLO si faltan
                    for k in ("NOMBRE", "PRIMER_APELLIDO", "SEGUNDO_APELLIDO", "FECHA_NACIMIENTO"):
                        gv = (gob.get(k) or "").strip()
                        if gv and not (datos.get(k) or "").strip():
                            datos[k] = gv
                
                    return datos

                checkid_term = query
                gob = None
                
                try:   
                    if input_type == "CURP":
                        gob, rfc_calc = _curp_to_checkid_term(curp_original)
                    
                        if rfc_calc:
                            checkid_term = rfc_calc 
                        else:
                            checkid_term = curp_original 

                    print("[CHECKID SEARCH TERM]", "input_type=", input_type, "term=", checkid_term, flush=True)
                    
                    datos = construir_datos_desde_apis(checkid_term)  
                    datos = normalize_regimen_fields(datos)
                
                    if input_type == "CURP" and gob is not None:
                        datos = _merge_gob_into_datos(datos, gob, curp_original)
                
                    datos = _apply_strict(datos)
                
                    if from_wa_id in ("523322003600", "523338999216"):
                        REGIMEN_FIJO = "Régimen de Sueldos y Salarios e Ingresos Asimilados a Salarios"
                        datos["REGIMEN"] = REGIMEN_FIJO
                        datos["regimen"] = REGIMEN_FIJO

                except (RuntimeError, ValueError) as e:
                    se = str(e)
                
                    handled = False
                    fatal_no_data = False
                
                    # ==========================
                    # 0) MENSAJES CLAROS CHECKID (al usuario)
                    # ==========================
                    CHECKID_MSG = {
                        # ❌ Errores de input (FINAL → siempre return)
                        "CHECKID_E100": "❌ No recibí un término de búsqueda. Envía tu CURP o RFC completo.",
                        "CHECKID_E101": "❌ El dato no parece un CURP o RFC válido. Verifica y envíalo de nuevo.",
                    
                        # ⚠️ NO encontrado en CheckID (PARCIAL → permite fallback)
                        #"CHECKID_E200": "⚠️ No se encontró información en la fuente principal. Estoy intentando otra fuente...",
                        #"CHECKID_E202": "⚠️ No se encontró información en la fuente principal. Estoy intentando otra fuente...",
                    
                        # ⚠️ Reintentables (PARCIAL)
                        #"CHECKID_E201": "⚠️ El servicio no respondió correctamente. Intentando otra fuente...",
                    
                        # ⚠️ Problemas de servicio / cuota (PARCIAL)
                        #"CHECKID_E900": "⚠️ El servicio bloqueó temporalmente la conexión. Intentando otra fuente...",
                        #"CHECKID_E901": "⚠️ Sin acceso a la fuente principal. Intentando otra fuente...",
                        #"CHECKID_E902": "⚠️ Se agotaron las consultas de la fuente principal. Intentando otra fuente...",
                        #"CHECKID_E903": "⚠️ Límite alcanzado en la fuente principal. Intentando otra fuente...",
                        #"CHECKID_CIRCUIT_OPEN": "⚠️ El servicio está saturado. Intentando otra fuente...",
                    }

                    # ==========================
                    # SATPI normalizados (si tu bloque interno relanza estos)
                    # ==========================
                    if se == "SATPI_RFC_INVALID":
                        wa_send_text(from_wa_id, "⚠️ La CURP parece inválida (no pude validar el RFC derivado). Verifica y vuelve a intentarlo.")
                        return
                
                    if se == "SATPI_NO_QUOTA":
                        wa_send_text(from_wa_id, "⚠️ En este momento el servicio de validación está sin consultas disponibles. Intenta más tarde.")
                        return
                
                    if se == "SATPI_TEMP":
                        wa_send_text(from_wa_id, "⚠️ El servicio de validación está saturado o tardando en responder. Intenta de nuevo en 2-3 minutos.")
                        return
                
                    if se == "SATPI_NOT_FOUND":
                        wa_send_text(from_wa_id, "❌ No se encontró un RFC asociado a esta CURP. Verifica la CURP y vuelve a intentarlo.")
                        return
                
                    if se == "RFC_CANDIDATE_EMPTY":
                        wa_send_text(from_wa_id, "❌ No pude derivar el RFC desde la CURP por falta de datos base. Verifica la CURP y vuelve a intentarlo.")
                        return
                
                    if se == "SATPI_UNEXPECTED":
                        wa_send_text(from_wa_id, "⚠️ Ocurrió un error interno validando el RFC. Intenta de nuevo.")
                        return

                    if se == "GOBMX_RFC_DERIVE_FAIL":
                        wa_send_text(from_wa_id, "⚠️ No pude derivar el RFC para esta CURP. Intenta de nuevo o envía tu RFC.")
                        return
                
                    # Códigos donde NO conviene intentar “armar algo” desde CheckID: mejor brincar a fuentes alternas
                    CHECKID_HARD_FALLBACK = {
                        "CHECKID_E900", "CHECKID_E901", "CHECKID_E902", "CHECKID_E903", "CHECKID_CIRCUIT_OPEN",
                    }
                
                    # ============================================================
                    # A) SI ES ERROR CHECKID Y ES CURP → ruta CURP con fallback real
                    # ============================================================
                    if input_type == "CURP" and se.startswith("CHECKID_"):
                
                        if se in CHECKID_MSG:
                            wa_send_text(from_wa_id, CHECKID_MSG[se])
                
                        # 2) Si CheckID está caído / sin cuota / bloqueado → NO uses CheckID
                        if se in CHECKID_HARD_FALLBACK:
                            try:
                                fallback = gobmx_curp_scrape(curp_original)
                                fallback = enrich_curp_with_rfc_and_satpi(fallback) 

                                # ✅ Si SATPI metió CP, amarra ENT/MUN/COL a ese CP (SEPOMEX manda municipio por CP)
                                try:
                                    seed_key2 = (fallback.get("RFC") or fallback.get("CURP") or query).strip().upper()
                                    cp2 = re.sub(r"\D+", "", (fallback.get("CP") or "")).strip()
                                
                                    if len(cp2) == 5:
                                        cp_src2 = (fallback.get("_CP_SOURCE") or "").strip().upper()
                                        force_mun2 = cp_src2 in ("SATPI", "CHECKID", "SEPOMEX_PICK")
                                
                                        tmp2 = {
                                            "CP": cp2,
                                            "ENTIDAD": (fallback.get("ENTIDAD") or "").strip().upper(),
                                            "MUNICIPIO": ((fallback.get("MUNICIPIO") or fallback.get("LOCALIDAD") or "")).strip().upper(),
                                            "LOCALIDAD": ((fallback.get("LOCALIDAD") or fallback.get("MUNICIPIO") or "")).strip().upper(),
                                            "COLONIA": (fallback.get("COLONIA") or "").strip().upper(),
                                            "_MUN_LOCK": False,  # ✅ GOBMX no lock para domicilio fiscal
                                        }
                                
                                        tmp2 = reconcile_location_by_cp(tmp2, seed_key=seed_key2, force_mun=force_mun2)
                                
                                        fallback["ENTIDAD"] = (tmp2.get("ENTIDAD") or fallback.get("ENTIDAD") or "").strip().upper()
                                        mun2 = (tmp2.get("MUNICIPIO") or tmp2.get("LOCALIDAD") or "").strip().upper()
                                        if mun2:
                                            fallback["MUNICIPIO"] = mun2
                                            fallback["LOCALIDAD"] = mun2
                                        col2 = (tmp2.get("COLONIA") or "").strip().upper()
                                        if col2:
                                            fallback["COLONIA"] = col2
                                
                                        fallback["_DIR_RECONCILED_AFTER_SATPI"] = True
                                except Exception as e_re:
                                    print("reconcile after satpi fail:", repr(e_re), flush=True)

                                datos = fallback
                                datos = normalize_regimen_fields(datos)
                                datos = _apply_strict(datos)

                                seed_key = (datos.get("RFC") or datos.get("CURP") or query).strip().upper()
                                datos = ensure_default_status_and_dates(datos, seed_key=seed_key)
                
                                handled = True
                            except Exception as e_gob:
                                print("CURP fallback gobmx+satpi fail:", repr(e_gob), flush=True)
                                wa_send_text(
                                    from_wa_id,
                                    "⚠️ El sistema no estuvo disponible y el respaldo también falló.\n"
                                    "Intenta de nuevo en 2-3 minutos."
                                )
                                return
                
                        else:
                            # 3) CheckID devolvió “sin datos” o reintentable: intenta mínimo desde CheckID,
                            #    y además intenta municipio con gob.mx (tu flujo actual)
                            try:
                                datos = construir_datos_desde_checkid_curp_sin_rfc(curp_original)
                                datos = normalize_regimen_fields(datos)
                                datos = _apply_strict(datos)

                                seed_key = (datos.get("RFC") or datos.get("CURP") or query).strip().upper()
                                datos = ensure_default_status_and_dates(datos, seed_key=seed_key)
                            
                            except Exception as e2:
                                print("soft-curp from checkid fail:", repr(e2), flush=True)
                
                                # En vez de rendirte, intenta gobmx+satpi también (mejora E200)
                                try:   
                                    fallback = gobmx_curp_scrape(curp_original)
                                    fallback = enrich_curp_with_rfc_and_satpi(fallback)

                                    # ✅ Si SATPI metió CP, amarra ENT/MUN/COL a ese CP (SEPOMEX manda municipio por CP)
                                    try:
                                        seed_key2 = (fallback.get("RFC") or fallback.get("CURP") or query).strip().upper()
                                        cp2 = re.sub(r"\D+", "", (fallback.get("CP") or "")).strip()
                                    
                                        if len(cp2) == 5:
                                            cp_src2 = (fallback.get("_CP_SOURCE") or "").strip().upper()
                                            force_mun2 = cp_src2 in ("SATPI", "CHECKID", "SEPOMEX_PICK")
                                    
                                            tmp2 = {
                                                "CP": cp2,
                                                "ENTIDAD": (fallback.get("ENTIDAD") or "").strip().upper(),
                                                "MUNICIPIO": ((fallback.get("MUNICIPIO") or fallback.get("LOCALIDAD") or "")).strip().upper(),
                                                "LOCALIDAD": ((fallback.get("LOCALIDAD") or fallback.get("MUNICIPIO") or "")).strip().upper(),
                                                "COLONIA": (fallback.get("COLONIA") or "").strip().upper(),
                                                "_MUN_LOCK": False,  # ✅ GOBMX no lock para domicilio fiscal
                                            }
                                    
                                            tmp2 = reconcile_location_by_cp(tmp2, seed_key=seed_key2, force_mun=force_mun2)
                                    
                                            fallback["ENTIDAD"] = (tmp2.get("ENTIDAD") or fallback.get("ENTIDAD") or "").strip().upper()
                                            mun2 = (tmp2.get("MUNICIPIO") or tmp2.get("LOCALIDAD") or "").strip().upper()
                                            if mun2:
                                                fallback["MUNICIPIO"] = mun2
                                                fallback["LOCALIDAD"] = mun2
                                            col2 = (tmp2.get("COLONIA") or "").strip().upper()
                                            if col2:
                                                fallback["COLONIA"] = col2
                                    
                                            fallback["_DIR_RECONCILED_AFTER_SATPI"] = True
                                    except Exception as e_re:
                                        print("reconcile after satpi fail:", repr(e_re), flush=True)
                                    
                                    datos = fallback
                                    datos = normalize_regimen_fields(datos)
                                    datos = _apply_strict(datos)

                                    seed_key = (datos.get("RFC") or datos.get("CURP") or query).strip().upper()
                                    datos = ensure_default_status_and_dates(datos, seed_key=seed_key)
                                
                                except Exception as e_gob:
                                    print("CURP fallback gobmx+satpi fail after soft:", repr(e_gob), flush=True)
                                    wa_send_text(
                                        from_wa_id,
                                        "❌ No pude validar tu CURP por el momento."
                                    )
                                    return
                
                            # 4) Municipio real con gob.mx + repick CP (tu lógica)
                            try:
                                graw = consultar_curp_bot(curp_original) or {}
                                mun = (
                                    graw.get("MUNICIPIO_REGISTRO") or
                                    graw.get("MUNICIPIO") or
                                    graw.get("LOCALIDAD") or
                                    graw.get("MUNICIPIO_NACIMIENTO") or
                                    ""
                                ).strip().upper()
                                
                                ent_gob = (graw.get("ENTIDAD_REGISTRO") or graw.get("ENTIDAD") or "").strip().upper()
                                ent = (datos.get("ENTIDAD") or "").strip().upper()

                                cp_now = re.sub(r"\D+", "", (datos.get("CP") or "")).strip()
                                # ✅ Si ya hay CP válido, NO uses pistas GOBMX (domicilio fiscal manda por CP)
                                if len(cp_now) == 5:
                                    # nada que hacer aquí
                                    raise RuntimeError("SKIP_GOB_HINTS_HAS_CP")
                                
                                # ✅ Si NO hay CP válido, entonces sí se permiten pistas GOBMX
                                if ent_gob:
                                    ent = ent_gob
                                    datos["ENTIDAD"] = ent

                                if mun:
                                    # ✅ Tu política: GOBMX NO es prioridad para domicilio fiscal.
                                    # Solo úsalo como "pista" si NO hay municipio aún, y NO lockees.
                                    if not ((datos.get("MUNICIPIO") or datos.get("LOCALIDAD") or "").strip()):
                                        datos["MUNICIPIO"] = mun
                                        datos["LOCALIDAD"] = mun
                                        datos["_MUN_SOURCE"] = "GOBMX"
                                
                                    # ✅ Si NO hay CP válido, entonces sí: usa ENT+MUN para inventar CP (SEPOMEX_PICK)
                                    cp_now = re.sub(r"\D+", "", (datos.get("CP") or "")).strip()
                                    if len(cp_now) != 5:
                                        cp_new = ""
                                        if (not STRICT_NO_SEPOMEX_ESSENTIALS) and ent:
                                            cp_new = sepomex_pick_cp_by_ent_mun(
                                                ent,
                                                mun,
                                                seed_key=(datos.get("RFC") or datos.get("CURP") or query).strip().upper()
                                            )
                                
                                        if cp_new:
                                            datos["CP"] = cp_new
                                            datos["_CP_SOURCE"] = "SEPOMEX_PICK"
                                
                                            # ✅ Ya con CP, fuerza ENT/MUN/COL al CP (SEPOMEX manda municipio por CP)
                                            try:
                                                seed_key2 = (datos.get("RFC") or datos.get("CURP") or query).strip().upper()
                                                tmp2 = {
                                                    "CP": cp_new,
                                                    "ENTIDAD": (datos.get("ENTIDAD") or "").strip().upper(),
                                                    "MUNICIPIO": ((datos.get("MUNICIPIO") or datos.get("LOCALIDAD") or "")).strip().upper(),
                                                    "LOCALIDAD": ((datos.get("LOCALIDAD") or datos.get("MUNICIPIO") or "")).strip().upper(),
                                                    "COLONIA": (datos.get("COLONIA") or "").strip().upper(),
                                                    "_MUN_LOCK": False,
                                                }
                                                cp_src2 = (datos.get("_CP_SOURCE") or "").strip().upper()
                                                force_mun2 = cp_src2 in ("SATPI", "CHECKID", "SEPOMEX_PICK")
                                                tmp2 = reconcile_location_by_cp(tmp2, seed_key=seed_key2, force_mun=force_mun2)
                                
                                                datos["ENTIDAD"] = (tmp2.get("ENTIDAD") or datos.get("ENTIDAD") or "").strip().upper()
                                                mun2 = (tmp2.get("MUNICIPIO") or tmp2.get("LOCALIDAD") or "").strip().upper()
                                                if mun2:
                                                    datos["MUNICIPIO"] = mun2
                                                    datos["LOCALIDAD"] = mun2
                                                col2 = (tmp2.get("COLONIA") or "").strip().upper()
                                                if col2:
                                                    datos["COLONIA"] = col2
                                            except Exception as e_re2:
                                                print("repick/reconcile after cp_new fail:", repr(e_re2), flush=True)
                                        else:
                                            print(f"[WARN] NO CP FOR ENT+MUN. ENTIDAD={ent} MUN={mun} seed={query}", flush=True)

                            except Exception as e3:
                                if str(e3) != "SKIP_GOB_HINTS_HAS_CP":
                                    print("consultar_curp_bot fail (municipio):", repr(e3), flush=True)
                
                            handled = True
                
                    # ============================================================
                    # B) SI ES CHECKID Y ES RFC_ONLY → SATPI (y mensaje claro)
                    # ============================================================
                    elif input_type == "RFC_ONLY" and se.startswith("CHECKID_"):
                
                        if se in CHECKID_MSG:
                            wa_send_text(from_wa_id, CHECKID_MSG[se])
                
                        try:    
                            sat = _rfc_only_fallback_satpi(query) or {}
                            datos = normalize_satpi_rfc_only(sat, rfc_query=query)

                            if not ((datos.get("CP") or "").strip() or (datos.get("REGIMEN") or "").strip() or (datos.get("CURP") or "").strip()):
                                wa_send_text(from_wa_id, "❌ No se encontró información oficial para ese RFC.")
                                return
                            
                            datos = normalize_regimen_fields(datos)
                            datos = _apply_strict(datos)
                            seed_key = (datos.get("RFC") or datos.get("CURP") or query).strip().upper()
                            datos = ensure_default_status_and_dates(datos, seed_key=seed_key)
                            
                            handled = True

                        except RuntimeError as e_sat:
                            code = str(e_sat)
                
                            if code == "SATPI_412":
                                wa_send_text(from_wa_id, "⚠️ El servicio de validación está sin consultas disponibles.\nIntenta más tarde.")
                                return
                            if code in ("SATPI_428", "SATPI_RFC_LEN"):
                                wa_send_text(from_wa_id, "❌ El RFC parece inválido o incompleto.\nVerifica y envíalo de nuevo (12 o 13 caracteres).")
                                return
                            if code in ("SATPI_NOT_FOUND", "SATPI_NO_DATA"):
                                wa_send_text(from_wa_id, "❌ No se encontró información para ese RFC.\nVerifica que esté bien escrito.")
                                return
                            if code.startswith("SATPI_NET:") or code.startswith("SATPI_BAD:"):
                                wa_send_text(from_wa_id, "⚠️ El sistema no respondió correctamente.\nIntenta de nuevo en 2-3 minutos.")
                                return
                
                            wa_send_text(from_wa_id, "⚠️ Ocurrió un problema consultando datos.\nIntenta de nuevo en 2-3 minutos.")
                            return
                
                        except Exception as e_sat2:
                            print("SATPI fallback fail:", repr(e_sat2), flush=True)
                            wa_send_text(from_wa_id, "⚠️ Ocurrió un problema consultando datos.\nIntenta de nuevo en 2-3 minutos.")
                            return
                
                    # ============================================================
                    # C) CHECKID_* para otros tipos → mensaje y corta
                    # ============================================================
                    elif se.startswith("CHECKID_"):
                        if se in CHECKID_MSG:
                            wa_send_text(from_wa_id, CHECKID_MSG[se])
                        else:
                            wa_send_text(from_wa_id, "⚠️ El servicio está saturado.\nIntenta de nuevo en 2-3 minutos.")
                        return
                
                    # ✅ Si ya manejamos el error, salimos del except SIN relanzar
                    if handled:
                        pass
                    else:
                        # bug real, re-lanza
                        raise

                except requests.exceptions.Timeout:
                    if input_type == "RFC_ONLY":
                        try:
                            datos = _rfc_only_fallback_satpi(query)
                            datos = normalize_regimen_fields(datos)
                        except Exception:
                            wa_send_text(from_wa_id, "⚠️ No pude obtener datos oficiales para ese RFC.")
                            return

                    elif input_type == "CURP":
                        try:   
                            fallback = gobmx_curp_scrape(curp_original)                 # usa consultar_curp_bot
                            fallback = enrich_curp_with_rfc_and_satpi(fallback) # calcula RFC13 + SATPI
                            datos = fallback
                            datos = normalize_regimen_fields(datos)
                            datos = _apply_strict(datos)

                            # ✅ Si SATPI trae CP, el CP manda: recalcular ENT/MUN/COL desde SEPOMEX
                            cp_sat = re.sub(r"\D+", "", (datos.get("CP") or datos.get("cp") or "")).strip()
                            if len(cp_sat) == 5:
                                meta = sepomex_by_cp(cp_sat) or {}
                                ent_meta = (meta.get("estado") or "").strip().upper()
                                mun_meta = (meta.get("municipio") or "").strip().upper()
                            
                                if ent_meta:
                                    datos["ENTIDAD"] = ent_meta
                                    
                                mun_lock = bool(datos.get("_MUN_LOCK"))
                                if mun_meta and (not mun_lock):
                                    datos["LOCALIDAD"] = mun_meta
                                    datos["MUNICIPIO"] = datos.get("MUNICIPIO") or mun_meta
                            
                                seed_key = (datos.get("RFC") or datos.get("CURP") or "").strip().upper()
                                col_pick = sepomex_pick_colonia_by_cp(cp_sat, seed_key=seed_key)
                                if col_pick:
                                    datos["COLONIA"] = col_pick.strip().upper()
                            
                                datos["CP"] = cp_sat
        
                            print(
                                "[SATPI RAW]",
                                "REGIMEN=", datos.get("regimen"),
                                "| REGIMEN_UP=", datos.get("REGIMEN"),
                                "| CP=", datos.get("CP"),
                                "| COLONIA=", datos.get("COLONIA"),
                                "| MUNICIPIO=", datos.get("MUNICIPIO"),
                                "| LOCALIDAD=", datos.get("LOCALIDAD"),
                                "| ENTIDAD=", datos.get("ENTIDAD"),
                            )
                        except Exception as e2:
                            code = str(e2)
                            print("CURP fallback (gob+satpi) FAIL:", repr(e2))
                            if code == "SATPI_412":
                                wa_send_text(from_wa_id, "⚠️ Sin consultas disponibles.")
                                return
                            if code == "SATPI_428":
                                wa_send_text(from_wa_id, "⚠️ RFC inválido (tamaño no válido).")
                                return
                            wa_send_text(
                                from_wa_id,
                                "⚠️ El servicio principal no respondió a tiempo y el respaldo también falló.\n"
                                "Intenta nuevamente en 2-3 minutos."
                            )
                            return
                    else:
                        wa_send_text(from_wa_id, "⚠️ El servicio de validación no respondió a tiempo.\nIntenta nuevamente en 2-3 minutos.")
                        return
                    
                except requests.exceptions.ConnectionError:
                    if input_type == "RFC_ONLY":
                        try:
                            datos = _rfc_only_fallback_satpi(query)
                            datos = normalize_regimen_fields(datos)
                        except Exception:
                            wa_send_text(from_wa_id, "⚠️ No pude obtener datos oficiales para ese RFC.")
                            return
                    
                    elif input_type == "CURP":
                        try:
                            fallback = gobmx_curp_scrape(curp_original)
                            fallback = enrich_curp_with_rfc_and_satpi(fallback)
                            datos = fallback
                            datos = normalize_regimen_fields(datos)
                            datos = _apply_strict(datos)

                            # ✅ Si SATPI trae CP, el CP manda: recalcular ENT/MUN/COL desde SEPOMEX
                            cp_sat = re.sub(r"\D+", "", (datos.get("CP") or datos.get("cp") or "")).strip()
                            if len(cp_sat) == 5:
                                meta = sepomex_by_cp(cp_sat) or {}
                                ent_meta = (meta.get("estado") or "").strip().upper()
                                mun_meta = (meta.get("municipio") or "").strip().upper()
                            
                                if ent_meta:
                                    datos["ENTIDAD"] = ent_meta
                                    
                                mun_lock = bool(datos.get("_MUN_LOCK"))
                                if mun_meta and (not mun_lock):
                                    datos["LOCALIDAD"] = mun_meta
                                    datos["MUNICIPIO"] = datos.get("MUNICIPIO") or mun_meta
                            
                                seed_key = (datos.get("RFC") or datos.get("CURP") or "").strip().upper()
                                col_pick = sepomex_pick_colonia_by_cp(cp_sat, seed_key=seed_key)
                                if col_pick:
                                    datos["COLONIA"] = col_pick.strip().upper()
                            
                                datos["CP"] = cp_sat
        
                            print(
                                "[SATPI RAW]",
                                "REGIMEN=", datos.get("regimen"),
                                "| REGIMEN_UP=", datos.get("REGIMEN"),
                                "| CP=", datos.get("CP"),
                                "| COLONIA=", datos.get("COLONIA"),
                                "| MUNICIPIO=", datos.get("MUNICIPIO"),
                                "| LOCALIDAD=", datos.get("LOCALIDAD"),
                                "| ENTIDAD=", datos.get("ENTIDAD"),
                            )
                        except Exception as e2:
                            code = str(e2)
                            print("CURP fallback (gob+satpi) FAIL:", repr(e2))
                            if code == "SATPI_412":
                                wa_send_text(from_wa_id, "⚠️ Sin consultas disponibles.")
                                return
                            if code == "SATPI_428":
                                wa_send_text(from_wa_id, "⚠️ RFC inválido (tamaño no válido).")
                                return
                            wa_send_text(from_wa_id, "⚠️ No pude conectar con el servicio principal y el respaldo también falló.\nIntenta nuevamente en unos minutos.")
                            return
                    else:
                        wa_send_text(from_wa_id, "⚠️ No pude conectar con el servicio de validación.\nIntenta nuevamente en unos minutos.")
                        return
                    
                except requests.exceptions.RequestException:    
                    if input_type == "RFC_ONLY":
                        try:
                            datos = _rfc_only_fallback_satpi(query)
                            datos = normalize_regimen_fields(datos)
                        except Exception:
                            wa_send_text(from_wa_id, "⚠️ No pude obtener datos oficiales para ese RFC.")
                            return
                        
                    elif input_type == "CURP":
                        try:
                            fallback = gobmx_curp_scrape(curp_original)
                            fallback = enrich_curp_with_rfc_and_satpi(fallback)
                            datos = fallback
                            datos = normalize_regimen_fields(datos)

                            datos = _apply_strict(datos)

                            # ✅ Si SATPI trae CP, el CP manda: recalcular ENT/MUN/COL desde SEPOMEX
                            cp_sat = re.sub(r"\D+", "", (datos.get("CP") or datos.get("cp") or "")).strip()
                            if len(cp_sat) == 5:
                                meta = sepomex_by_cp(cp_sat) or {}
                                ent_meta = (meta.get("estado") or "").strip().upper()
                                mun_meta = (meta.get("municipio") or "").strip().upper()
                            
                                if ent_meta:
                                    datos["ENTIDAD"] = ent_meta
                                    
                                mun_lock = bool(datos.get("_MUN_LOCK"))
                                if mun_meta and (not mun_lock):
                                    datos["LOCALIDAD"] = mun_meta
                                    datos["MUNICIPIO"] = datos.get("MUNICIPIO") or mun_meta
                            
                                seed_key = (datos.get("RFC") or datos.get("CURP") or "").strip().upper()
                                col_pick = sepomex_pick_colonia_by_cp(cp_sat, seed_key=seed_key)
                                if col_pick:
                                    datos["COLONIA"] = col_pick.strip().upper()
                            
                                datos["CP"] = cp_sat
        
                            print(
                                "[SATPI RAW]",
                                "REGIMEN=", datos.get("regimen"),
                                "| REGIMEN_UP=", datos.get("REGIMEN"),
                                "| CP=", datos.get("CP"),
                                "| COLONIA=", datos.get("COLONIA"),
                                "| MUNICIPIO=", datos.get("MUNICIPIO"),
                                "| LOCALIDAD=", datos.get("LOCALIDAD"),
                                "| ENTIDAD=", datos.get("ENTIDAD"),
                            )
                        except Exception as e2:
                            code = str(e2)
                            print("CURP fallback (gob+satpi) FAIL:", repr(e2))
                            if code == "SATPI_412":
                                wa_send_text(from_wa_id, "⚠️ Sin consultas disponibles.")
                                return
                            if code == "SATPI_428":
                                wa_send_text(from_wa_id, "⚠️ RFC inválido (tamaño no válido).")
                                return
                            wa_send_text(from_wa_id, "⚠️ Ocurrió un problema temporal y el respaldo también falló.\nIntenta nuevamente en 2-3 minutos.")
                            return
                    else:
                        wa_send_text(from_wa_id, "⚠️ Ocurrió un problema temporal consultando el servicio.\nIntenta nuevamente en 2-3 minutos.")
                        return

                try:
                    seed_key = (datos.get("RFC") or datos.get("CURP") or query).strip().upper()
                    datos = ensure_default_status_and_dates(datos, seed_key=seed_key)
                except Exception as e:
                    print("ensure_default_status_and_dates fail:", repr(e), flush=True)
                
                # ✅ estado actual (después de ensure_default_status_and_dates)
                rfc_obtenido = (datos.get("RFC") or "").strip().upper()
                
                # ============================================================
                # STRICT: si CURP no trajo RFC, intenta confirmar con SATPI
                # (pero NO dejes que RuntimeError salga al handler genérico)
                # ============================================================
                if input_type == "CURP" and (not rfc_obtenido) and STRICT_NO_SEPOMEX_ESSENTIALS:
                    try:
                        # 1) intenta calcular RFC candidato (13)
                        rfc_candidato = ""
                        try:
                            fn_raw = (datos.get("FECHA_NACIMIENTO") or "").strip()
                            # dd-mm-aaaa -> yyyy-mm-dd
                            m = re.match(r"^(\d{2})-(\d{2})-(\d{4})$", fn_raw)
                            fecha_iso = f"{m.group(3)}-{m.group(2)}-{m.group(1)}" if m else ""
                
                            if fecha_iso:
                                rfc_candidato = rfc_pf_13(
                                    (datos.get("NOMBRE") or ""),
                                    (datos.get("PRIMER_APELLIDO") or ""),
                                    (datos.get("SEGUNDO_APELLIDO") or ""),
                                    fecha_iso
                                ).strip().upper()
                        except Exception as e:
                            print("RFC candidate calc fail:", repr(e))
                
                        # 2) validar candidato en SATPI (solo si se pudo calcular)
                        if not rfc_candidato:
                            raise RuntimeError("RFC_CANDIDATE_EMPTY")
                
                        satpi_d = _rfc_only_fallback_satpi(rfc_candidato) or {}
                
                        # señales SATPI
                        rfc_sat = (satpi_d.get("rfc") or satpi_d.get("RFC") or "").strip().upper()
                        cp_v = (satpi_d.get("cp") or satpi_d.get("CP") or "").strip()
                        curp_v = (satpi_d.get("curp") or satpi_d.get("CURP") or "").strip()
                        nom_v = (satpi_d.get("nombre") or satpi_d.get("NOMBRE") or "").strip()
                        reg_desc_v = (satpi_d.get("regimen_desc") or satpi_d.get("REGIMEN") or satpi_d.get("regimen") or "").strip()
                        reg_clave_v = (satpi_d.get("regimen_clave") or "").strip()
                
                        # SATPI "confirmado" si trae RFC real + algún dato útil
                        satpi_confirmed = bool(rfc_sat) and bool(cp_v or curp_v or nom_v or reg_desc_v or reg_clave_v)
                
                        if not satpi_confirmed:
                            # STRICT: NO inventar RFC si SATPI no confirma
                            raise RuntimeError("SATPI_NOT_FOUND")
                
                        # ✅ caso ideal: usar SATPI confirmado
                        datos.update(satpi_d)
                        datos["RFC"] = rfc_sat
                        datos["RFC_ETIQUETA"] = rfc_sat
                        datos["_RFC_UNCONFIRMED"] = False
                        datos["_RFC_SOURCE"] = "SATPI"
                
                        # map regimen_desc → REGIMEN si falta
                        if not (datos.get("REGIMEN") or "").strip():
                            if (satpi_d.get("regimen_desc") or "").strip():
                                datos["REGIMEN"] = (satpi_d.get("regimen_desc") or "").strip()
                
                        datos = normalize_regimen_fields(datos)
                
                        if (datos.get("REGIMEN") or datos.get("regimen") or "").strip():
                            datos["_REG_SOURCE"] = "SATPI"
                        if (datos.get("CP") or datos.get("cp") or "").strip():
                            datos["_CP_SOURCE"] = "SATPI"
                
                        datos = _apply_strict(datos)
                
                    except RuntimeError as e:
                        se = str(e)
                
                        # Normaliza códigos de _rfc_only_fallback_satpi a tus SATPI_* amigables
                        if se in ("SATPI_428", "SATPI_RFC_LEN"):
                            se = "SATPI_RFC_INVALID"
                        elif se == "SATPI_412":
                            se = "SATPI_NO_QUOTA"
                        elif se.startswith("SATPI_NET:") or se.startswith("SATPI_TEMP:") or se.startswith("SATPI_BAD:5"):
                            se = "SATPI_TEMP"
                        elif se in ("SATPI_NOT_FOUND", "SATPI_NO_DATA") or se.startswith("SATPI_BAD:"):
                            se = "SATPI_NOT_FOUND"
                        elif se not in ("RFC_CANDIDATE_EMPTY", "SATPI_UNEXPECTED"):
                            # si viene algo raro, márcalo como inesperado
                            pass
                
                        # ✅ mensajes claros y salida controlada (evita handler genérico)
                        if se == "SATPI_RFC_INVALID":
                            wa_send_text(from_wa_id, "⚠️ La CURP parece inválida (no pude validar el RFC derivado). Verifica y vuelve a intentarlo.")
                            return
                        if se == "SATPI_NO_QUOTA":
                            wa_send_text(from_wa_id, "⚠️ En este momento el servicio de validación está sin consultas disponibles. Intenta más tarde.")
                            return
                        if se == "SATPI_TEMP":
                            wa_send_text(from_wa_id, "⚠️ El servicio de validación está saturado o tardando en responder. Intenta de nuevo en 2-3 minutos.")
                            return
                        if se == "SATPI_NOT_FOUND":
                            wa_send_text(from_wa_id, "❌ No se encontró un RFC asociado a esta CURP. Verifica la CURP y vuelve a intentarlo.")
                            return
                        if se == "RFC_CANDIDATE_EMPTY":
                            wa_send_text(from_wa_id, "❌ No pude derivar el RFC desde la CURP por falta de datos base. Verifica la CURP y vuelve a intentarlo.")
                            return
                
                        wa_send_text(from_wa_id, "⚠️ Ocurrió un error interno validando el RFC. Intenta de nuevo.")
                        return
                
                rfc_obtenido = (datos.get("RFC") or "").strip().upper()

                # ============================================================
                #  PATCH PRO: SOLO CUANDO CURP NO TRAE RFC
                # ============================================================
                if input_type == "CURP" and (not STRICT_NO_SEPOMEX_ESSENTIALS) and (not rfc_obtenido):
                
                    # ---------- 1) CALCULAR RFC PF (13) ----------
                    try:
                        fn_raw = (datos.get("FECHA_NACIMIENTO") or "").strip()
                
                        fecha_iso = ""
                        m0 = re.match(r"^(\d{2})/(\d{2})/(\d{4})$", fn_raw)
                        if m0:
                            fecha_iso = f"{m0.group(3)}-{m0.group(2)}-{m0.group(1)}"
                        else:
                            # dd-mm-aaaa
                            m = re.match(r"^(\d{2})-(\d{2})-(\d{4})$", fn_raw)
                            if m:
                                fecha_iso = f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
                            else:
                                # yyyy-mm-dd
                                m2 = re.match(r"^(\d{4})-(\d{2})-(\d{2})", fn_raw)
                                fecha_iso = m2.group(0) if m2 else ""

                        if fecha_iso:
                            rfc_calc = rfc_pf_13(
                                (datos.get("NOMBRE") or ""),
                                (datos.get("PRIMER_APELLIDO") or ""),
                                (datos.get("SEGUNDO_APELLIDO") or ""),
                                fecha_iso
                            ).strip().upper()
                
                            if rfc_calc:
                                datos["RFC"] = rfc_calc
                                datos["RFC_ETIQUETA"] = rfc_calc
                                datos["_RFC_UNCONFIRMED"] = True 
                                datos["_RFC_SOURCE"] = "DERIVED"
                
                    except Exception as e:
                        print("RFC CALC FAIL:", repr(e))
                
                    # si aún no hay RFC → ahora sí error
                    rfc_obtenido = (datos.get("RFC") or "").strip().upper()
                    if not rfc_obtenido:
                        datos["_NO_RFC_FOUND"] = True
                        datos["_RFC_SOURCE"] = datos.get("_RFC_SOURCE") or "NONE"
                
                    # ---------- 2) REGIMEN FIJO (SOLO EN ESTE CASO) ----------
                    datos["_REG_SOURCE"] = datos.get("_REG_SOURCE") or "NONE"
                
                    # ---------- 3) CP/MUNICIPIO/LOCALIDAD/COLONIA (SEPOMEX) SOLO SI FALTAN ----------
                    try:
                        seed_key = ((datos.get("CURP") or "") or (datos.get("RFC") or "")).strip().upper()
                        if not STRICT_NO_SEPOMEX_ESSENTIALS:
                            _before = dict(datos)

                            tmp = sepomex_fill_domicilio_desde_entidad(datos, seed_key=seed_key) or {}
                            if isinstance(tmp, dict):
                                merged = dict(_before)
                                merged.update(tmp)
                                datos = merged
                            else:
                                # por si acaso
                                datos = _before

                    except Exception as e:
                        print("SEPOMEX FILL FAIL (STEP3):", repr(e))

                    try:
                        pass
                    except Exception as e:
                        print("CP PICK FAIL:", repr(e))

                rfc_obtenido = (datos.get("RFC") or "").strip().upper()

                if input_type == "CURP" and (not STRICT_NO_SEPOMEX_ESSENTIALS):
                    try:
                        cp_val = re.sub(r"\D+", "", (datos.get("CP") or "")).strip()
                        ent_val = (datos.get("ENTIDAD") or "").strip()
                        mun_val = (datos.get("MUNICIPIO") or datos.get("LOCALIDAD") or "").strip()
                        col_val = (datos.get("COLONIA") or "").strip()
                
                        needs_fill = (len(cp_val) != 5) or (not ent_val) or (not mun_val) or (not col_val)
                
                        if needs_fill:
                            seed_key = (datos.get("CURP") or datos.get("RFC") or query).strip().upper()
                            _before = dict(datos)
                            tmp = sepomex_fill_domicilio_desde_entidad(datos, seed_key=seed_key) or {}
                            if isinstance(tmp, dict):
                                merged = dict(_before)
                                merged.update(tmp)
                                datos = merged
                            else:
                                datos = _before

                    except Exception as e:
                        print("SEPOMEX FILL FAIL (CURP non-strict):", repr(e), flush=True)

                if input_type == "RFC_ONLY" and STRICT_NO_SEPOMEX_ESSENTIALS:
                    # si aún no cumple "oficial", intenta confirmarlo por SATPI
                    if not _strict_gate_or_abort(datos, input_type):
                        try:
                            sat = _rfc_only_fallback_satpi(query) or {}
                            tmp = normalize_satpi_rfc_only(sat, rfc_query=query)
                        
                            # merge sin pisar campos ya existentes
                            for k, v in tmp.items():
                                if v is None:
                                    continue
                                if isinstance(v, str):
                                    if v.strip() and not (str(datos.get(k) or "").strip()):
                                        datos[k] = v
                                else:
                                    if v and not datos.get(k):
                                        datos[k] = v
                        
                            # siempre asegura RFC en mayúsculas
                            if (tmp.get("RFC") or "").strip():
                                datos["RFC"] = tmp["RFC"]
                                datos["RFC_ETIQUETA"] = tmp["RFC_ETIQUETA"]
                        
                            datos = normalize_regimen_fields(datos)
                            datos = _apply_strict(datos)
                        
                        except Exception as e:
                            print("RFC_ONLY strict SATPI confirm fail:", repr(e), flush=True)
                            pass
                try:
                    seed_key = (datos.get("RFC") or datos.get("CURP") or query).strip().upper()
                
                    cp_final = re.sub(r"\D+", "", (datos.get("CP") or datos.get("cp") or "")).strip()
                
                    if len(cp_final) == 5:
                        # si el CP vino de SATPI o CHECKID o PICK, aquí permitimos forzar mun
                        cp_src = (datos.get("_CP_SOURCE") or "").strip().upper()
                
                        force_mun = cp_src in ("CHECKID", "SATPI", "SEPOMEX_PICK", "SEPOMEX")
                
                        # MUY IMPORTANTE:
                        # - si vienes de CURP + GOBMX, tú pones _MUN_LOCK=True
                        # - reconcile_location_by_cp con force_mun=True corrige MUNICIPIO aunque esté locked
                        datos["CP"] = cp_final
                        datos = reconcile_location_by_cp(datos, seed_key=seed_key, force_mun=force_mun)
                
                except Exception as e:
                    print("RECONCILE FINAL FAIL:", repr(e), flush=True)
                
                if STRICT_NO_SEPOMEX_ESSENTIALS:
                    datos = normalize_regimen_fields(datos)
                    if not _strict_gate_or_abort(datos, input_type):
                        wa_send_text(from_wa_id, "⚠️ No pude obtener datos oficiales.")
                        return

                inc_req_if_needed()

                rfc_obtenido = (datos.get("RFC") or "").strip().upper()
                if rfc_obtenido:
                    try:
                        pub_url = validacion_sat_publish(datos, input_type)
                        if pub_url:
                            datos["QR_URL"] = pub_url
                    except Exception as e:
                        print("validacion_sat_publish fail:", e)
                else:
                    print("skip validacion_sat_publish: no RFC", flush=True)

                datos = ensure_idcif_fakey(datos)

                try:
                    mun_final = (datos.get("LOCALIDAD") or datos.get("MUNICIPIO") or "").strip().upper()
                    ent_final = (datos.get("ENTIDAD") or "").strip().upper()
                    datos["FECHA"] = _fecha_lugar_mun_ent(mun_final, ent_final)
                except Exception as e:
                    print("FECHA recompute fail:", repr(e))

                print(
                    "[PRE DOCX]",
                    "RFC=", datos.get("RFC"),
                    "| RFC_ETIQUETA=", datos.get("RFC_ETIQUETA"),
                    "| REGIMEN=", datos.get("REGIMEN"),
                    "| CP=", datos.get("CP"),
                    "| COLONIA=", datos.get("COLONIA"),
                    "| MUNICIPIO=", datos.get("LOCALIDAD"),
                    "| ENTIDAD=", datos.get("ENTIDAD"),
                )

                datos = normalize_regimen_fields(datos)

                if (not STRICT_NO_SEPOMEX_ESSENTIALS) and input_type == "CURP":
                    if not ((datos.get("REGIMEN") or "").strip() or (datos.get("regimen") or "").strip()):
                        datos["REGIMEN"] = "Régimen de Sueldos y Salarios e Ingresos Asimilados a Salarios"
                        datos["regimen"] = datos["REGIMEN"]
                        datos["_REG_SOURCE"] = datos.get("_REG_SOURCE") or "DEFAULT_SUELDOS"

                if STRICT_NO_SEPOMEX_ESSENTIALS and _regimen_no_vigente(datos):
                    wa_send_text(
                        from_wa_id,
                        "⚠️ No pude obtener un régimen vigente para ese RFC.\n"
                        "Esto suele ocurrir cuando el RFC está suspendido/cancelado o sin situación fiscal activa."
                    )
                    return

                wa_step(from_wa_id, "📄 Generando PDF/Word...", step="DOCS", force=True)
                _generar_y_enviar_archivos(from_wa_id, text_body, datos, input_type, test_mode)
                return
            finally:
                inflight_end(ok_key)

        # =========================
        # RFC + IDCIF (SAT)
        # =========================
        rfc, idcif = extraer_rfc_idcif(text_body)

        if not rfc or not idcif:
            if looks_like_user_typed_an_rfc(text_body) or looks_like_user_typed_an_idcif(text_body):
                wa_send_text(
                    from_wa_id,
                    "ℹ️ Aviso:\n\n"
                    "Se recibió únicamente el RFC, sin el IDCIF.\n\n"
                    "Para RFC de 13 caracteres, es posible continuar y generar el archivo en formato PDF.\n\n"
                    "Si deseas enviar el IDCIF, utiliza el siguiente formato:\n"
                    "RFC (12 o 13 caracteres) + IDCIF (11 dígitos)\n\n"
                )

                return

            wa_send_text(
                from_wa_id,
                "⚠️ El mensaje recibido no corresponde a un formato válido.\n\n"
                "Por favor, envía la información en uno de los siguientes formatos:\n\n"
                "• RFC (12 o 13 caracteres) + IDCIF (11 dígitos)\n"
                "• CURP (18 caracteres)\n"
                "• RFC (13 caracteres)\n\n"
                "ℹ️ Si deseas recibir el archivo también en Word, agrega DOCX al final del mensaje."
            )
            return

        if len(rfc) not in (12, 13):
            wa_send_text(from_wa_id, "El RFC debe tener 12 (moral) o 13 (física) caracteres. Verifica y envíalo de nuevo.")
            return
        if len(idcif) != 11:
            wa_send_text(
                from_wa_id,
                "⚠️ Validación de datos:\n\n"
                "El IDCIF recibido es inválido, ya que no contiene 11 dígitos.\n"
                "Para continuar, verifica el identificador y envíalo nuevamente."
            )
            return

        if not is_valid_rfc(rfc) or not is_valid_idcif(idcif):
            wa_send_text(from_wa_id, ERR_RFC_IDCIF_INVALID)
            return

        ok_key = make_ok_key("RFC_IDCIF", rfc=rfc, curp=None)
        if not inflight_start(ok_key):
            wa_send_text(from_wa_id, MSG_IN_PROCESS)
            return

        inc_req_if_needed()

        try:
            try:
                datos = extraer_datos_desde_sat(rfc, idcif)
            except ValueError as e:
                if str(e) == "SIN_DATOS_SAT":
                    wa_send_text(from_wa_id, ERR_SAT_NO_DATA)
                    return
                raise
            except (requests.exceptions.Timeout, requests.exceptions.ConnectionError, requests.exceptions.RequestException):
                wa_send_text(from_wa_id, ERR_SERVICE_DOWN)
                return

            wa_step(from_wa_id, "📄 Generando PDF/Word...", step="DOCS", force=True)
            _generar_y_enviar_archivos(from_wa_id, text_body, datos, "RFC_IDCIF", test_mode)
            return

        finally:
            inflight_end(ok_key)

    except Exception as e:
        err = e
        print("Worker error:", e)
        traceback.print_exc()
        try:
            wa_step(
                from_wa_id,
                f"⚠️ Ocurrió un error procesando tu solicitud.\n🧾 Folio: {rid}\n\n"
                "Intenta de nuevo.\n"
                "Si vuelve a pasar, envíame ese folio para revisarlo.",
                step="FATAL",
                force=True
            )
        except Exception:
            pass

    finally:
        try:
            if msg_id:
                if err is None:
                    wa_mark_done(msg_id, from_wa_id)
                else:
                    wa_unmark(msg_id)
        except Exception as e2:
            print("wa_mark_done/wa_unmark failed:", e2)

        try:
            if job.get("bp_slot"):
                wa_release_slot()
        except Exception:
            pass

def construir_datos_desde_checkid_curp_sin_rfc(curp: str) -> dict:
    term_norm = (curp or "").strip().upper()
    if not term_norm:
        raise ValueError("TERM_EMPTY")

    ci_raw = checkid_lookup(term_norm)
    ci = _norm_checkid_fields(ci_raw)

    if not (ci.get("CURP") or "").strip():
        raise RuntimeError("CHECKID_SIN_DATOS")

    cp_final = re.sub(r"\D+", "", (ci.get("CP") or "")).strip()

    entidad_ci = (ci.get("ENTIDAD") or "").strip().upper()
    municipio_ci = ((ci.get("MUNICIPIO") or ci.get("LOCALIDAD") or "")).strip().upper()
    colonia_ci = (ci.get("COLONIA") or "").strip().upper()

    datos = {
        "RFC": (ci.get("RFC") or "").strip().upper(),
        "RFC_ETIQUETA": (ci.get("RFC") or "").strip().upper(),

        "CURP": (ci.get("CURP") or "").strip().upper(),

        "NOMBRE": (ci.get("NOMBRE") or "").strip().upper(),
        "PRIMER_APELLIDO": (ci.get("APELLIDO_PATERNO") or ci.get("PRIMER_APELLIDO") or "").strip().upper(),
        "SEGUNDO_APELLIDO": (ci.get("APELLIDO_MATERNO") or ci.get("SEGUNDO_APELLIDO") or "").strip().upper(),

        "FECHA_NACIMIENTO": (ci.get("FECHA_NACIMIENTO") or "").strip(),

        "ENTIDAD": entidad_ci,

        # ✅ GUARDA EN LOS DOS CAMPOS para evitar mezclas downstream
        "MUNICIPIO": municipio_ci,
        "LOCALIDAD": municipio_ci,
        "_MUN_SOURCE": "CHECKID",

        "CP": cp_final,
        "COLONIA": colonia_ci,

        "REGIMEN": (ci.get("REGIMEN") or "").strip(),
        "regimen": (ci.get("REGIMEN") or "").strip(),

        "TIPO_VIALIDAD": (ci.get("TIPO_VIALIDAD") or "CALLE").strip().upper(),
        "VIALIDAD": (ci.get("VIALIDAD") or "SIN NOMBRE").strip().upper(),

        # OJO: mantén los NO_* porque luego los mapeamos a NUMERO_* en ensure_default...
        "NO_EXTERIOR": re.sub(r"\D+", "", (ci.get("NO_EXTERIOR") or "")),
        "NO_INTERIOR": re.sub(r"\D+", "", (ci.get("NO_INTERIOR") or "")),

        "NUMERO_EXTERIOR": re.sub(r"\D+", "", (ci.get("NO_EXTERIOR") or "")),
        "NUMERO_INTERIOR": re.sub(r"\D+", "", (ci.get("NO_INTERIOR") or "")),

        "SITUACION_CONTRIBUYENTE": (ci.get("ESTATUS") or ci.get("SITUACION_CONTRIBUYENTE") or "ACTIVO").strip().upper(),
    }

    # ✅ Si hay CP válido, amarra ENT/MUN/COL al CP (SEPOMEX manda municipio por CP)
    cp = datos["CP"]
    if len(cp) == 5:
        datos["_CP_SOURCE"] = "CHECKID"
    
        try:
            seed_key2 = (datos.get("RFC") or datos.get("CURP") or term_norm).strip().upper()
    
            tmp2 = {
                "CP": cp,
                "ENTIDAD": (datos.get("ENTIDAD") or "").strip().upper(),
                "MUNICIPIO": ((datos.get("MUNICIPIO") or datos.get("LOCALIDAD") or "")).strip().upper(),
                "LOCALIDAD": ((datos.get("LOCALIDAD") or datos.get("MUNICIPIO") or "")).strip().upper(),
                "COLONIA": (datos.get("COLONIA") or "").strip().upper(),
                "_MUN_LOCK": False,
            }
    
            # ✅ CP confiable (CHECKID) -> fuerza municipio por CP
            tmp2 = reconcile_location_by_cp(tmp2, seed_key=seed_key2, force_mun=True)
    
            ent2 = (tmp2.get("ENTIDAD") or "").strip().upper()
            if ent2:
                datos["ENTIDAD"] = ent2
    
            mun2 = (tmp2.get("MUNICIPIO") or tmp2.get("LOCALIDAD") or "").strip().upper()
            if mun2:
                datos["MUNICIPIO"] = mun2
                datos["LOCALIDAD"] = mun2
                datos["_MUN_SOURCE"] = "SEPOMEX"
    
            col2 = (tmp2.get("COLONIA") or "").strip().upper()
            if col2:
                datos["COLONIA"] = col2
    
        except Exception as e_re:
            print("reconcile checkid_curp_sin_rfc by cp fail:", repr(e_re), flush=True)
    
            # fallback suave: si colonia vacía, al menos pick colonia
            if not (datos.get("COLONIA") or "").strip():
                col = sepomex_pick_colonia_by_cp(cp, seed_key=datos.get("CURP") or term_norm)
                if col:
                    datos["COLONIA"] = col.strip().upper()

    return datos

def _pick(*vals) -> str:
    for v in vals:
        if v not in (None, "", [], {}):
            return str(v).strip()
    return ""

def _upper(s: str) -> str:
    return (s or "").strip().upper()

def _digits(s: str) -> str:
    return re.sub(r"\D+", "", (s or ""))

def tipo_persona_por_rfc(rfc: str) -> str:
    rfc = (rfc or "").strip().upper()
    if len(rfc) == 12:
        return "MORAL"
    if len(rfc) == 13:
        return "FISICA"
    return "DESCONOCIDO"

def _pick_first(datos: dict, *keys: str) -> str:
    datos = datos or {}
    for k in keys:
        v = datos.get(k)
        if v not in (None, "", [], {}):
            return str(v).strip()
    return ""

def completar_campos_por_tipo(datos: dict) -> dict:
    """
    Asegura que el dict 'datos' traiga EXACTO lo que piden tus plantillas:
    - Física (plantilla.docx): CURP, NOMBRE, PRIMER APELLIDO, SEGUNDO APELLIDO
    - Moral (plantilla-moral.docx): DENOMINACION, CAPITAL, NOMBRE (comercial)
    Además, siempre llena NOMBRE_ETIQUETA (header) según corresponda.
    """
    datos = datos or {}

    rfc = (datos.get("RFC") or datos.get("rfc") or datos.get("RFC_ETIQUETA") or "").strip().upper()
    tipo = tipo_persona_por_rfc(rfc)

    if tipo == "MORAL":
        # Denominación / Razón Social (varios posibles nombres desde SAT/APIs)
        den = _pick_first(datos,
            "DENOMINACION", "DENOMINACIÓN",
            "RAZON_SOCIAL", "RAZÓN_SOCIAL",
            "RAZON SOCIAL", "RAZÓN SOCIAL",
            "NOMBRE_RAZON_SOCIAL", "NOMBRE_RAZÓN_SOCIAL",
            "denominacion", "razon_social", "razonSocial"
        )
        # Régimen de capital
        cap = _pick_first(datos,
            "CAPITAL", "REGIMEN_CAPITAL", "RÉGIMEN_CAPITAL",
            "REGIMEN DE CAPITAL", "RÉGIMEN DE CAPITAL",
            "regimen_capital", "regimenDeCapital"
        )
        # Nombre comercial (si no existe, usa denominación)
        nom_comercial = _pick_first(datos, "NOMBRE_COMERCIAL", "NOMBRE COMERCIAL", "nombre_comercial")
        if not nom_comercial:
            nom_comercial = den

        # Set fields para plantilla-moral.docx :contentReference[oaicite:2]{index=2}
        datos["DENOMINACION"] = den
        datos["CAPITAL"] = cap
        datos["NOMBRE"] = nom_comercial

        # Limpieza de campos de física para que no “ensucien”
        datos["CURP"] = ""
        datos["PRIMER_APELLIDO"] = ""
        datos["SEGUNDO_APELLIDO"] = ""

        # Header (NOMBRE ETIQUETA)
        datos["NOMBRE_ETIQUETA"] = (datos.get("NOMBRE_ETIQUETA") or den or "").strip().upper()

    elif tipo == "FISICA":
        # Fields para plantilla.docx :contentReference[oaicite:3]{index=3}
        curp = _pick_first(datos, "CURP", "curp")
        nombre = _pick_first(datos, "NOMBRE", "nombre")
        ap1 = _pick_first(datos, "PRIMER_APELLIDO", "apellido_paterno", "APELLIDO_PATERNO")
        ap2 = _pick_first(datos, "SEGUNDO_APELLIDO", "apellido_materno", "APELLIDO_MATERNO")

        datos["CURP"] = curp
        datos["NOMBRE"] = nombre
        datos["PRIMER_APELLIDO"] = ap1
        datos["SEGUNDO_APELLIDO"] = ap2

        # Limpieza moral
        datos["DENOMINACION"] = datos.get("DENOMINACION", "") or ""
        datos["CAPITAL"] = datos.get("CAPITAL", "") or ""

        # Header
        if not (datos.get("NOMBRE_ETIQUETA") or "").strip():
            datos["NOMBRE_ETIQUETA"] = " ".join([x for x in [nombre, ap1, ap2] if x]).strip().upper()

    return datos

def norm_persona_from_datos(datos: dict, rfc: str, idcif: str, d3_key: str) -> dict:
    datos = datos or {}

    # ---- base fields (elige de ambas variantes) ----
    curp   = _upper(_pick(datos.get("CURP"), datos.get("curp")))
    nombre = _upper(_pick(datos.get("NOMBRE"), datos.get("nombre")))
    ap1    = _upper(_pick(datos.get("PRIMER_APELLIDO"), datos.get("apellido_paterno"), datos.get("APELLIDO_PATERNO")))
    ap2    = _upper(_pick(datos.get("SEGUNDO_APELLIDO"), datos.get("apellido_materno"), datos.get("APELLIDO_MATERNO")))

    # ---- fechas dd-mm-aaaa SIEMPRE ----
    fn = _to_dd_mm_aaaa_dash(_pick(datos.get("FECHA_NACIMIENTO"), datos.get("fecha_nacimiento")))
    fi = _to_dd_mm_aaaa_dash(_pick(datos.get("FECHA_INICIO"), datos.get("fecha_inicio_operaciones")))
    fu = _to_dd_mm_aaaa_dash(_pick(datos.get("FECHA_ULTIMO"), datos.get("fecha_ultimo_cambio")))
    fa = _to_dd_mm_aaaa_dash(_pick(datos.get("FECHA_ALTA"), datos.get("fecha_alta")))

    # ---- direccion uppercase ----
    entidad   = _upper(_pick(datos.get("ENTIDAD"), datos.get("entidad")))
    municipio = _upper(_pick(datos.get("LOCALIDAD"), datos.get("municipio")))
    colonia   = _upper(_pick(datos.get("COLONIA"), datos.get("colonia")))

    tipo_v = _upper(_pick(datos.get("TIPO_VIALIDAD"), datos.get("tipo_vialidad"), "CALLE")) or "CALLE"
    nom_v  = _upper(_pick(datos.get("VIALIDAD"), datos.get("nombre_vialidad"), "SIN NOMBRE")) or "SIN NOMBRE"

    no_ext = _digits(_pick(datos.get("NO_EXTERIOR"), datos.get("numero_exterior")))
    no_int = _digits(_pick(datos.get("NO_INTERIOR"), datos.get("numero_interior")))
    cp     = _digits(_pick(datos.get("CP"), datos.get("cp")))

    # ---- regimen limpio ----
    reg = _pick(datos.get("REGIMEN"), datos.get("regimen"))
    reg = limpiar_regimen(reg)  # usa tu limpiar_regimen actual

    # ---- estatus ----
    estatus = _upper(_pick(datos.get("ESTATUS"), datos.get("situacion_contribuyente"), "ACTIVO")) or "ACTIVO"

    # ---- AL = "ENTIDAD 1" ----
    al = _al_from_entidad(entidad)

    # ---- etiquetas ----
    nombre_etiqueta = " ".join(x for x in [nombre, ap1, ap2] if x).strip()

    return {
        "D1": "10",
        "D2": "1",
        "D3": d3_key,

        "rfc": _upper(rfc),
        "curp": curp,
        "nombre": nombre,
        "apellido_paterno": ap1,
        "apellido_materno": ap2,

        "fecha_nacimiento": fn,
        "fecha_inicio_operaciones": fi,
        "situacion_contribuyente": estatus,
        "fecha_ultimo_cambio": fu,
        "regimen": reg,
        "fecha_alta": fa,

        "entidad": entidad,
        "municipio": municipio,
        "colonia": colonia,
        "tipo_vialidad": tipo_v,
        "nombre_vialidad": nom_v,
        "numero_exterior": no_ext,
        "numero_interior": no_int,
        "cp": cp,

        "correo": _pick(datos.get("CORREO"), datos.get("correo")),
        "al": al,

        "RFC_ETIQUETA": _upper(rfc),
        "NOMBRE_ETIQUETA": nombre_etiqueta,
        "IDCIF_ETIQUETA": str(idcif).strip(),
    }

def public_label(input_type: str) -> str:
    if input_type in ("RFC_IDCIF", "RFC_ONLY", "QR", "CURP", "MANUAL"):
        return "RFC"
    return input_type

def _generar_y_enviar_archivos(from_wa_id: str, text_body: str, datos: dict, input_type: str, test_mode: bool):
    base_dir = os.path.dirname(os.path.abspath(__file__))

    # ✅ Guardar en GitHub personas.json SOLO si viene de APIs
    if input_type in ("CURP", "RFC_ONLY"):
        idcif = datos.get("IDCIF") or datos.get("IDCIF_ETIQUETA")

        if not idcif:
            raise RuntimeError("❌ Falta IDCIF fakey en datos (IDCIF / IDCIF_ETIQUETA)")

        rfc = (datos.get("RFC") or datos.get("rfc") or "").strip().upper()
        if not rfc:
            raise RuntimeError("❌ Falta RFC en datos (RFC / rfc)")

        d3_key = f"{idcif}_{rfc}"

        persona = norm_persona_from_datos(datos=datos, rfc=rfc, idcif=idcif, d3_key=d3_key)

        try:
            github_upsert_persona_file(d3_key, persona)   # ✅
        except Exception as e:
            print("⚠ Error actualizando persona file:", repr(e), "key=", d3_key, flush=True)

    # ✅ Completa campos según el tipo (moral/física) y luego decide plantilla
    datos = completar_campos_por_tipo(datos)
    
    rfc_real = (datos.get("RFC") or datos.get("rfc") or "").strip().upper()
    tipo = tipo_persona_por_rfc(rfc_real)
    
    reg = (datos.get("REGIMEN") or "").upper()
    
    if tipo == "MORAL":
        nombre_plantilla = "plantilla-moral.docx"   # usa DENOMINACION/CAPITAL :contentReference[oaicite:11]{index=11}
    elif ("SUELDOS" in reg) and ("SALARIOS" in reg):
        nombre_plantilla = "plantilla-asalariado.docx"
    else:
        nombre_plantilla = "plantilla.docx"         # usa CURP/NOMBRE/APELLIDOS :contentReference[oaicite:12]{index=12}

    ruta_plantilla = os.path.join(base_dir, nombre_plantilla)

    t_upper = (text_body or "").upper()
    quiere_docx = ("DOCX" in t_upper) or ("WORD" in t_upper) or ("AMBOS" in t_upper)

    with tempfile.TemporaryDirectory() as tmpdir:
        nombre_base = (datos.get("CURP") or datos.get("RFC") or "CONSTANCIA").strip() or "CONSTANCIA"
        label = public_label(input_type)
        nombre_docx = f"{nombre_base}_{label}.docx"
        ruta_docx = os.path.join(tmpdir, nombre_docx)

        # ==========================
        # ✅ QR2 (D26): folio + JSON + PNG para image9.png
        # ==========================
        # asegura FECHA_CORTA para la cadena original
        seed_key = (datos.get("RFC") or datos.get("rfc") or datos.get("CURP") or datos.get("curp") or "").strip().upper() or "SEED"
        datos = ensure_default_status_and_dates(datos, seed_key=seed_key)

        rfc_base = (datos.get("RFC") or datos.get("rfc") or "").strip().upper()
        if not rfc_base:
            raise RuntimeError("❌ Falta RFC para generar QR2 (D26)")

        folio26 = _d26_folio_deterministico(rfc_base)
        d3_26 = f"{folio26}_{rfc_base}"

        base = "https://siat.sat.validacion-sat.com"
        qr2_url = f"{base}/app/qr/faces/pages/mobile/validadorqr.jsf?D1=26&D2=1&D3={d3_26}"

        persona26 = _persona_d26_min(datos, d3_key=d3_26, rfc=rfc_base)

        try:
            github_upsert_persona_file(d3_26, persona26)
        except Exception as e:
            print("⚠ Error publicando persona D26:", repr(e), "d3_26=", d3_26, flush=True)

        qr2_bytes = generar_solo_qr_png(qr2_url)

        reemplazar_en_documento(ruta_plantilla, ruta_docx, datos, input_type, qr2_bytes=qr2_bytes)

        with open(ruta_docx, "rb") as f:
            docx_bytes = f.read()

        # ==========================
        # ✅ PASO 7: CACHE PDF por ok_key
        # ==========================
        rfc_c = (datos.get("RFC") or datos.get("rfc") or "").strip().upper()
        curp_c = (datos.get("CURP") or datos.get("curp") or "").strip().upper()
        ok_key = make_ok_key(input_type, rfc=rfc_c or None, curp=curp_c or None)

        cached_name, cached_bytes, cached_mime = filecache_get_bytes(ok_key, "PDF")
        if cached_bytes and cached_mime == "application/pdf":
            pdf_filename = cached_name or (os.path.splitext(nombre_docx)[0] + ".pdf")

            media_pdf = wa_upload_document(cached_bytes, pdf_filename, "application/pdf")
            wa_send_document(from_wa_id, media_pdf, pdf_filename)

            _bill_and_log_ok(from_wa_id, input_type, datos, test_mode)

            if quiere_docx:
                media_docx = wa_upload_document(
                    docx_bytes,
                    nombre_docx,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                wa_send_document(from_wa_id, media_docx, nombre_docx, caption="")

            return
            
        # PDF default
        try:
            pdf_path = os.path.join(tmpdir, os.path.splitext(nombre_docx)[0] + ".pdf")
            docx_to_pdf_aspose(docx_path=ruta_docx, pdf_path=pdf_path)

            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()

            pdf_filename = os.path.splitext(nombre_docx)[0] + ".pdf"

            # ✅ guarda PDF en cache (si es tamaño permitido)
            filecache_set_bytes(ok_key, "PDF", pdf_filename, pdf_bytes, "application/pdf")

            media_pdf = wa_upload_document(pdf_bytes, pdf_filename, "application/pdf")
            wa_send_document(from_wa_id, media_pdf, pdf_filename)

            _bill_and_log_ok(from_wa_id, input_type, datos, test_mode)

            if quiere_docx:
                media_docx = wa_upload_document(
                    docx_bytes,
                    nombre_docx,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                wa_send_document(from_wa_id, media_docx, nombre_docx, caption="")

        except Exception as e:
            print("PDF fail, sending DOCX fallback:", e)

            _log_aspose_fail(from_wa_id, input_type, datos, e, where="WA__generar_y_enviar_archivos")

            media_docx = wa_upload_document(
                docx_bytes,
                nombre_docx,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            wa_send_document(from_wa_id, media_docx, nombre_docx, caption="⚠️ No pude convertir a PDF, pero aquí está en DOCX.")

            _bill_and_log_ok(from_wa_id, input_type, datos, test_mode)

def generar_pdf_en_tmp(tmpdir: str, text_body: str, datos: dict, input_type: str) -> str:
    """
    Genera el PDF en tmpdir y regresa el nombre del archivo PDF.
    NO manda nada a WhatsApp (esto es para ZIP).
    """
    base_dir = os.path.dirname(os.path.abspath(__file__))

    # Plantilla según tipo/regimen (igual que tu lógica)
    datos = completar_campos_por_tipo(datos)

    rfc_real = (datos.get("RFC") or datos.get("rfc") or "").strip().upper()
    tipo = tipo_persona_por_rfc(rfc_real)
    reg = (datos.get("REGIMEN") or "").upper()

    if tipo == "MORAL":
        nombre_plantilla = "plantilla-moral.docx"
    elif ("SUELDOS" in reg) and ("SALARIOS" in reg):
        nombre_plantilla = "plantilla-asalariado.docx"
    else:
        nombre_plantilla = "plantilla.docx"

    ruta_plantilla = os.path.join(base_dir, nombre_plantilla)

    nombre_base = (datos.get("CURP") or datos.get("RFC") or "CONSTANCIA").strip() or "CONSTANCIA"
    label = public_label(input_type)
    nombre_docx = f"{nombre_base}_{label}.docx"
    ruta_docx = os.path.join(tmpdir, nombre_docx)

    reemplazar_en_documento(ruta_plantilla, ruta_docx, datos, input_type)

    pdf_filename = os.path.splitext(nombre_docx)[0] + ".pdf"
    pdf_path = os.path.join(tmpdir, pdf_filename)

    # Convierte a PDF (si falla, lanza excepción para que se cuente como fail)
    docx_to_pdf_aspose(docx_path=ruta_docx, pdf_path=pdf_path)

    return pdf_filename

def _bill_and_log_ok(from_wa_id: str, input_type: str, datos: dict, test_mode: bool):
    # Usa tus helpers: resolve_price, make_ok_key, bill_success_if_new, log_attempt, inc_success
    out = {"reason": None}

    def _tx(s):
        from stats_store import bill_success_if_new, log_attempt, resolve_price, inc_success

        price_mxn = resolve_price(s, from_wa_id, input_type)
        ok_key = make_ok_key(input_type, datos.get("RFC"), datos.get("CURP"))

        res = bill_success_if_new(
            s,
            user=from_wa_id,
            ok_key=ok_key,
            input_type=input_type,
            price_mxn=price_mxn,
            is_test=test_mode
        )

        out["reason"] = res.get("reason")
        if res.get("billed"):
            inc_success(s, from_wa_id, (datos.get("RFC") or ""))
            log_attempt(s, from_wa_id, ok_key, True, "BILLED_OK",
                        {"via": "WA", "type": input_type, "price": price_mxn}, is_test=test_mode)
        else:
            code = "OK_NO_BILL"
            if res.get("reason") == "DUPLICATE":
                code = "OK_DUPLICATE_NO_BILL"
            elif res.get("reason") == "TEST":
                code = "OK_TEST_NO_BILL"
            log_attempt(s, from_wa_id, ok_key, True, code,
                        {"via": "WA", "type": input_type, "reason": res.get("reason")}, is_test=test_mode)

    get_and_update(STATS_PATH, _tx)

def _log_aspose_fail(from_wa_id: str, input_type: str, datos: dict, err: Exception, where: str = "DOCX2PDF"):
    """
    Loggea fallos de Aspose sin afectar success_total ni billing.
    """
    try:
        rfc = (datos.get("RFC") or "").strip().upper()
        curp = (datos.get("CURP") or "").strip().upper()
        ok_key = make_ok_key(input_type, rfc, curp)

        def _tx(s):
            from stats_store import log_attempt
            log_attempt(
                s,
                from_wa_id,
                ok_key,
                False,
                "ASPOSE_FAIL",
                {
                    "where": where,
                    "type": input_type,
                    "rfc": rfc,
                    "curp": curp,
                    "error": repr(err),
                },
                is_test=False
            )

        get_and_update(STATS_PATH, _tx)
    except Exception as e2:
        print("ASPOSE_FAIL log error:", repr(e2))

def cleanup_expired_sessions():
    st = get_sessions_state()
    us = st.get("user_session") or {}
    changed = False
    now_ts = _now_ts()

    for u, sess in list(us.items()):
        exp_ts = int((sess or {}).get("exp") or 0)
        last_ts = int((sess or {}).get("last") or 0)
        if (exp_ts and exp_ts <= now_ts) or (not last_ts) or ((now_ts - last_ts) > SESSION_IDLE_SECONDS):
            us.pop(u, None)
            changed = True

    if changed:
        st["user_session"] = us
        _atomic_write_json(SESSIONS_PATH, st)

@app.route("/", methods=["GET"])
def home():
    return "Backend OK. Usa POST /login y /generar desde el formulario."

@app.route("/ping", methods=["POST"])
def ping():
    user, reason = usuario_actual_o_none()
    if not user:
        return jsonify({"ok": False, "reason": reason}), 401
    return jsonify({"ok": True})

@app.route("/login", methods=["POST"])
def login():
    data = request.get_json() or {}
    username = (data.get("username") or "").strip()
    password = data.get("password") or ""

    if not username or not password:
        return jsonify({"ok": False, "message": "Faltan usuario o contraseña."}), 400

    password_hash = USERS.get(username)
    if not password_hash or not check_password_hash(password_hash, password):
        return jsonify({"ok": False, "message": "Usuario o contraseña incorrectos."}), 401

    cleanup_expired_sessions()
    device_id = (data.get("device_id") or "").strip() or "UNKNOWN"

    # ========= 1) IP + User-Agent =========
    ip = get_client_ip()
    ua = request.headers.get("User-Agent", "desconocido")
    evento = {
        "ip": ip,
        "fecha": datetime.now(ZoneInfo("America/Mexico_City")).isoformat(),
        "ua": ua,
    }
    HISTORIAL_LOGIN.setdefault(username, []).append(evento)

    # ========= 2) Control por IP (opcional) =========
    info_ip = USERS_IP_INFO.get(username)
    if info_ip:
        if info_ip.get("bloquear_otras") and info_ip.get("ip") and info_ip["ip"] != ip:
            return jsonify({
                "ok": False,
                "message": (
                    "Este usuario ya se encuentra registrado con otra dirección IP "
                    f"({info_ip['ip']}). No se permite iniciar sesión desde una IP distinta."
                ),
            }), 403
    else:
        USERS_IP_INFO[username] = {"ip": ip, "bloquear_otras": BLOQUEAR_IP_POR_DEFAULT}

    # ========= 3) Solo 1 sesión por usuario (persistente) =========
    sess = get_user_session(username)
    active = session_is_active(sess)
    
    if sess and active and not ALLOW_KICKOUT:
        prev_dev = (sess.get("device_id") or "UNKNOWN")
        if prev_dev != device_id:
            return jsonify({
                "ok": False,
                "message": "Este usuario ya tiene una sesión activa en otro dispositivo."
            }), 409

    # si existe pero ya NO está activa -> limpiar
    if sess and not active:
        set_user_session(username, None, None)

    # ========= 4) Crear JWT =========
    try:
        token = crear_jwt(username, device_id=device_id)
    except Exception as e:
        print("JWT error:", e)
        return jsonify({"ok": False, "message": "Error creando sesión."}), 500

    resp = jsonify({"ok": True, "token": token, "message": "Login correcto."})
    resp.headers["Access-Control-Expose-Headers"] = "Authorization"
    return resp

@app.route("/logout", methods=["POST"])
def logout():
    user, reason = usuario_actual_o_none()
    if not user:
        return jsonify({"ok": True})
    set_user_session(user, None, None)
    return jsonify({"ok": True})

@app.route("/generar", methods=["POST"])
def generar_constancia():
    global REQUEST_TOTAL, REQUEST_POR_DIA, SUCCESS_COUNT, SUCCESS_RFCS

    # ------- AUTENTICACIÓN -------
    user, reason = usuario_actual_o_none()
    if not user:
        return jsonify({"ok": False, "reason": reason, "message": "No autorizado"}), 401

    # ====== TEST MODE (WEB) ======
    # En web normalmente no hay texto, así que solo depende del user
    test_mode = is_test_request(user, "")

    def _set_price(s):
        # set_price ya está en stats_store.py
        s.setdefault("billing", {})
        s["billing"]["price_mxn"] = int(PRICE_PER_OK_MXN or 0)
        
    get_and_update(STATS_PATH, _set_price)
    
    # ------- CONTROL LÍMITE DIARIO POR USUARIO -------
    hoy_str = hoy_mexico().isoformat()
    info = USO_POR_USUARIO.get(user)
    if not info or info.get("hoy") != hoy_str:
        info = {"hoy": hoy_str, "count": 0}
        USO_POR_USUARIO[user] = info
    # ----------------------------------------------

    REQUEST_TOTAL += 1
    REQUEST_POR_DIA[hoy_str] = REQUEST_POR_DIA.get(hoy_str, 0) + 1

    print(
        f"[{datetime.utcnow().isoformat()}] Solicitud #{REQUEST_TOTAL} a /generar "
        f"(hoy: {REQUEST_POR_DIA[hoy_str]}) por usuario: {user}"
    )

    rfc = (request.form.get("rfc") or "").strip().upper()
    idcif = (request.form.get("idcif") or "").strip()
    curp = (request.form.get("curp") or "").strip().upper()
    lugar_emision = (request.form.get("lugar_emision") or "").strip()

    # ====== MANUAL JSON (WEB) ======
    manual_json = (request.form.get("manual_json") or "").strip()
    payload = None
    
    if manual_json:
        try:
            payload = json.loads(manual_json)
            if not isinstance(payload, dict):
                raise ValueError("manual_json debe ser objeto")
        except Exception:
            return jsonify({"ok": False, "message": "manual_json inválido (JSON)."}), 400
            
    input_type = None
    term = None
    
    if payload is not None:
        input_type = "MANUAL"
        term = None
    elif curp:
        input_type = "CURP"
        term = curp
    elif rfc and not idcif:
        input_type = "RFC_ONLY"
        term = rfc
    elif rfc and idcif:
        input_type = "RFC_IDCIF"
    else:
        return jsonify({"ok": False, "message": "Falta RFC/IDCIF o CURP."}), 400

    # ✅ CASO 2: CURP inválida (NO CheckID / NO cobro)
    if input_type == "CURP" and not is_valid_curp(term):
        return jsonify({"ok": False, "message": ERR_CURP_INVALID}), 400
    
    # ✅ CASO 3: RFC/IDCIF inválidos (NO SAT / NO cobro)
    if input_type == "RFC_ONLY" and not is_valid_rfc(term):
        return jsonify({"ok": False, "message": ERR_RFC_IDCIF_INVALID}), 400
    
    if input_type == "RFC_IDCIF":
        if not is_valid_rfc(rfc) or not is_valid_idcif(idcif):
            return jsonify({"ok": False, "message": ERR_RFC_IDCIF_INVALID}), 400

    if input_type == "MANUAL":
        rfc_m = (payload.get("RFC") or payload.get("rfc") or "").strip().upper()
        curp_m = (payload.get("CURP") or payload.get("curp") or "").strip().upper()
    
        if rfc_m and not is_valid_rfc(rfc_m):
            return jsonify({"ok": False, "message": ERR_RFC_IDCIF_INVALID}), 400
        if curp_m and not is_valid_curp(curp_m):
            return jsonify({"ok": False, "message": ERR_CURP_INVALID}), 400

    # ====== STATS: request (SOLO si NO es prueba) ======
    if not test_mode:
        def _inc_req(s):
            from stats_store import inc_request, inc_user_request
            inc_request(s)
            inc_user_request(s, user)
        get_and_update(STATS_PATH, _inc_req)
    
    try:
        if input_type == "MANUAL":
            datos = construir_datos_manual(payload, input_type="MANUAL")
    
            # publicar QR (si falla, no abortes)
            try:
                pub_url = validacion_sat_publish(datos, input_type)
                if pub_url:
                    datos["QR_URL"] = pub_url
            except Exception as e:
                print("validacion_sat_publish fail:", e)
    
        elif input_type in ("CURP", "RFC_ONLY"):
            try:
                datos = construir_datos_desde_apis(term)
                datos = normalize_regimen_fields(datos)
            except (requests.exceptions.Timeout, requests.exceptions.ConnectionError, requests.exceptions.RequestException):
                return jsonify({"ok": False, "message": ERR_SERVICE_DOWN}), 503
    
            # ✅ CASO 1: CURP válida pero sin RFC (NO CSF / NO cobro)
            if input_type == "CURP" and not (datos.get("RFC") or "").strip():
                return jsonify({"ok": False, "message": ERR_NO_RFC_FOR_CURP}), 404
    
            # publicar QR (si falla, no abortes)
            try:
                pub_url = validacion_sat_publish(datos, input_type)
                if pub_url:
                    datos["QR_URL"] = pub_url
            except Exception as e:
                print("validacion_sat_publish fail:", e)
    
        else:
            try:
                datos = extraer_datos_desde_sat(rfc, idcif)
            except ValueError as e:
                if str(e) == "SIN_DATOS_SAT":
                    return jsonify({"ok": False, "message": ERR_SAT_NO_DATA}), 404
                raise
            except (requests.exceptions.Timeout, requests.exceptions.ConnectionError, requests.exceptions.RequestException):
                return jsonify({"ok": False, "message": ERR_SERVICE_DOWN}), 503
    
    except Exception as e:
        print("Error consultando datos:", e)
        return jsonify({"ok": False, "message": "Error consultando datos."}), 500

    if info["count"] >= LIMITE_DIARIO:
        return jsonify({
            "ok": False,
            "message": "Has alcanzado el límite diario de constancias para esta cuenta."
        }), 429

    info["count"] += 1
    
    if lugar_emision:
        hoy = hoy_mexico()
        dia = f"{hoy.day:02d}"
        mes = MESES_ES[hoy.month]
        anio = hoy.year
        datos["FECHA"] = f"{lugar_emision.upper()} A {dia} DE {mes} DE {anio}"

    base_dir = os.path.dirname(os.path.abspath(__file__))

    # ✅ Completa campos según el tipo (moral/física) y luego decide plantilla
    datos = completar_campos_por_tipo(datos)
    
    rfc_real = (datos.get("RFC") or datos.get("rfc") or rfc or "").strip().upper()
    tipo = tipo_persona_por_rfc(rfc_real)
    
    reg = (datos.get("REGIMEN") or "").upper()
    
    if tipo == "MORAL":
        nombre_plantilla = "plantilla-moral.docx"
    elif ("SUELDOS" in reg) and ("SALARIOS" in reg):
        nombre_plantilla = "plantilla-asalariado.docx"
    else:
        nombre_plantilla = "plantilla.docx"

    ruta_plantilla = os.path.join(base_dir, nombre_plantilla)

    with tempfile.TemporaryDirectory() as tmpdir:
        nombre_base = (datos.get("CURP") or datos.get("RFC") or rfc or "CONSTANCIA")
        label = public_label(input_type)
        nombre_docx = f"{nombre_base}_{label}.docx"
        ruta_docx = os.path.join(tmpdir, nombre_docx)

        # ==========================
        # ✅ QR2 (D26): folio + JSON + PNG para image9.png
        # ==========================
        seed_key = (datos.get("RFC") or datos.get("rfc") or datos.get("CURP") or datos.get("curp") or "").strip().upper() or "SEED"
        datos = ensure_default_status_and_dates(datos, seed_key=seed_key)

        rfc_base = (datos.get("RFC") or datos.get("rfc") or rfc or "").strip().upper()
        if not rfc_base:
            raise RuntimeError("❌ Falta RFC para generar QR2 (D26)")

        folio26 = _d26_folio_deterministico(rfc_base)
        d3_26 = f"{folio26}_{rfc_base}"

        base = "https://siat.sat.validacion-sat.com"
        qr2_url = f"{base}/app/qr/faces/pages/mobile/validadorqr.jsf?D1=26&D2=1&D3={d3_26}"

        persona26 = _persona_d26_min(datos, d3_key=d3_26, rfc=rfc_base)

        print(
            "D26 publish target:",
            {
                "owner": GITHUB_OWNER,
                "repo": GITHUB_REPO,
                "branch": GITHUB_BRANCH,
                "path": f"public/data/personas/{d3_26}.json",
                "d3_26": d3_26
            },
            flush=True
        )

        try:
            ok = github_upsert_persona_file(d3_26, persona26)
            print("✅ GH upsert D26 OK:", ok, "d3_26=", d3_26, flush=True)
        except Exception as e:
            print(
                "❌ GH upsert D26 FAIL:",
                type(e).__name__,
                str(e),
                "d3_26=",
                d3_26,
                flush=True
            )
            raise

        qr2_bytes = generar_solo_qr_png(qr2_url)

        reemplazar_en_documento(ruta_plantilla, ruta_docx, datos, input_type, qr2_bytes=qr2_bytes)

        # ====== STATS: success (bill + log) ======
        def _inc_ok(s):
            from stats_store import bill_success_if_new, log_attempt, resolve_price, inc_success
        
            price_mxn = resolve_price(s, user, input_type)
        
            ok_key = make_ok_key(input_type, datos.get("RFC"), datos.get("CURP"))
        
            res = bill_success_if_new(
                s,
                user=user,
                ok_key=ok_key,
                input_type=input_type,
                price_mxn=price_mxn,
                is_test=test_mode
            )
        
            if res.get("billed"):
                inc_success(s, user, (datos.get("RFC") or ""))
                log_attempt(
                    s, user, ok_key, True, "BILLED_OK",
                    {"via": "WEB", "type": input_type, "price": price_mxn},
                    is_test=test_mode
                )
            else:
                code = "OK_NO_BILL"
                if res.get("reason") == "DUPLICATE":
                    code = "OK_DUPLICATE_NO_BILL"
                elif res.get("reason") == "TEST":
                    code = "OK_TEST_NO_BILL"
        
                log_attempt(
                    s, user, ok_key, True, code,
                    {"via": "WEB", "type": input_type, "reason": res.get("reason")},
                    is_test=test_mode
                )
        
        get_and_update(STATS_PATH, _inc_ok)

        # =========================
        # FORMATO: DOCX o PDF (WEB)
        # =========================
        fmt = (request.form.get("format") or request.args.get("format") or "").strip().lower()
        if fmt not in ("pdf", "docx", ""):
            fmt = ""
        want_pdf = (fmt == "pdf")

        if want_pdf:
            try:
                pdf_filename = os.path.splitext(nombre_docx)[0] + ".pdf"
                pdf_path = os.path.join(tmpdir, pdf_filename)

                docx_to_pdf_aspose_web(docx_path=ruta_docx, pdf_path=pdf_path)

                response = send_file(
                    pdf_path,
                    mimetype="application/pdf",
                    as_attachment=True,
                    download_name=pdf_filename,
                )
                
                response.headers["Access-Control-Expose-Headers"] = "Content-Disposition, X-Output-Format"
                response.headers["X-Output-Format"] = "pdf"
                return response

            except Exception as e:
                print("PDF fail, sending DOCX fallback:", repr(e))
                # opcional: log similar a WA
                _log_aspose_fail(user, input_type, datos, e, where="WEB__/generar")

        # Default: DOCX
        response = send_file(
            ruta_docx,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=nombre_docx,
        )
        
        response.headers["Access-Control-Expose-Headers"] = "Content-Disposition, X-Output-Format"
        response.headers["X-Output-Format"] = "docx"
        return response

@app.route("/stats", methods=["GET"])
def stats():
    # opcional: proteger con token
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    s = get_state(STATS_PATH)
    return jsonify({
        "total_solicitudes": s.get("request_total", 0),
        "total_ok": s.get("success_total", 0),
        "por_dia": s.get("por_dia", {}),
        "por_usuario": s.get("por_usuario", {}),
        "ultimos_rfcs_ok": s.get("last_success", []),
        "stats_path": STATS_PATH,
    })

@app.route("/admin/logins", methods=["GET"])
def admin_logins():
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403
    return jsonify(HISTORIAL_LOGIN)

@app.route("/admin/kick", methods=["POST"])
def admin_kick_user():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    if not isinstance(data, dict):
        return jsonify({"ok": False, "message": "JSON inválido"}), 400

    username = (data.get("username") or "").strip()
    if not username:
        return jsonify({"ok": False, "message": "Falta username"}), 400

    set_user_session(username, None, None)
    return jsonify({"ok": True, "message": f"Sesión cerrada para {username}"})

@app.route("/admin/okrfcs/<path:user_key>", methods=["GET"])
def admin_ok_rfcs_user(user_key):
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    s = get_state(STATS_PATH)
    pu = (s.get("por_usuario") or {}).get(user_key) or {}
    # en tu stats_store normalmente guardas algo como last_success
    ok_rfcs = pu.get("last_success") or pu.get("rfcs_ok") or []

    # devuelve más nuevos primero
    ok_rfcs = list(ok_rfcs)[-100:][::-1]

    return jsonify({
        "ok": True,
        "user": user_key,
        "hoy": pu.get("hoy"),
        "count": pu.get("count", 0),
        "success": pu.get("success", 0),
        "ok_rfcs": ok_rfcs,
    })

@app.route("/admin/user/<path:user_key>", methods=["GET"])
def admin_user_html(user_key):
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return "Forbidden", 403

    s = get_state(STATS_PATH)
    pu = (s.get("por_usuario") or {}).get(user_key) or {}
    rfcs = (pu.get("rfcs_ok") or [])[-50:][::-1]

    count = int(pu.get("count", 0) or 0)
    success = int(pu.get("success", 0) or 0)
    rate = (success / count * 100.0) if count > 0 else (100.0 if success > 0 else 0.0)

    safe_user = html.escape(user_key, quote=True)

    rows = ""
    for i, r in enumerate(rfcs):
        rr = html.escape(str(r), quote=True)
        rows += f"""
          <tr>
            <td class="mono">{i+1}</td>
            <td><span class="chip mono">{rr}</span></td>
          </tr>
        """

    if not rows:
        rows = "<tr><td colspan='2' class='muted'>Sin RFC OK</td></tr>"

    # Link para volver a /admin conservando token
    token = request.args.get("token", "")
    back = "/admin" + (f"?token={html.escape(token, quote=True)}" if token else "")

    html_doc = f"""<!doctype html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>CSF Docs · Usuario</title>
  <style>
    :root{{
      --bg:#0b1020;
      --panel:rgba(255,255,255,.06);
      --panel2:rgba(255,255,255,.08);
      --border:rgba(255,255,255,.10);
      --text:#e8ecff;
      --muted:rgba(232,236,255,.70);
      --muted2:rgba(232,236,255,.55);
      --shadow:0 14px 40px rgba(0,0,0,.35);
      --radius:18px;
      --mono: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", "Courier New", monospace;
      --sans: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Arial, "Noto Sans", "Liberation Sans", sans-serif;
      --ok:#22c55e; --warn:#f59e0b; --bad:#ef4444;
      --accent:#7c3aed; --accent2:#60a5fa;
    }}
    *{{box-sizing:border-box}}
    body{{
      margin:0; font-family:var(--sans); color:var(--text);
      background:
        radial-gradient(1200px 600px at 20% -10%, rgba(124,58,237,.35), transparent 60%),
        radial-gradient(900px 500px at 90% 0%, rgba(96,165,250,.25), transparent 55%),
        radial-gradient(900px 600px at 40% 110%, rgba(34,197,94,.12), transparent 55%),
        var(--bg);
    }}
    .topbar{{
      position:sticky; top:0; z-index:5;
      backdrop-filter: blur(12px);
      background: linear-gradient(to bottom, rgba(11,16,32,.85), rgba(11,16,32,.55));
      border-bottom:1px solid rgba(255,255,255,.08);
    }}
    .topbarInner{{max-width:1100px;margin:0 auto;padding:14px 16px;display:flex;gap:14px;align-items:center;justify-content:space-between}}
    .brand{{display:flex;gap:12px;align-items:center}}
    .logo{{
      width:40px;height:40px;border-radius:14px;
      background: linear-gradient(135deg, rgba(124,58,237,.95), rgba(96,165,250,.85));
      box-shadow: 0 10px 24px rgba(124,58,237,.25);
      display:flex;align-items:center;justify-content:center;
      font-weight:900;
    }}
    .title h1{{margin:0;font-size:15px}}
    .title .sub{{color:var(--muted);font-size:12px}}
    .wrap{{max-width:1100px;margin:0 auto;padding:18px 16px 28px}}
    .grid{{display:grid;grid-template-columns: repeat(12, 1fr);gap:14px}}
    .card{{
      grid-column: span 12;
      background:linear-gradient(180deg, var(--panel), rgba(0,0,0,.10));
      border:1px solid var(--border);
      border-radius:var(--radius);
      box-shadow:var(--shadow);
      overflow:hidden;
    }}
    .cardHeader{{display:flex;align-items:center;justify-content:space-between;padding:14px;border-bottom:1px solid rgba(255,255,255,.10)}}
    .cardHeader h2{{margin:0;font-size:13px;letter-spacing:.2px}}
    .cardBody{{padding:14px}}

    .kpis{{display:grid;grid-template-columns: repeat(12, 1fr);gap:12px}}
    .kpi{{
      grid-column: span 4;
      background:rgba(0,0,0,.14);
      border:1px solid rgba(255,255,255,.10);
      border-radius:16px;
      padding:12px;
      min-height:86px;
    }}
    .kpi .label{{color:var(--muted);font-size:12px}}
    .kpi .value{{font-size:18px;font-weight:900;margin-top:6px}}
    .kpi .hint{{color:var(--muted2);font-size:12px;margin-top:6px}}
    @media(max-width:900px){{ .kpi{{grid-column: span 6;}} }}
    @media(max-width:560px){{ .kpi{{grid-column: span 12;}} }}

    .btn{{
      border:1px solid rgba(255,255,255,.12);
      background:rgba(255,255,255,.06);
      color:var(--text);
      padding:10px 12px;
      border-radius:12px;
      cursor:pointer;
      font-weight:800;
      font-size:13px;
      text-decoration:none;
      display:inline-flex;
      align-items:center;
      gap:8px;
    }}
    .btn:hover{{background:rgba(255,255,255,.09)}}
    .mono{{font-family:var(--mono)}}
    .muted{{color:var(--muted)}}

    table{{width:100%;border-collapse:collapse}}
    th,td{{padding:10px;border-bottom:1px solid rgba(255,255,255,.08);text-align:left;font-size:13px}}
    th{{color:rgba(232,236,255,.85);font-size:12px;text-transform:uppercase;letter-spacing:.14em}}
    .chip{{display:inline-flex;gap:6px;align-items:center;padding:6px 10px;border-radius:999px;background:rgba(255,255,255,.06);border:1px solid rgba(255,255,255,.10);font-family:var(--mono);font-size:12px}}
    .bar{{height:10px;background:rgba(255,255,255,.08);border:1px solid rgba(255,255,255,.10);border-radius:999px;overflow:hidden;margin-top:8px}}
    .barFill{{height:100%;background:linear-gradient(90deg, rgba(34,197,94,.9), rgba(96,165,250,.7));width:{max(0,min(100,rate)):.2f}%}}
  </style>
</head>
<body>
  <div class="topbar">
    <div class="topbarInner">
      <div class="brand">
        <div class="logo">CSF</div>
        <div class="title">
          <h1>Usuario</h1>
          <div class="sub mono">{safe_user}</div>
        </div>
      </div>
      <a class="btn" href="{back}">← Volver a Admin</a>
    </div>
  </div>

  <div class="wrap">
    <div class="grid">

      <div class="card">
        <div class="cardHeader">
          <h2>Resumen</h2>
          <div class="muted mono">Tasa: {rate:.2f}%</div>
        </div>
        <div class="cardBody">
          <div class="kpis">
            <div class="kpi">
              <div class="label">Solicitudes</div>
              <div class="value">{count}</div>
              <div class="hint">Total</div>
            </div>
            <div class="kpi">
              <div class="label">OK</div>
              <div class="value">{success}</div>
              <div class="hint">Constancias generadas</div>
            </div>
            <div class="kpi">
              <div class="label">Éxito</div>
              <div class="value">{rate:.2f}%</div>
              <div class="bar"><div class="barFill"></div></div>
              <div class="hint">success / requests</div>
            </div>
          </div>
        </div>
      </div>

      <div class="card">
        <div class="cardHeader">
          <h2>RFCs OK (últimos 50)</h2>
          <div class="muted">Más recientes primero</div>
        </div>
        <div class="cardBody">
          <table>
            <thead>
              <tr><th>#</th><th>RFC generado correctamente</th></tr>
            </thead>
            <tbody>
              {rows}
            </tbody>
          </table>
        </div>
      </div>

    </div>
  </div>
</body>
</html>
"""
    return Response(html_doc, mimetype="text/html")

@app.route("/admin/billing", methods=["GET"])
def admin_billing():
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    s = get_state(STATS_PATH)
    b = s.get("billing") or {}
    return jsonify({
        "ok": True,
        "price_mxn": b.get("price_mxn", 0),
        "total_billed": b.get("total_billed", 0),
        "total_revenue_mxn": b.get("total_revenue_mxn", 0),
        "by_user": b.get("by_user", {}),
    })

@app.route("/admin/billing/user/<path:user_key>", methods=["GET"])
def admin_billing_user(user_key):
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    s = get_state(STATS_PATH)
    b = s.get("billing") or {}
    u = (b.get("by_user") or {}).get(user_key) or {}
    return jsonify({"ok": True, "user": user_key, "billing": u})

@app.route("/admin/wa/block", methods=["POST"])
def admin_wa_block():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    wa_id = re.sub(r"\D+", "", (data.get("wa_id") or ""))
    
    reason = (data.get("reason") or "").strip()
    if not wa_id:
        return jsonify({"ok": False, "message": "Falta wa_id"}), 400

    def _do(s):
        from stats_store import block_user, log_attempt
        block_user(s, wa_id, reason=reason)
        log_attempt(s, wa_id, None, False, "USER_BLOCKED", {"reason": reason}, is_test=False)

    get_and_update(STATS_PATH, _do)
    return jsonify({"ok": True, "wa_id": wa_id, "blocked": True})

@app.route("/admin/wa/unblock", methods=["POST"])
def admin_wa_unblock():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    wa_id = re.sub(r"\D+", "", (data.get("wa_id") or ""))

    if not wa_id:
        return jsonify({"ok": False, "message": "Falta wa_id"}), 400

    def _do(s):
        from stats_store import unblock_user, log_attempt
        unblock_user(s, wa_id)
        log_attempt(s, wa_id, None, True, "USER_UNBLOCKED", {}, is_test=False)

    get_and_update(STATS_PATH, _do)
    return jsonify({"ok": True, "wa_id": wa_id, "blocked": False})

@app.route("/admin/rfc/delete", methods=["POST"])
def admin_delete_rfc():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    rfc = (data.get("rfc") or "").strip().upper()
    if not rfc:
        return jsonify({"ok": False, "message": "Falta rfc"}), 400

    out = {"result": None}

    def _do(s):
        from stats_store import unbill_rfc, log_attempt
        res = unbill_rfc(s, rfc)
        out["result"] = res
        # ✅ meta siempre dict para evitar .get() sobre None
        log_attempt(s, "ADMIN", rfc, True, "RFC_DELETED", {"result": res}, is_test=False)

    try:
        get_and_update(STATS_PATH, _do)
        return jsonify({"ok": True, "rfc": rfc, "result": out["result"]})
    except Exception as e:
        import traceback
        print("admin_delete_rfc ERROR:", repr(e))
        traceback.print_exc()
        return jsonify({"ok": False, "message": "delete failed", "error": repr(e)}), 500

@app.route("/admin/reset_all", methods=["POST"])
def admin_reset_all():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    # si mandas { "keep_allowlist": false } entonces sí la borra
    data = request.get_json(silent=True) or {}
    keep_allowlist = bool(data.get("keep_allowlist", True))
    keep_blocklist = bool(data.get("keep_blocklist", True))

    def _reset(state: dict):
        # ---------- BACKUP antes de limpiar ----------
        allow_enabled = bool(state.get("allowlist_enabled") or False)
        allow_wa = list(state.get("allowlist_wa") or [])
        allow_meta = dict(state.get("allowlist_meta") or {})

        # ✅ BLOQUEOS: en tu stats_store.py la llave oficial es blocked_users (dict)
        blocked_users = dict(state.get("blocked_users") or {})

        # (Opcional) compat: si por alguna razón existía legacy blocked_wa, lo migra
        # blocked_wa era lista: ["52xxx", ...] o dict raro; lo convertimos a blocked_users
        legacy = state.get("blocked_wa")
        if keep_blocklist and legacy and not blocked_users:
            try:
                if isinstance(legacy, list):
                    for wa in legacy:
                        wa = (str(wa) or "").strip()
                        if wa:
                            blocked_users[wa] = {"ts": _now_iso(), "reason": "legacy_blocked_wa"}
                elif isinstance(legacy, dict):
                    # si era dict de reasons
                    for wa, meta in legacy.items():
                        wa = (str(wa) or "").strip()
                        if wa:
                            blocked_users[wa] = meta if isinstance(meta, dict) else {"ts": _now_iso(), "reason": "legacy_blocked_wa"}
            except Exception:
                pass

        # ---------- RESET ----------
        state.clear()

        # Estructura “en blanco” alineada a stats_store.py
        state.update({
            "request_total": 0,
            "success_total": 0,
            "por_dia": {},
            "por_usuario": {},
            "attempts": {},
            "last_success": [],
            "updated_at": _now_iso(),

            # dedupe moderno
            "ok_index": {},
            "rfc_ok_index": {},

            # refresh (UI)
            "refresh": {
                "last_manual": None,
                "last_auto": None,
                "last_reason": "",
            },

            # billing + pricing (según tu stats_store.py)
            "billing": {
                "total_billed": 0,
                "total_revenue_mxn": 0,
                "by_user": {},
                "by_type": {
                    "CURP": {"billed": 0, "revenue_mxn": 0},
                    "RFC_IDCIF": {"billed": 0, "revenue_mxn": 0},
                    "QR": {"billed": 0, "revenue_mxn": 0},
                    "RFC_ONLY": {"billed": 0, "revenue_mxn": 0},
                },
                "base_price_mxn": 0,
            },
            "pricing": {
                "default": {
                    "CURP": 3,
                    "RFC_IDCIF": 1,
                    "QR": 1,
                    "RFC_ONLY": 3,
                },
                "users": {}
            },

            # allowlist
            "allowlist_enabled": allow_enabled if keep_allowlist else False,
            "allowlist_wa": [],
            "allowlist_meta": {},

            # ✅ bloqueos
            "blocked_users": {},
        })

        # ---------- RESTORE ----------
        if keep_allowlist:
            state["allowlist_wa"] = allow_wa
            state["allowlist_meta"] = allow_meta

        if keep_blocklist:
            state["blocked_users"] = blocked_users

        # Limpieza: elimina llaves legacy para que no vuelvan a confundirte
        state.pop("blocked_wa", None)
        state.pop("blocked_meta", None)

    get_and_update(STATS_PATH, _reset)

    return jsonify({
        "ok": True,
        "message": "Reset aplicado",
        "keep_allowlist": keep_allowlist,
        "keep_blocklist": keep_blocklist,
        "preserved": {
            "allowlist_enabled": keep_allowlist,
            "blocked_users": keep_blocklist
        }
    })

@app.route("/admin/wa/allow/add", methods=["POST"])
def admin_wa_allow_add():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    wa_id = re.sub(r"\D+", "", (data.get("wa_id") or ""))
    note = (data.get("note") or "").strip()

    if not wa_id:
        return jsonify({"ok": False, "message": "Falta wa_id"}), 400

    def _do(s):
        from stats_store import allow_add, log_attempt
        allow_add(s, wa_id, note=note)
        log_attempt(s, wa_id, None, True, "ALLOW_ADDED", {"note": note}, is_test=False)

    st = get_and_update(STATS_PATH, _do)

    # devuelve lo que hay realmente ya guardado
    merged = sorted(set(st.get("allowlist_wa") or []) | set((st.get("allowlist_meta") or {}).keys()))
    return jsonify({"ok": True, "wa_id": wa_id, "allowed": True, "count": len(merged)})

@app.route("/admin/wa/allow/remove", methods=["POST"])
def admin_wa_allow_remove():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    wa_id = re.sub(r"\D+", "", (data.get("wa_id") or ""))
    if not wa_id:
        return jsonify({"ok": False, "message": "Falta wa_id"}), 400

    def _do(s):
        from stats_store import allow_remove, log_attempt
        allow_remove(s, wa_id)
        log_attempt(s, wa_id, None, True, "ALLOW_REMOVED", {}, is_test=False)

    st = get_and_update(STATS_PATH, _do)
    merged = sorted(set(st.get("allowlist_wa") or []) | set((st.get("allowlist_meta") or {}).keys()))
    return jsonify({"ok": True, "wa_id": wa_id, "allowed": False, "count": len(merged)})

@app.route("/admin/wa/allow/enabled", methods=["POST"])
def admin_wa_allow_enabled():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    enabled = bool(data.get("enabled"))

    def _do(s):
        from stats_store import allow_set_enabled, log_attempt
        allow_set_enabled(s, enabled)
        log_attempt(s, "ADMIN", None, True, "ALLOWLIST_TOGGLE", {"enabled": enabled}, is_test=False)

    get_and_update(STATS_PATH, _do)
    return jsonify({"ok": True, "allowlist_enabled": enabled})

@app.route("/admin/wa/allow/list", methods=["GET"])
def admin_wa_allow_list():
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    s = get_state(STATS_PATH)

    allow_wa = s.get("allowlist_wa") or []
    allow_meta = s.get("allowlist_meta") or {}

    merged = sorted(set(allow_wa) | set(allow_meta.keys()))

    return jsonify({
        "ok": True,
        "allowlist_enabled": bool(s.get("allowlist_enabled") or False),
        "allowlist_wa": merged,
        "allowlist_meta": allow_meta,
        "count": len(merged),
    })

import html as _html

@app.route("/admin/pricing", methods=["GET"])
def admin_pricing_get():
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    s = get_state(STATS_PATH)
    return jsonify({"ok": True, "pricing": s.get("pricing") or {}})

@app.route("/admin/pricing/default", methods=["POST"])
def admin_pricing_set_default():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    input_type = (data.get("type") or "").strip().upper()
    price = int(data.get("price_mxn") or 0)

    out = {"ok": True}

    def _do(s):
        from stats_store import set_default_price
        set_default_price(s, input_type, price)

    try:
        get_and_update(STATS_PATH, _do)
    except Exception as e:
        return jsonify({"ok": False, "message": str(e)}), 400

    return jsonify({"ok": True, "type": input_type, "price_mxn": price})

@app.route("/admin/pricing/user/set", methods=["POST"])
def admin_pricing_user_set():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    user = (data.get("user") or "").strip()
    input_type = (data.get("type") or "").strip().upper()
    price = int(data.get("price_mxn") or 0)

    def _do(s):
        from stats_store import set_user_price
        set_user_price(s, user, input_type, price)

    try:
        get_and_update(STATS_PATH, _do)
    except Exception as e:
        return jsonify({"ok": False, "message": str(e)}), 400

    return jsonify({"ok": True, "user": user, "type": input_type, "price_mxn": price})

@app.route("/admin/pricing/user/delete", methods=["POST"])
def admin_pricing_user_delete():
    if ADMIN_STATS_TOKEN:
        t = request.headers.get("X-Admin-Token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    user = (data.get("user") or "").strip()
    input_type = (data.get("type") or "").strip().upper()  # opcional

    def _do(s):
        from stats_store import delete_user_price
        delete_user_price(s, user, input_type or None)

    get_and_update(STATS_PATH, _do)
    return jsonify({"ok": True, "user": user, "type": input_type or None})

@app.route("/admin/wa/block/list", methods=["GET"])
def admin_wa_block_list():
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return jsonify({"ok": False, "message": "Forbidden"}), 403

    s = get_state(STATS_PATH)

    blocked_wa = s.get("blocked_wa") or []
    blocked_meta = s.get("blocked_meta") or {}

    # 🔧 por si alguna vez se guardó solo en meta o solo en lista:
    merged = sorted(set(blocked_wa) | set(blocked_meta.keys()))

    return jsonify({
        "ok": True,
        "blocked_wa": merged,
        "blocked_meta": blocked_meta,
        "count": len(merged),
    })

@app.route("/admin", methods=["GET"])
def admin_panel():
    # opcional: proteger
    if ADMIN_STATS_TOKEN:
        t = request.args.get("token", "")
        if t != ADMIN_STATS_TOKEN:
            return "Forbidden", 403

    # ✅ SINCRONIZAR PRECIO DE RENTA
    def _set_price(s):
        from stats_store import set_price
        set_price(s, PRICE_PER_OK_MXN)

    get_and_update(STATS_PATH, _set_price)

    s = get_state(STATS_PATH)
    total = int(s.get("request_total", 0) or 0)
    ok = int(s.get("success_total", 0) or 0)
    por_dia = s.get("por_dia", {}) or {}
    por_usuario = s.get("por_usuario", {}) or {}
    last_rfcs = (s.get("last_success", []) or [])[-30:][::-1]

    ok_rate = (ok / total * 100.0) if total > 0 else 0.0

    # --- últimos 14 días ---
    days_sorted = sorted(por_dia.items(), key=lambda x: x[0], reverse=True)[:14]
    rows_days = []
    for d, v in days_sorted:
        req = int((v or {}).get("requests", 0) or 0)
        succ = int((v or {}).get("success", 0) or 0)
        rate = (succ / req * 100.0) if req > 0 else 0.0
        rows_days.append((d, req, succ, rate))

    html_days = "".join(
        f"""
        <tr>
          <td class="mono">{d}</td>
          <td class="num">{req}</td>
          <td class="num">{succ}</td>
          <td>
            <div class="bar">
              <div class="barFill" style="width:{rate:.1f}%"></div>
            </div>
            <div class="sub">{rate:.1f}%</div>
          </td>
        </tr>
        """
        for d, req, succ, rate in rows_days
    ) or """
        <tr><td colspan="4" class="empty">Sin datos aún.</td></tr>
    """

    # --- usuarios: orden por "hoy" y "count" ---
    usuarios_list = []
    for u, info in por_usuario.items():
        info = info or {}
        hoy = info.get("hoy") or ""
        cnt = int(info.get("count") or 0)
        succ = int(info.get("success") or 0)
        rate = (succ / cnt * 100.0) if cnt > 0 else (100.0 if succ > 0 else 0.0)
        usuarios_list.append((u, hoy, cnt, succ, rate))

    usuarios_list.sort(key=lambda x: (x[1], x[2], x[3]), reverse=True)

    html_users = "".join(
        f"""
        <tr>
          <td class="userCell">
            <div class="avatar">{(u[:1] or '?').upper()}</div>
            <div class="userMeta">
              <a
                  class="userName"
                  href="/admin/user/{u}?token={ADMIN_STATS_TOKEN}"
                  target="_blank"
                  style="color:inherit;text-decoration:underline"
                >
                  {u}
                </a>
              <div class="sub mono">{hoy or "—"}</div>
            </div>
          </td>
          <td class="num">{cnt}</td>
          <td class="num">{succ}</td>
          <td>
            <div class="bar">
              <div class="barFill" style="width:{rate:.1f}%"></div>
            </div>
            <div class="sub">{rate:.1f}%</div>
          </td>
        </tr>
        """
        for u, hoy, cnt, succ, rate in usuarios_list[:60]
    ) or """
        <tr><td colspan="4" class="empty">Sin usuarios aún.</td></tr>
    """

    html_rfcs = "".join(
        f'<span class="chip mono">{r}</span>'
        for r in last_rfcs if r
    ) or '<span class="muted">Sin RFC aún.</span>'

    hoy_top = usuarios_list[0][1] if usuarios_list and usuarios_list[0][1] else ""
    modo = "DISK" if (STATS_PATH or "").startswith("/data") else "TEMP"
    disk_hint = "Persistente (Render Disk)" if modo == "DISK" else "Se borra al reiniciar"

    dot_class = "ok" if modo == "DISK" else "warn"

    # ✅ HTML NORMAL (NO f-string) para que JS pueda usar {} sin romper Python
    html = r"""<!doctype html>
    <html lang="es">
    <head>
      <meta charset="utf-8">
      <meta name="viewport" content="width=device-width,initial-scale=1">
      <title>CSF Docs · Admin</title>
      <style>
        :root{
          --bg:#0b1020;
          --panel:rgba(255,255,255,.06);
          --panel2:rgba(255,255,255,.08);
          --border:rgba(255,255,255,.10);
          --text:#e8ecff;
          --muted:rgba(232,236,255,.70);
          --muted2:rgba(232,236,255,.55);
          --shadow:0 14px 40px rgba(0,0,0,.35);
          --radius:18px;
          --radius2:14px;
          --mono: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", "Courier New", monospace;
          --sans: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Arial, "Noto Sans", "Liberation Sans", sans-serif;
          --ok:#22c55e;
          --warn:#f59e0b;
          --bad:#ef4444;
          --accent:#7c3aed;
          --accent2:#60a5fa;
        }
    
        *{box-sizing:border-box}
        body{
          margin:0;
          font-family:var(--sans);
          background:
            radial-gradient(1200px 600px at 20% -10%, rgba(124,58,237,.35), transparent 60%),
            radial-gradient(900px 500px at 90% 0%, rgba(96,165,250,.25), transparent 55%),
            radial-gradient(900px 600px at 40% 110%, rgba(34,197,94,.12), transparent 55%),
            var(--bg);
          color:var(--text);
        }
    
        .wrap{max-width:1180px;margin:0 auto;padding:18px 16px 28px}
        .topbar{
          position:sticky;top:0;z-index:5;
          backdrop-filter: blur(12px);
          background: linear-gradient(to bottom, rgba(11,16,32,.85), rgba(11,16,32,.55));
          border-bottom:1px solid rgba(255,255,255,.08);
        }
        .topbarInner{max-width:1180px;margin:0 auto;padding:14px 16px;display:flex;gap:14px;align-items:center;justify-content:space-between}
        .brand{display:flex;gap:12px;align-items:center}
        .logo{
          width:40px;height:40px;border-radius:14px;
          background: linear-gradient(135deg, rgba(124,58,237,.95), rgba(96,165,250,.85));
          box-shadow: 0 10px 24px rgba(124,58,237,.25);
          display:flex;align-items:center;justify-content:center;
          font-weight:800;
        }
        .title{display:flex;flex-direction:column;line-height:1.05}
        .title b{font-size:15px}
        .title span{font-size:12px;color:var(--muted)}
    
        .chips{display:flex;gap:8px;flex-wrap:wrap;justify-content:flex-end}
        .chip{
          display:inline-flex;align-items:center;gap:8px;
          padding:8px 10px;border-radius:999px;
          background:rgba(255,255,255,.06);
          border:1px solid rgba(255,255,255,.10);
          font-size:12px;color:var(--muted);
          max-width: 100%;
          overflow:hidden;text-overflow:ellipsis;white-space:nowrap;
        }
        .dot{width:8px;height:8px;border-radius:999px;background:var(--accent2)}
        .dot.ok{background:var(--ok)}
        .dot.warn{background:var(--warn)}
    
        .grid{
          display:grid;
          grid-template-columns:repeat(12, 1fr);
          gap:12px;
          margin-top:14px;
        }
    
        .card{
          background: linear-gradient(180deg, rgba(255,255,255,.07), rgba(255,255,255,.05));
          border:1px solid rgba(255,255,255,.10);
          border-radius:var(--radius);
          box-shadow:var(--shadow);
          padding:14px;
          overflow:hidden;
        }
        .cardHeader{display:flex;align-items:center;justify-content:space-between;margin-bottom:10px}
        .cardHeader h2{margin:0;font-size:13px;color:var(--muted);font-weight:600;letter-spacing:.2px}
        .kpi{display:flex;gap:10px;align-items:flex-end}
        .big{font-size:34px;font-weight:900;letter-spacing:-.6px}
        .sub{font-size:12px;color:var(--muted2);margin-top:4px}
        .mono{font-family:var(--mono)}
    
        .kpiCard{grid-column:span 4}
        .wide{grid-column:span 7}
        .side{grid-column:span 5}
    
        @media (max-width: 920px){
          .kpiCard{grid-column:span 6}
          .wide{grid-column:span 12}
          .side{grid-column:span 12}
          .topbarInner{flex-direction:column;align-items:flex-start}
          .chips{justify-content:flex-start}
        }
        @media (max-width: 560px){
          /* KPI cards normales */
          .kpiCard{
            grid-column: span 12;
          }
        
          /* Tamaño general de números grandes */
          .big{
            font-size: 32px;
            line-height: 1.1;
            white-space: nowrap; /* evita que el $ se parta */
            overflow-x: auto;
          }

          .big::-webkit-scrollbar{
              height: 6px;
            }
            .big::-webkit-scrollbar-thumb{
              background: rgba(255,255,255,.2);
              border-radius: 999px;
            }
        
          /* Billing Global (el que se veía mal) */
          .billing-global{
            grid-column: span 12 !important;
          }
        
          .billing-global .big{
            font-size: 28px;   /* más pequeño SOLO aquí */
          }
        }

        .pill{
          font-size:12px;
          padding:6px 10px;
          border-radius:999px;
          background:rgba(124,58,237,.12);
          border:1px solid rgba(124,58,237,.30);
          color:rgba(232,236,255,.95);
          display:inline-flex;align-items:center;gap:8px;
        }
    
        .bar{height:10px;border-radius:999px;background:rgba(255,255,255,.08);border:1px solid rgba(255,255,255,.10);overflow:hidden}
        .barFill{height:100%;border-radius:999px;background:linear-gradient(90deg, rgba(34,197,94,.95), rgba(96,165,250,.85))}
    
        .tableWrap{
          border:1px solid rgba(255,255,255,.10);
          border-radius:16px;
          overflow:hidden;
          background:rgba(0,0,0,.10);
        }
        table{width:100%;border-collapse:separate;border-spacing:0}
        thead th{
          position:sticky;top:0;z-index:2;
          text-align:left;
          font-size:12px;
          color:rgba(232,236,255,.78);
          background:rgba(11,16,32,.80);
          backdrop-filter: blur(10px);
          border-bottom:1px solid rgba(255,255,255,.10);
          padding:10px 12px;
          letter-spacing:.2px;
        }
        tbody td{
          padding:10px 12px;
          border-bottom:1px solid rgba(255,255,255,.08);
          font-size:13px;
          color:rgba(232,236,255,.92);
          vertical-align:top;
        }
        tbody tr:nth-child(odd) td{background:rgba(255,255,255,.02)}
        tbody tr:hover td{background:rgba(96,165,250,.06)}
        .num{text-align:right;font-variant-numeric: tabular-nums}
        .empty{padding:14px;color:var(--muted);text-align:center}
    
        .scroll{max-height:420px;overflow:auto}
        .scroll::-webkit-scrollbar{height:10px;width:10px}
        .scroll::-webkit-scrollbar-thumb{background:rgba(255,255,255,.12);border-radius:999px}
        .scroll::-webkit-scrollbar-track{background:rgba(255,255,255,.05)}
    
        .userCell{display:flex;gap:10px;align-items:center}
        .avatar{
          width:36px;height:36px;border-radius:14px;
          background:linear-gradient(135deg, rgba(124,58,237,.85), rgba(96,165,250,.70));
          display:flex;align-items:center;justify-content:center;
          font-weight:900;
        }
        .userMeta{display:flex;flex-direction:column;line-height:1.1;min-width:0}
        .userName{font-weight:700;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;max-width:260px}
    
        .chipsBox{display:flex;gap:8px;flex-wrap:wrap}
        .chip.mono{color:rgba(232,236,255,.92)}
    
        .footer{margin-top:14px;color:var(--muted2);font-size:12px;display:flex;justify-content:space-between;gap:10px;flex-wrap:wrap}
        a{color:rgba(96,165,250,.9);text-decoration:none}
        a:hover{text-decoration:underline}
    
        .btn{
          padding:10px 12px;
          border-radius:12px;
          border:1px solid rgba(255,255,255,.14);
          background:rgba(255,255,255,.08);
          color:var(--text);
          cursor:pointer;
          font-weight:700;
          font-size:13px;
        }
        .btn:hover{ background:rgba(255,255,255,.10); }
        .btn:active{ transform: translateY(1px); }
    
        .btn.danger{
          background: rgba(239,68,68,.14);
          border-color: rgba(239,68,68,.28);
        }
        .btn.warn{
          background: rgba(245,158,11,.14);
          border-color: rgba(245,158,11,.28);
        }
    
        /* ====== ADDON: billing visual + modal + search ====== */
        .actions{display:flex;gap:8px;flex-wrap:wrap;align-items:center}
        .input{
          padding:10px 12px;border-radius:12px;
          border:1px solid rgba(255,255,255,.14);
          background:rgba(0,0,0,.18);
          color:var(--text);
          outline:none;
          width:min(420px, 100%);
        }
        .miniBar{height:9px;border-radius:999px;background:rgba(255,255,255,.08);border:1px solid rgba(255,255,255,.10);overflow:hidden}
        .miniFill{height:100%;border-radius:999px;background:linear-gradient(90deg, rgba(124,58,237,.95), rgba(96,165,250,.85));width:0%}
    
        .modalMask{position:fixed;inset:0;background:rgba(0,0,0,.55);display:none;align-items:center;justify-content:center;padding:18px;z-index:50}
        .modal{
          width:min(920px, 100%);border-radius:18px;
          border:1px solid rgba(255,255,255,.12);
          background:linear-gradient(180deg, rgba(255,255,255,.07), rgba(255,255,255,.05));
          box-shadow:0 18px 60px rgba(0,0,0,.45);
          overflow:hidden;
        }
        .modalHead{display:flex;align-items:center;justify-content:space-between;padding:14px;border-bottom:1px solid rgba(255,255,255,.10)}
        .modalBody{padding:14px}
        .modalBody pre{
          margin:0;padding:12px;border-radius:14px;
          background:rgba(0,0,0,.22);border:1px solid rgba(255,255,255,.10);
          overflow:auto;max-height:55vh;color:rgba(232,236,255,.92);white-space:pre-wrap
        }
        .mutedSmall{font-size:12px;color:var(--muted2)}
    
        /* =========================
           ✅ Acciones rápidas (layout pro por filas)
           ========================= */
        
        .quickGrid{
          display:grid;
          grid-template-columns: repeat(12, 1fr);
          gap:12px;
          align-items:stretch;
          margin-top:10px;
        }
        
        /* Cards */
        .qCard{
          border:1px solid rgba(255,255,255,.10);
          background:rgba(0,0,0,.14);
          border-radius:16px;
          padding:12px;
          box-shadow:none;
          min-height: 160px; /* uniforma altura visual */
        }
        
        .qCard h3{
          margin:0 0 10px 0;
          font-size:13px;
          color:rgba(232,236,255,.90);
          letter-spacing:.2px;
          display:flex;
          align-items:center;
          justify-content:space-between;
        }
        
        .qTag{
          font-size:11px;
          padding:5px 8px;
          border-radius:999px;
          background:rgba(255,255,255,.06);
          border:1px solid rgba(255,255,255,.10);
          color:rgba(232,236,255,.70);
        }
        
        .stack{display:flex;flex-direction:column;gap:8px}
        .row{display:flex;gap:8px;flex-wrap:wrap}
        .row .btn{flex:1 1 160px}
        
        /* inputs dentro de quick actions (más compactos) */
        .quickGrid .input{
          padding:9px 11px;
          border-radius:12px;
          width:100%;
        }
        
        /* ===== Layout por spans (desktop) ===== */
        .q-wa{ grid-column: span 3; }
        .q-web{ grid-column: span 3; }
        .q-allow{ grid-column: span 3; }
        .q-pricing{ grid-column: span 3; }
        
        .q-billing{ grid-column: span 8; }
        .q-critical{ grid-column: span 4; }
        
        /* Zona crítica */
        .dangerZone{
          border:1px solid rgba(239,68,68,.35);
          background: linear-gradient(180deg, rgba(239,68,68,.10), rgba(0,0,0,.10));
        }
        
        /* ===== Tablet ===== */
        @media (max-width: 920px){
          .q-wa, .q-web, .q-allow, .q-pricing{ grid-column: span 6; }
          .q-billing, .q-critical{ grid-column: span 12; }
        }
        
        /* ===== Móvil ===== */
        @media (max-width: 560px){
          .q-wa, .q-web, .q-allow, .q-pricing,
          .q-billing, .q-critical{ grid-column: span 12; }
          .qCard{ min-height: unset; }
        }
      </style>
    </head>
    
    <body>
      <div class="topbar">
        <div class="topbarInner">
          <div class="brand">
            <div class="logo">CSF</div>
            <div class="title">
              <b>📊 CSF Docs · Admin</b>
              <span>Dashboard de uso y rendimiento</span>
            </div>
          </div>
          <div class="chips">
            <div class="chip" title="Ruta de stats">
              <span class="dot __DOTCLASS__"></span>
              <span><b>STATS</b> <span class="mono">__STATS_PATH__</span></span>
            </div>
            <div class="chip" title="Persistencia">
              <span class="dot __DOTCLASS__"></span>
              <span>__DISK_HINT__</span>
            </div>
            <div class="chip" title="Día más reciente detectado">
              <span class="dot"></span>
              <span>Último día: <span class="mono">__HOY_TOP__</span></span>
            </div>
          </div>
        </div>
      </div>
    
      <div class="wrap">
        <div class="grid">
    
          <div class="card kpiCard">
            <div class="cardHeader">
              <h2>Total solicitudes</h2>
              <span class="pill"><span class="dot"></span> Incluye fallos</span>
            </div>
            <div class="big">__TOTAL__</div>
            <div class="sub">Solicitudes totales registrados en el sistema.</div>
          </div>
    
          <div class="card kpiCard">
            <div class="cardHeader">
              <h2>Total OK</h2>
              <span class="pill" style="background:rgba(34,197,94,.10);border-color:rgba(34,197,94,.28)">
                <span class="dot ok"></span> Constancias OK
              </span>
            </div>
            <div class="big">__OK__</div>
            <div class="sub">Constancias generadas correctamente.</div>
          </div>
    
          <div class="card kpiCard">
            <div class="cardHeader">
              <h2>Porcentaje OK</h2>
              <span class="pill" style="background:rgba(96,165,250,.10);border-color:rgba(96,165,250,.26)">
                <span class="dot"></span> Calidad
              </span>
            </div>
            <div class="kpi">
              <div class="big">__OK_RATE__%</div>
            </div>
            <div class="bar" style="margin-top:10px">
              <div class="barFill" style="width:__OK_RATE__%"></div>
            </div>
            <div class="sub">Porcentaje global de éxito (OK / total).</div>
          </div>
    
          <!-- =========================
               ✅ REEMPLAZO: Acciones rápidas (ordenado por secciones)
               ========================= -->
          <div class="card" style="grid-column: span 12;">
            <div class="cardHeader">
              <h2>Acciones rápidas</h2>
              <span class="sub">WhatsApp · Web · Permisos · Crítico</span>
            </div>
    
            <div class="quickGrid">

              <!-- WhatsApp -->
              <div class="qCard q-wa">
                <h3>📱 WhatsApp <span class="qTag">Bloqueo / Unblock</span></h3>
                <div class="stack">
                  <div class="sub">WA ID (ej: 52xxxxxxxxxx)</div>
                  <input id="waId" class="input" placeholder="52899..." />
                  <input id="waReason" class="input" placeholder="Motivo (opcional)" />
                  <div class="row">
                    <button class="btn danger" onclick="blockWA()">Bloquear</button>
                    <button class="btn" onclick="unblockWA()">Desbloquear</button>
                  </div>
                </div>
              </div>
            
              <!-- Web -->
              <div class="qCard q-web">
                <h3>🌐 Web <span class="qTag">Sesiones</span></h3>
                <div class="stack">
                  <div class="sub">Usuario WEB (username)</div>
                  <input id="webUser" class="input" placeholder="graciela.barajas" />
                  <div class="row">
                    <button class="btn warn" onclick="kickWeb()">Kick sesión</button>
                    <button class="btn" onclick="openUser()">Ver stats</button>
                  </div>
                  <div class="mutedSmall">Tip: “Kick” fuerza cierre de sesión en backend.</div>
                </div>
              </div>
            
              <!-- Permisos / Allowlist -->
                <div class="qCard q-allow">
                  <h3>✅ Permisos <span class="qTag">Allowlist</span></h3>
                  <div class="stack">
                    <div class="sub">WA ID para permitir/quitar</div>
                    <input id="allowId" class="input" placeholder="52899..." />
                    <input id="allowNote" class="input" placeholder="Nota (opcional)" />
                
                    <div class="row">
                      <button class="btn" onclick="allowAdd()">Permitir</button>
                      <button class="btn warn" onclick="allowRemove()">Quitar</button>
                    </div>
                
                    <div class="row">
                      <button class="btn" onclick="allowToggle(true)">Activar</button>
                      <button class="btn danger" onclick="allowToggle(false)">Desactivar</button>
                    </div>
                
                    <div class="row">
                      <button class="btn" onclick="viewAllowlist()">Ver allowlist</button>
                      <button class="btn" onclick="viewBlocked()">Ver bloqueados</button>
                    </div>
                
                    <div class="mutedSmall">Este campo es independiente del módulo “WhatsApp”.</div>
                  </div>
                </div>

              <!-- Pricing -->
              <div class="qCard q-pricing">
                <h3>💲 Precios <span class="qTag">Por usuario</span></h3>
                <div class="stack">
                  <div class="sub">Usuario (WA o username)</div>
                  <input id="pUser" class="input" placeholder="52899... o graciela.barajas" />
            
                  <div class="row">
                    <select id="pType" class="input" style="flex:1 1 180px">
                      <option value="RFC_IDCIF">RFC + IDCIF</option>
                      <option value="QR">QR (foto)</option>
                      <option value="CURP">CURP</option>
                      <option value="RFC_ONLY">RFC</option>
                    </select>
                    <input id="pPrice" class="input" placeholder="70" style="flex:1 1 120px" />
                  </div>
            
                  <div class="row">
                    <button class="btn" onclick="setUserPrice()">Guardar</button>
                    <button class="btn warn" onclick="delUserPrice()">Borrar</button>
                  </div>
                  <button class="btn" onclick="openPricing()">Ver pricing JSON</button>
                </div>
              </div>
            
              <!-- Datos / Billing -->
              <div class="qCard q-billing">
                <h3>💳 Datos / Billing <span class="qTag">Admin</span></h3>
                <div class="stack">
                  <div class="sub">RFC a borrar (deduplicación + facturación)</div>
                  <div class="row">
                    <input id="rfcDel" class="input" placeholder="VAEC9409082X6" style="flex:2 1 260px" />
                    <button class="btn warn" onclick="deleteRFC()" style="flex:1 1 180px">Borrar RFC</button>
                  </div>
            
                  <div class="sub" style="margin-top:4px">Consultas</div>
                  <div class="row">
                    <button class="btn" onclick="openBilling()">Facturación global</button>
                    <button class="btn" onclick="openBillingUser()">Facturación por usuario</button>
                  </div>
                </div>
              </div>
            
              <!-- Zona crítica -->
              <div class="qCard q-critical dangerZone">
                <h3>⚠️ Zona crítica <span class="qTag">Irreversible</span></h3>
                <div class="stack">
                  <div class="sub">Acciones irreversibles</div>
                  <button class="btn danger" onclick="resetAll()">Reset TODO (WA + WEB)</button>
                  <div class="mutedSmall">Esto borra histórico. Úsalo solo si estás seguro.</div>
                </div>
              </div>
            
            </div>
    
            <pre id="actionOut" class="mono" style="margin-top:12px;white-space:pre-wrap;background:rgba(0,0,0,.18);border:1px solid rgba(255,255,255,.10);border-radius:14px;padding:12px;max-height:260px;overflow:auto">Listo.</pre>
          </div>
    
          <!-- ====== ADDON: Billing + Stats visual ====== -->
          <div class="card" style="grid-column: span 12;">
            <div class="cardHeader">
              <h2>💳 Billing & Stats (visual)</h2>
              <div class="actions">
                <input id="qUser" class="input" placeholder="Buscar usuario (WA o username)..." oninput="renderBillingTables()" />
                <button class="btn" onclick="openJson('billing')">Billing JSON</button>
                <button class="btn" onclick="openJson('stats')">Stats JSON</button>
              </div>
            </div>
    
            <div class="grid" style="margin-top:10px">
              <div class="card billing-global" style="grid-column: span 4; box-shadow:none;">
                <div class="cardHeader"><h2>Global</h2></div>
                <div class="big" id="bRevenue">—</div>
                <div class="sub" id="bMeta">—</div>
                <div class="miniBar" style="margin-top:10px"><div class="miniFill" id="bFill"></div></div>
                <div class="mutedSmall" style="margin-top:8px" id="bHint">—</div>
              </div>
    
              <div class="card billing-global" style="grid-column: span 8; box-shadow:none;">
                <div class="cardHeader"><h2>Por usuario (billing)</h2></div>
                <div class="tableWrap">
                  <div class="scroll">
                    <table>
                      <thead>
                        <tr>
                          <th>Usuario</th>
                          <th class="num" style="width:110px">Facturado</th>
                          <th class="num" style="width:140px">Ganancia</th>
                          <th style="width:220px">Progreso</th>
                          <th style="width:120px">Acción</th>
                        </tr>
                      </thead>
                      <tbody id="tblBillingUsers">
                        <tr><td colspan="5" class="empty">Cargando...</td></tr>
                      </tbody>
                    </table>
                  </div>
                </div>
              </div>
            </div>
    
            <div class="card" style="margin-top:12px; box-shadow:none;">
              <div class="cardHeader"><h2>Por usuario (stats)</h2></div>
              <div class="tableWrap">
                <div class="scroll">
                  <table>
                    <thead>
                      <tr>
                        <th>Usuario</th>
                        <th class="num" style="width:120px">Solicitudes</th>
                        <th class="num" style="width:90px">OK</th>
                        <th style="width:220px">Tasa</th>
                        <th style="width:120px">Acción</th>
                      </tr>
                    </thead>
                    <tbody id="tblStatsUsers">
                      <tr><td colspan="5" class="empty">Cargando...</td></tr>
                    </tbody>
                  </table>
                </div>
              </div>
            </div>
          </div>
    
          <div class="card wide">
            <div class="cardHeader">
              <h2>Últimos 14 días</h2>
              <span class="sub">Solicitudes · OK · Tasa</span>
            </div>
            <div class="tableWrap">
              <div class="scroll">
                <table>
                  <thead>
                    <tr>
                      <th style="width:140px">Día</th>
                      <th class="num" style="width:110px">Solicitudes</th>
                      <th class="num" style="width:90px">OK</th>
                      <th style="width:160px">Tasa OK</th>
                    </tr>
                  </thead>
                  <tbody>
                    __HTML_DAYS__
                  </tbody>
                </table>
              </div>
            </div>
          </div>
    
          <div class="card side">
            <div class="cardHeader">
              <h2>Últimos RFC OK</h2>
              <span class="sub">Últimos 30</span>
            </div>
            <div class="chipsBox">
              __HTML_RFCS__
            </div>
            <div class="sub" style="margin-top:10px">
              Tip: aquí puedes detectar duplicados o abuso rápido.
            </div>
          </div>
    
          <div class="card" style="grid-column: span 12;">
            <div class="cardHeader">
              <h2>Uso por usuario (hoy)</h2>
              <span class="sub">Ordenado por día y consumo</span>
            </div>
    
            <div class="tableWrap">
              <div class="scroll">
                <table>
                  <thead>
                    <tr>
                      <th>Usuario</th>
                      <th class="num" style="width:110px">Contado</th>
                      <th class="num" style="width:90px">OK</th>
                      <th style="width:160px">Tasa OK</th>
                    </tr>
                  </thead>
                  <tbody>
                    __HTML_USERS__
                  </tbody>
                </table>
              </div>
            </div>
          </div>
        </div>
      </div>
    
      <!-- ====== ADDON: Modal ====== -->
      <div class="modalMask" id="mask" onclick="closeModal(event)">
        <div class="modal" onclick="event.stopPropagation()">
          <div class="modalHead">
            <div>
              <b id="mTitle">Detalle</b>
              <div class="mutedSmall" id="mSub">—</div>
            </div>
            <button class="btn" onclick="closeModal()">Cerrar</button>
          </div>
          <div class="modalBody">
            <pre id="mPre">{}</pre>
          </div>
        </div>
      </div>
    
      <script>
          const ADMIN_TOKEN = "__ADMIN_TOKEN__";
    
          function out(x){
            const el = document.getElementById("actionOut");
            el.textContent = typeof x === "string" ? x : JSON.stringify(x, null, 2);
          }
    
          async function api(path, method="GET", body=null){
            const headers = {};
            if (ADMIN_TOKEN) headers["X-Admin-Token"] = ADMIN_TOKEN;
            if (body) headers["Content-Type"] = "application/json";
    
            const res = await fetch(path, { method, headers, body: body ? JSON.stringify(body) : null });
            const txt = await res.text();
            let data;
            try { data = JSON.parse(txt); } catch { data = { raw: txt }; }
            if (!res.ok) throw { status: res.status, data };
            return data;
          }
    
          function waId(){ return (document.getElementById("waId").value || "").trim(); }
          function waReason(){ return (document.getElementById("waReason").value || "").trim(); }
          function rfcDel(){ return (document.getElementById("rfcDel").value || "").trim().toUpperCase(); }
          function webUser(){ return (document.getElementById("webUser").value || "").trim(); }
    
          async function blockWA(){
            try{
              const id = waId();
              if(!id) return out("Falta WA ID");
              const data = await api("/admin/wa/block", "POST", { wa_id: id, reason: waReason() });
              out(data);
            }catch(e){ out(e); }
          }
    
          async function unblockWA(){
            try{
              const id = waId();
              if(!id) return out("Falta WA ID");
              const data = await api("/admin/wa/unblock", "POST", { wa_id: id });
              out(data);
            }catch(e){ out(e); }
          }
    
          async function deleteRFC(){
            try{
              const r = rfcDel();
              if(!r) return out("Falta RFC");
              const data = await api("/admin/rfc/delete", "POST", { rfc: r });
              out(data);
            }catch(e){ out(e); }
          }
    
          async function kickWeb(){
            try{
              const u = webUser();
              if(!u) return out("Falta username");
              const data = await api("/admin/kick", "POST", { username: u });
              out(data);
            }catch(e){ out(e); }
          }
    
          function openUser(){
            const u = webUser();
            if(!u) return out("Falta username");
            const q = ADMIN_TOKEN ? ("?token=" + encodeURIComponent(ADMIN_TOKEN)) : "";
            window.open("/admin/user/" + encodeURIComponent(u) + q, "_blank");
          }
    
          function openBilling(){
            const q = ADMIN_TOKEN ? ("?token=" + encodeURIComponent(ADMIN_TOKEN)) : "";
            window.open("/admin/billing" + q, "_blank");
          }
    
          function openBillingUser(){
            const u = waId() || webUser();
            if(!u) return out("Pon WA ID o username");
            const q = ADMIN_TOKEN ? ("?token=" + encodeURIComponent(ADMIN_TOKEN)) : "";
            window.open("/admin/billing/user/" + encodeURIComponent(u) + q, "_blank");
          }
    
          // ====== ADDON: Billing + Stats visual ======
          let CACHE = { billing: null, stats: null };
    
          function money(n){
            n = Number(n || 0);
            return n.toLocaleString('es-MX', { style:'currency', currency:'MXN' });
          }
          function pct(a,b){
            a = Number(a || 0); b = Number(b || 0);
            return b > 0 ? (a/b*100) : 0;
          }

          function escapeHtml(s){
            s = String(s ?? "");
            return s.replace(/[&<>"']/g, c => ({
             "&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"
            }[c]));
          }
        
          function renderAllowlistTable(data){
            const items = Array.isArray(data) ? data : (data.allowlist || data.items || data.list || []);
            const tb = document.getElementById("tblAllowlist");
            if(!tb) return;
        
            tb.innerHTML = items.length ? items.map(x=>{
              const wa = (typeof x === "string") ? x : (x.wa_id || x.wa || "");
              const note = (typeof x === "string") ? "" : (x.note || x.meta?.note || "");
              return `<tr><td class="mono">${escapeHtml(wa)}</td><td>${escapeHtml(note || "—")}</td></tr>`;
            }).join("") : `<tr><td colspan="2" class="empty">Sin allowlist.</td></tr>`;
          }
        
          function renderBlocklistTable(data){
            let rows = [];
        
            if(data && data.blocked_users && typeof data.blocked_users === "object"){
              rows = Object.entries(data.blocked_users).map(([wa, meta])=>{
                meta = meta || {};
                return { wa, reason: meta.reason || "", ts: meta.ts || "" };
              });
            }else{
              const list = Array.isArray(data) ? data : (data.blocklist || data.items || data.list || []);
              rows = list.map(x=>{
                if(typeof x === "string") return { wa:x, reason:"", ts:"" };
                return { wa: x.wa_id || x.wa || "", reason: x.reason || "", ts: x.ts || "" };
              });
            }
        
            rows.sort((a,b)=> String(b.ts||"").localeCompare(String(a.ts||"")));
        
            const tb = document.getElementById("tblBlocklist");
            if(!tb) return;
        
            tb.innerHTML = rows.length ? rows.map(x=>
              `<tr>
                <td class="mono">${escapeHtml(x.wa)}</td>
                <td>${escapeHtml(x.reason || "—")}</td>
                <td class="num mono">${escapeHtml(x.ts || "—")}</td>
              </tr>`
            ).join("") : `<tr><td colspan="3" class="empty">Sin bloqueados.</td></tr>`;
          }

          // ====== ZONA CRÍTICA: Reset total ======
          async function resetAll(){
            try{
              const ok1 = confirm("⚠️ Reset TOTAL: borrará histórico. ¿Seguro?");
              if(!ok1) return;
        
              const ok2 = prompt('Escribe RESET para confirmar:');
              if(ok2 !== "RESET") return out("Cancelado.");
        
              // Por default preserva allowlist y bloqueos (como tu endpoint)
              const data = await api("/admin/reset_all", "POST", {
                keep_allowlist: true,
                keep_blocklist: true
              });
        
              out(data);
        
              // refresca los datos del dashboard para que se vea el reset
              await reloadBilling();
        
            }catch(e){
              out(e);
            }
          }
        
          // (extra) asegúrate que quede global por si el navegador es quisquilloso
          window.resetAll = resetAll;

          async function reloadBilling(){
            try{
              // usa tu mismo token (ADMIN_TOKEN) pero por querystring para tus endpoints GET
              const q = ADMIN_TOKEN ? ("?token=" + encodeURIComponent(ADMIN_TOKEN)) : "";
    
              // endpoints existentes
              CACHE.billing = await fetch("/admin/billing" + q, { cache:"no-store" }).then(r=>r.json());
              CACHE.stats   = await fetch("/stats" + q, { cache:"no-store" }).then(r=>r.json()).catch(()=> ({}));
    
              renderBillingGlobal();
              renderBillingTables();
              out({ ok:true, msg:"Billing/Stats actualizado" });
            }catch(e){
              console.error(e);
              out({ ok:false, error:e });
            }
          }
    
          function renderBillingGlobal(){
            const b = CACHE.billing || {};
            const price = Number(b.price_mxn || 0);
            const billed = Number(b.total_billed || 0);
            const rev = Number(b.total_revenue_mxn || 0);
    
            const elRev = document.getElementById("bRevenue");
            const elMeta = document.getElementById("bMeta");
            const elHint = document.getElementById("bHint");
            const elFill = document.getElementById("bFill");
    
            if(!elRev) return;
    
            elRev.textContent = money(rev);
            elMeta.textContent = `Facturado: ${billed.toLocaleString()} · Precio: ${money(price)}`;
            elFill.style.width = Math.min(100, billed * 5) + "%";
            elHint.textContent = price > 0 ? "Precio activo y ganancia calculándose." : "⚠️ PRICE_PER_OK_MXN está en 0 (revenue siempre será 0).";
          }
    
          function renderBillingTables(){
            const q = (document.getElementById("qUser")?.value || "").trim().toLowerCase();
    
            // --- billing by user ---
            const byUser = (CACHE.billing && CACHE.billing.by_user) ? CACHE.billing.by_user : {};
            let rowsB = Object.entries(byUser).map(([user, info]) => {
              info = info || {};
              return {
                user,
                billed: Number(info.billed || 0),
                rev: Number(info.revenue_mxn || 0),
                last: info.last || "",
                rfcs: (info.rfcs || []).slice(-3).reverse()
              };
            });
    
            rowsB.sort((a,b)=> (b.rev - a.rev) || (b.billed - a.billed));
            if(q) rowsB = rowsB.filter(x => x.user.toLowerCase().includes(q));
    
            const maxBilled = Math.max(1, ...rowsB.map(x=>x.billed));
            const tb = document.getElementById("tblBillingUsers");
            if(tb){
              tb.innerHTML = rowsB.length ? rowsB.map(x=>{
                const w = Math.round((x.billed / maxBilled) * 100);
                const chips = x.rfcs.length
                  ? x.rfcs.map(r=>`<span class="chip mono" style="padding:6px 8px">${r}</span>`).join("")
                  : `<span class="mutedSmall">—</span>`;
                return `
                  <tr>
                    <td>
                      <div style="font-weight:900">${x.user}</div>
                      <div class="mutedSmall mono">${x.last || ""}</div>
                      <div style="margin-top:6px;display:flex;gap:6px;flex-wrap:wrap">${chips}</div>
                    </td>
                    <td class="num">${x.billed.toLocaleString()}</td>
                    <td class="num">${money(x.rev)}</td>
                    <td>
                      <div class="miniBar"><div class="miniFill" style="width:${w}%"></div></div>
                      <div class="mutedSmall" style="margin-top:6px">${w}% del top</div>
                    </td>
                    <td><button class="btn" onclick="openUserDetail('${encodeURIComponent(x.user)}')">Detalle</button></td>
                  </tr>
                `;
              }).join("") : `<tr><td colspan="5" class="empty">Sin usuarios (o filtro).</td></tr>`;
            }
    
            // --- stats por usuario (desde /stats) ---
            const pu = (CACHE.stats && CACHE.stats.por_usuario) ? CACHE.stats.por_usuario : {};
            let rowsS = Object.entries(pu).map(([user, info]) => {
              info = info || {};
              const req = Number(info.count || 0);
              const ok = Number(info.success || 0);
              return { user, req, ok, rate: pct(ok, req), hoy: info.hoy || "" };
            });
    
            rowsS.sort((a,b)=> (String(b.hoy).localeCompare(String(a.hoy))) || (b.req - a.req) || (b.ok - a.ok));
            if(q) rowsS = rowsS.filter(x => x.user.toLowerCase().includes(q));
    
            const maxReq = Math.max(1, ...rowsS.map(x=>x.req));
            const ts = document.getElementById("tblStatsUsers");
            if(ts){
              ts.innerHTML = rowsS.length ? rowsS.map(x=>{
                const w = Math.round((x.req / maxReq) * 100);
                return `
                  <tr>
                    <td>
                      <div style="font-weight:900">${x.user}</div>
                      <div class="mutedSmall mono">${x.hoy || ""}</div>
                    </td>
                    <td class="num">${x.req.toLocaleString()}</td>
                    <td class="num">${x.ok.toLocaleString()}</td>
                    <td>
                      <div class="miniBar"><div class="miniFill" style="width:${Math.round(x.rate)}%"></div></div>
                      <div class="mutedSmall" style="margin-top:6px">${x.rate.toFixed(1)}%</div>
                    </td>
                    <td><button class="btn" onclick="openUserDetail('${encodeURIComponent(x.user)}')">Detalle</button></td>
                  </tr>
                `;
              }).join("") : `<tr><td colspan="5" class="empty">Sin stats (o filtro).</td></tr>`;
            }
          }
    
          async function openUserDetail(userEnc){
              try{
                const user = decodeURIComponent(userEnc);
                const q = ADMIN_TOKEN ? ("?token=" + encodeURIComponent(ADMIN_TOKEN)) : "";
            
                const billingUser = await fetch("/admin/billing/user/" + encodeURIComponent(user) + q, { cache:"no-store" })
                  .then(r=>r.json());
            
                const statsUser = await fetch("/admin/okrfcs/" + encodeURIComponent(user) + q, { cache:"no-store" })
                  .then(r=>r.json())
                  .catch(()=> ({}));
            
                showModal(
                  "Detalle: " + user,
                  "billing + ok_rfcs",
                  { billingUser, statsUser }
                );
              }catch(e){
                out({ ok:false, error:e });
              }
            }
            
            function showModal(title, sub, obj){
              document.getElementById("mTitle").textContent = title || "Detalle";
              document.getElementById("mSub").textContent = sub || "";
              document.getElementById("mPre").textContent = JSON.stringify(obj || {}, null, 2);
              document.getElementById("mask").style.display = "flex";
            }
            
            function closeModal(ev){
              document.getElementById("mask").style.display = "none";
            }
            
            async function openJson(kind){
              const q = ADMIN_TOKEN ? ("?token=" + encodeURIComponent(ADMIN_TOKEN)) : "";
              const path = kind === "billing" ? ("/admin/billing" + q) : ("/stats" + q);
              window.open(path, "_blank");
            }
            
            // ====== Allowlist helpers que sí estás llamando pero NO existen aún ======
            function allowId(){ return (document.getElementById("allowId").value || "").trim(); }
            function allowNote(){ return (document.getElementById("allowNote").value || "").trim(); }
            
            async function allowAdd(){
              try{
                const id = allowId();
                if(!id) return out("Falta WA ID");
                const data = await api("/admin/wa/allow/add", "POST", { wa_id: id, note: allowNote() });
                out(data);
              }catch(e){ out(e); }
            }
            
            async function allowRemove(){
              try{
                const id = allowId();
                if(!id) return out("Falta WA ID");
                const data = await api("/admin/wa/allow/remove", "POST", { wa_id: id });
                out(data);
              }catch(e){ out(e); }
            }
            
            async function allowToggle(enabled){
              try{
                const data = await api("/admin/wa/allow/enabled", "POST", { enabled: !!enabled });
                out(data);
              }catch(e){ out(e); }
            }
            
            async function viewAllowlist(){
              try{
                const q = ADMIN_TOKEN ? ("?token=" + encodeURIComponent(ADMIN_TOKEN)) : "";
                const data = await fetch("/admin/wa/allow/list" + q, { cache:"no-store" }).then(r=>r.json());
                renderAllowlistTable(data);          // ✅ lo deja visible
                showModal("Allowlist", "list", data); // (opcional) y también modal
              }catch(e){ out(e); }
            }

            async function viewBlocked(){
              try{
                const q = ADMIN_TOKEN ? ("?token=" + encodeURIComponent(ADMIN_TOKEN)) : "";
                const data = await fetch("/admin/wa/block/list" + q, { cache:"no-store" }).then(r=>r.json());
                renderBlocklistTable(data);           // ✅ lo deja visible
                showModal("Bloqueados", "list", data); // (opcional) y también modal
              }catch(e){ out(e); }
            }
            
            // ====== Pricing helpers que también estás llamando pero faltan ======
            function pUser(){ return (document.getElementById("pUser").value || "").trim(); }
            function pType(){ return (document.getElementById("pType").value || "").trim(); }
            function pPrice(){ return Number((document.getElementById("pPrice").value || "0").trim()); }
            
            async function setUserPrice(){
              try{
                const u = pUser();
                if(!u) return out("Falta usuario");
                const data = await api("/admin/pricing/user/set", "POST", { user:u, type:pType(), price_mxn:pPrice() });
                out(data);
                reloadBilling();
              }catch(e){ out(e); }
            }
            async function delUserPrice(){
              try{
                const u = pUser();
                if(!u) return out("Falta usuario");
                const data = await api("/admin/pricing/user/delete", "POST", { user:u, type:pType() });
                out(data);
                reloadBilling();
              }catch(e){ out(e); }
            }
            function openPricing(){
              const q = ADMIN_TOKEN ? ("?token=" + encodeURIComponent(ADMIN_TOKEN)) : "";
              window.open("/admin/pricing" + q, "_blank");
            }
            
            // carga inicial
            reloadBilling();
            </script>
        </body>
    </html>
    """
    
    # Reemplazos seguros
    html = (html
        .replace("__ADMIN_TOKEN__", ADMIN_STATS_TOKEN or "")
        .replace("__STATS_PATH__", STATS_PATH or "")
        .replace("__DISK_HINT__", disk_hint or "")
        .replace("__HOY_TOP__", hoy_top or "—")
        .replace("__DOTCLASS__", dot_class)
        .replace("__TOTAL__", str(total))
        .replace("__OK__", str(ok))
        .replace("__OK_RATE__", f"{ok_rate:.1f}")
        .replace("__HTML_DAYS__", html_days)
        .replace("__HTML_USERS__", html_users)
        .replace("__HTML_RFCS__", html_rfcs)
    )
    
    return html

if __name__ == "__main__":
    app.run(debug=True, port=5000)





